import copy
import hashlib
import hmac
import io
import logging
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from typing import Any
from uuid import uuid4

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.config import settings

logger = logging.getLogger("wordreplace")


@dataclass
class ReplacedFile:
    item_id: str
    filename: str
    data: io.BytesIO
    row_idx: int
    replace_count: int


RUN_CACHE: dict[str, dict[str, Any]] = {}


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _prune_cache(now: datetime | None = None) -> None:
    now = now or _utcnow()
    expired = [
        rid for rid, item in RUN_CACHE.items()
        if item.get("expires_at") and item["expires_at"] <= now
    ]
    for rid in expired:
        RUN_CACHE.pop(rid, None)

    if len(RUN_CACHE) <= settings.run_cache_max_entries:
        return
    oldest = sorted(RUN_CACHE.items(), key=lambda x: x[1].get("created_at", now))
    for rid, _ in oldest[: max(0, len(RUN_CACHE) - settings.run_cache_max_entries)]:
        RUN_CACHE.pop(rid, None)


def clean_text(text: str) -> str:
    if not isinstance(text, str):
        return ""
    return " ".join(str(text).strip().split())


def clean_filename(filename: str) -> str:
    bad = '\\/:*?"<>|\x00\x01\x02\x03\x04\x05\x06\x07\x08\x09\x0a\x0b\x0c\x0d\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f'
    table = str.maketrans({ch: "_" for ch in bad})
    value = str(filename).translate(table).strip().strip(".")
    return value or "未命名.docx"


def precompute_replace_patterns(rules: list[tuple[str, str]], excel_row: pd.Series) -> list[tuple[str, str, str, str]]:
    patterns: list[tuple[str, str, str, str]] = []
    for old_text, col_name in rules:
        replacement = str(excel_row[col_name]).strip() if col_name in excel_row.index else ""
        cleaned = clean_text(old_text)
        if cleaned:
            patterns.append((old_text, col_name, cleaned, replacement))
    return patterns


def process_paragraph(paragraph, replace_patterns: list[tuple[str, str, str, str]]) -> int:
    para_text = paragraph.text
    if not para_text or not replace_patterns:
        return 0

    new_text = para_text
    count = 0
    for _, _, key, replacement in replace_patterns:
        hit = new_text.count(key)
        if hit:
            count += hit
            new_text = new_text.replace(key, replacement)

    if count > 0 and len(paragraph.runs) > 0:
        paragraph.runs[0].text = new_text
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ""

    return count


def replace_word_with_format(word_bytes: bytes, excel_row: pd.Series, rules: list[tuple[str, str]]) -> tuple[io.BytesIO, int]:
    doc = Document(io.BytesIO(word_bytes))
    patterns = precompute_replace_patterns(rules, excel_row)
    total_replace = 0

    for paragraph in doc.paragraphs:
        total_replace += process_paragraph(paragraph, patterns)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_replace += process_paragraph(paragraph, patterns)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output, total_replace


def merge_word_documents(replaced_files: list[ReplacedFile]) -> io.BytesIO:
    if not replaced_files:
        raise ValueError("没有可合并文档")

    main_doc = Document(io.BytesIO(replaced_files[0].data.getvalue()))
    main_body = main_doc._body._element

    for file in replaced_files[1:]:
        sub_doc = Document(io.BytesIO(file.data.getvalue()))
        page_break_para = OxmlElement("w:p")
        page_break_run = OxmlElement("w:r")
        page_break = OxmlElement("w:br")
        page_break.set(qn("w:type"), "page")
        page_break_run.append(page_break)
        page_break_para.append(page_break_run)
        main_body.append(page_break_para)

        for element in sub_doc._body._element:
            main_body.append(copy.deepcopy(element))

    output = io.BytesIO()
    main_doc.save(output)
    output.seek(0)
    return output


def _build_filename(row: pd.Series, file_name_column: str, row_idx: int) -> str:
    if file_name_column in row.index and str(row[file_name_column]).strip():
        base = str(row[file_name_column]).strip()
    else:
        base = f"文件_{row_idx + 1}"
    return clean_filename(f"{base}.docx")


def _apply_filename_options(
    file_name: str,
    seq: int,
    use_seq_prefix: bool,
    seq_format: str,
    name_prefix: str,
    name_suffix: str,
    seq_value: str = "",
) -> str:
    stem, ext = file_name.rsplit(".", 1) if "." in file_name else (file_name, "docx")
    if seq_format == "01":
        seq_text = f"{seq:02d}"
    elif seq_format == "0001":
        seq_text = f"{seq:04d}"
    elif seq_format == "1.":
        seq_text = f"{seq}."
    elif seq_format == "一":
        cn_nums = "零一二三四五六七八九"
        if seq < 10:
            seq_text = cn_nums[seq]
        elif seq < 100:
            seq_text = "十" + (cn_nums[seq % 10] if seq % 10 else "") if seq < 20 else cn_nums[seq // 10] + "十" + (cn_nums[seq % 10] if seq % 10 else "")
        else:
            seq_text = str(seq)
    else:
        seq_text = str(seq)
    prefix_value = seq_value.strip() if seq_value.strip() else seq_text
    prefix = f"{prefix_value}_" if use_seq_prefix else ""
    final_name = f"{prefix}{name_prefix}{stem}{name_suffix}.{ext}"
    return clean_filename(final_name)


def run_replace(
    word_bytes: bytes,
    excel_bytes: bytes,
    rules: list[dict[str, str]],
    start_row: int,
    end_row: int,
    file_name_column: str,
    use_seq_prefix: bool = False,
    seq_format: str = "1",
    seq_column: str = "",
    name_prefix: str = "",
    name_suffix: str = "",
) -> dict[str, Any]:
    excel_df = pd.read_excel(io.BytesIO(excel_bytes)).fillna("")
    excel_df.columns = [str(c).strip() for c in excel_df.columns]

    max_row = len(excel_df)
    start_idx = max(0, start_row - 1)
    end_idx = min(end_row - 1, max_row - 1)
    if start_idx > end_idx:
        return {"files": [], "total": 0, "success": 0, "failed": 0, "replacements": 0}

    rule_pairs = [(r["keyword"], r["excel_column"]) for r in rules]
    out_files: list[ReplacedFile] = []
    total_replace = 0
    details: list[dict[str, Any]] = []

    for seq, row_idx in enumerate(range(start_idx, end_idx + 1), start=1):
        row = excel_df.iloc[row_idx]
        base_name = _build_filename(row, file_name_column, row_idx)
        seq_value = str(row.get(seq_column, "")).strip() if seq_column and seq_column in excel_df.columns else ""
        filename = _apply_filename_options(
            file_name=base_name,
            seq=seq,
            use_seq_prefix=use_seq_prefix,
            seq_format=seq_format,
            name_prefix=name_prefix,
            name_suffix=name_suffix,
            seq_value=seq_value,
        )
        try:
            out, cnt = replace_word_with_format(word_bytes, row, rule_pairs)
            item_id = uuid4().hex[:8]
            out_files.append(ReplacedFile(item_id=item_id, filename=filename, data=out, row_idx=row_idx, replace_count=cnt))
            total_replace += cnt
            details.append(
                {
                    "item_id": item_id,
                    "seq": seq,
                    "row_number": row_idx + 1,
                    "file_name": filename,
                    "status": "成功",
                    "replace_count": cnt,
                    "message": "",
                }
            )
        except Exception as exc:
            logger.warning("replace row failed: row_idx=%s err=%s", row_idx, exc)
            details.append(
                {
                    "item_id": uuid4().hex[:8],
                    "seq": seq,
                    "row_number": row_idx + 1,
                    "file_name": filename,
                    "status": "失败",
                    "replace_count": 0,
                    "message": str(exc),
                }
            )
            continue

    return {
        "files": out_files,
        "total": end_idx - start_idx + 1,
        "success": len(out_files),
        "failed": (end_idx - start_idx + 1) - len(out_files),
        "replacements": total_replace,
        "details": details,
    }


def cache_run(result: dict[str, Any]) -> str:
    _prune_cache()
    run_id = uuid4().hex[:12]
    now = _utcnow()
    RUN_CACHE[run_id] = {
        **result,
        "created_at": now,
        "expires_at": now + timedelta(seconds=settings.run_cache_ttl_seconds),
    }
    return run_id


def get_run(run_id: str) -> dict[str, Any] | None:
    _prune_cache()
    return RUN_CACHE.get(run_id)


def get_run_file(run_id: str, item_id: str) -> ReplacedFile | None:
    run = get_run(run_id)
    if not run:
        return None
    for f in run.get("files", []):
        if f.item_id == item_id:
            return f
    return None


def delete_run_file(run_id: str, item_id: str) -> dict[str, Any] | None:
    run = get_run(run_id)
    if not run:
        return None
    original_len = len(run.get("files", []))
    run["files"] = [f for f in run.get("files", []) if f.item_id != item_id]
    if len(run["files"]) == original_len:
        return run
    new_details: list[dict[str, Any]] = []
    seq = 1
    for d in run.get("details", []):
        if d.get("item_id") == item_id:
            continue
        d["seq"] = seq
        seq += 1
        new_details.append(d)
    run["details"] = new_details
    run["total"] = len(new_details)
    run["success"] = sum(1 for d in new_details if d.get("status") == "成功")
    run["failed"] = run["total"] - run["success"]
    run["replacements"] = sum(int(d.get("replace_count", 0)) for d in new_details)
    return run


def sign_export_token(run_id: str) -> str:
    digest = hmac.new(
        settings.export_token_secret.encode("utf-8"),
        run_id.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return digest


def verify_export_token(run_id: str, token: str) -> bool:
    expected = sign_export_token(run_id)
    return hmac.compare_digest(expected, token)


def export_zip(files: list[ReplacedFile]) -> io.BytesIO:
    zip_buffer = io.BytesIO()
    used: dict[str, int] = defaultdict(int)
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            name = f.filename
            used[name] += 1
            if used[name] > 1:
                stem, ext = name.rsplit(".", 1)
                name = f"{stem}_{used[name]-1}.{ext}"
            zf.writestr(name, f.data.getvalue())
    zip_buffer.seek(0)
    return zip_buffer
