import io
from dataclasses import dataclass

import pandas as pd
from docx import Document

from app.services import (
    clean_excel_types,
    get_replace_params,
    export_statistics_to_csv,
    merge_word_documents,
    replace_word_with_format,
)


@dataclass
class DummyFile:
    filename: str
    data: io.BytesIO
    row_idx: int
    log: str
    replace_count: int


def test_clean_excel_types_normalizes_columns_and_values():
    df = pd.DataFrame({1: [' a ', None], 'b': [' x ', ' y ']})
    out = clean_excel_types(df)
    assert '1' in out.columns
    assert out.iloc[0]['1'] == 'a'
    assert out.iloc[1]['1'] == ''


def test_get_replace_params_contains_rule_hash_and_counts():
    df = pd.DataFrame({'a': [1, 2]})
    params = get_replace_params(None, df, 1, 2, 'a', 'P-', '', [('k', 'a')])
    assert params['excel_rows'] == 2
    assert params['rule_count'] == 1
    assert 'rule_hash' in params
    assert isinstance(params['rule_hash'], str)
    assert len(params['rule_hash']) == 16


def test_get_replace_params_rule_hash_is_stable():
    df = pd.DataFrame({'a': [1]})
    rules = [('【姓名】', 'a')]
    p1 = get_replace_params(None, df, 1, 1, 'a', '', '', rules)
    p2 = get_replace_params(None, df, 1, 1, 'a', '', '', rules)
    assert p1['rule_hash'] == p2['rule_hash']


def test_export_statistics_to_csv_has_headers():
    files = [DummyFile('a.docx', io.BytesIO(b'abc'), 0, 'ok', 2)]
    csv_text = export_statistics_to_csv(files)
    assert '文件名' in csv_text
    assert 'a.docx' in csv_text


class UploadLike:
    def __init__(self, data: bytes, name: str = "template.docx"):
        self._data = data
        self.name = name

    def getvalue(self):
        return self._data


class ReplacedFileLike:
    def __init__(self, filename: str, data: io.BytesIO):
        self.filename = filename
        self.data = data


def test_replace_word_with_format_replaces_across_runs():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("您好，")
    p.add_run("【姓名】")
    p.add_run("，欢迎")
    source = io.BytesIO()
    doc.save(source)
    upload = UploadLike(source.getvalue())

    row = pd.Series({'姓名': '张三'})
    out_file, log, count = replace_word_with_format(upload, row, [('【姓名】', '姓名')], '替换完整关键词')
    out_doc = Document(io.BytesIO(out_file.getvalue()))
    assert "张三" in out_doc.paragraphs[0].text
    assert count == 1
    assert "✓" in log


def test_replace_word_with_format_replaces_table_cell_text():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "部门：〔部门〕"
    source = io.BytesIO()
    doc.save(source)
    upload = UploadLike(source.getvalue())

    row = pd.Series({'部门': '研发'})
    out_file, _, count = replace_word_with_format(upload, row, [('〔部门〕', '部门')], '仅替换括号内内容')
    out_doc = Document(io.BytesIO(out_file.getvalue()))
    assert "〔研发〕" in out_doc.tables[0].cell(0, 0).text
    assert count == 1


def test_merge_word_documents_does_not_insert_blank_top_paragraph():
    doc1 = Document()
    doc1.add_paragraph("第一页内容")
    b1 = io.BytesIO()
    doc1.save(b1)

    doc2 = Document()
    doc2.add_paragraph("第二页首行")
    b2 = io.BytesIO()
    doc2.save(b2)

    merged = merge_word_documents([
        ReplacedFileLike("a.docx", io.BytesIO(b1.getvalue())),
        ReplacedFileLike("b.docx", io.BytesIO(b2.getvalue())),
    ])

    out_doc = Document(io.BytesIO(merged.getvalue()))
    non_empty = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert "第一页内容" in non_empty
    assert "第二页首行" in non_empty
