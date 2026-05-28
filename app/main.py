"""
Word+Excel批量替换工具 v1.6.3
功能：Word模板与Excel数据批量替换，保留格式，支持合并导出
特性：规范的缓存管理、高性能预览、全面Bug修复
"""

# ==================== 导入库 ====================
import os
import warnings
import json
import io
import zipfile
import re
import unicodedata
import copy
from datetime import datetime
import hashlib
import logging

# 数据处理库
import streamlit as st
import pandas as pd

# Word处理库
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 数据结构和类型提示
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple, Set
from collections import defaultdict
try:
    from core_utils import (
    clean_text,
    clean_filename,
    sanitize_cache_key,
    generate_safe_filename,
    get_replace_blockers,
    dedupe_filename,
)
except ImportError:
    from app.core_utils import (
        clean_text,
        clean_filename,
        sanitize_cache_key,
        generate_safe_filename,
        get_replace_blockers,
        dedupe_filename,
    )

try:
    from services import (
    replace_word_with_format,
    merge_word_documents,
    get_replace_params,
    clean_excel_types,
    export_statistics_to_csv,
    get_keyword_statistics,
)
except ImportError:
    from app.services import (
        replace_word_with_format,
        merge_word_documents,
        get_replace_params,
        clean_excel_types,
        export_statistics_to_csv,
        get_keyword_statistics,
    )

# ==================== 配置和常量 ====================

VERSION = "v1.6.3"

# 页面配置常量
PAGE_SIZE = 10
WIDGET_HEIGHT = 250
PREVIEW_ROWS = 50
MAX_FILENAME_LENGTH = 200
MAX_WORD_FILE_SIZE = 50 * 1024 * 1024
MAX_EXCEL_FILE_SIZE = 50 * 1024 * 1024
MAX_HISTORY_ITEMS = 30
MAX_RULE_IMPORT_SIZE = 2 * 1024 * 1024
MAX_RULE_CACHE_ITEMS = 1000
MAX_EXPORT_FILES = 5000

# ===== 缓存目录管理 =====
# 获取用户的本地缓存目录（跨平台兼容）
if os.name == 'nt':  # Windows
    CACHE_BASE_DIR = os.path.join(os.environ.get('APPDATA', ''), 'BatchReplacer')
else:  # Linux/Mac
    CACHE_BASE_DIR = os.path.expanduser('~/.cache/batch_replacer')

# 创建缓存子目录
CACHE_RULES_DIR = os.path.join(CACHE_BASE_DIR, 'rules')  # 规则缓存目录
CACHE_HISTORY_DIR = os.path.join(CACHE_BASE_DIR, 'history')  # 历史记录目录
CACHE_TEMP_DIR = os.path.join(CACHE_BASE_DIR, 'temp')  # 临时文件目录

# 历史记录文件（放在缓存目录）
HISTORY_FILE = os.path.join(CACHE_HISTORY_DIR, 'operation_history.json')

# 规范化缓存目录结构
for directory in [CACHE_BASE_DIR, CACHE_RULES_DIR, CACHE_HISTORY_DIR, CACHE_TEMP_DIR]:
    if not os.path.exists(directory):
        os.makedirs(directory, exist_ok=True)

# 过滤警告消息
warnings.filterwarnings("ignore", category=UserWarning)

logger = logging.getLogger("wordreplace")

# 环境变量配置
os.environ["STREAMLIT_VERSION"] = "1.52.2"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# ==================== Streamlit页面配置 ====================
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 全局样式 ====================
st.markdown("""
<style>
    /* ===== 设计令牌 ===== */
    :root {
        --wr-bg: #f5f7fb;
        --wr-card: #ffffff;
        --wr-ink: #1f2937;
        --wr-muted: #64748b;
        --wr-line: #e2e8f0;
        --wr-brand: #0f766e;
        --wr-brand-soft: #ccfbf1;
        --wr-warn: #b45309;
        --wr-shadow: 0 10px 30px rgba(2, 6, 23, 0.08);
    }

    /* ===== 全局间距优化 ===== */
    .main {
        padding: 0.65rem 1rem !important;
        background:
            radial-gradient(circle at 15% 5%, #ecfeff 0%, transparent 38%),
            radial-gradient(circle at 85% 0%, #dcfce7 0%, transparent 30%),
            var(--wr-bg);
        color: var(--wr-ink);
    }

    [data-testid="stMainBlockContainer"] {
        padding-top: 0.5rem !important;
        padding-bottom: 0.5rem !important;
    }

    /* 块容器紧凑 */
    .stContainer {
        padding: 0.8rem !important;
        margin-bottom: 0.5rem !important;
        border-radius: 10px;
        background-color: var(--wr-card);
        border: 1px solid var(--wr-line);
        box-shadow: var(--wr-shadow);
    }

    /* 删除元素间多余间距 */
    .element-container {
        margin-bottom: 0.3rem !important;
    }

    .stColumn {
        gap: 0.5rem !important;
    }

    /* ===== 按钮样式 ===== */
    .stButton > button {
        border-radius: 8px;
        font-weight: 600;
        padding: 0.4rem 0.8rem !important;
        font-size: 13px !important;
        margin-bottom: 0.2rem !important;
        border: 1px solid var(--wr-line);
    }

    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 8px 18px rgba(15, 23, 42, 0.12);
    }

    /* ===== 输入框样式 ===== */
    .stTextInput, .stTextArea, .stSelectbox, .stNumberInput {
        margin-bottom: 0.3rem !important;
    }

    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select,
    .stNumberInput > div > div > input {
        border-radius: 5px;
        border: 1px solid var(--wr-line);
        font-size: 13px;
        padding: 0.5rem !important;
    }

    /* ===== 标题和文字样式 ===== */
    h1 {
        padding-bottom: 0.5rem;
        border-bottom: 2px solid var(--wr-brand);
        margin-bottom: 0.5rem;
        line-height: 1.2;
        letter-spacing: 0.01em;
    }

    h2 {
        margin-top: 0.5rem;
        margin-bottom: 0.3rem;
        color: var(--wr-brand);
        font-size: 1.2rem;
    }

    h3 {
        margin-top: 0.3rem;
        margin-bottom: 0.2rem;
        color: #333;
        font-size: 1.05rem;
    }

    .stSubheader {
        margin-bottom: 0.5rem !important;
        padding-bottom: 0.3rem;
        border-bottom: 1.5px solid var(--wr-line);
        font-size: 1.1rem !important;
    }

    /* ===== 展开器样式 ===== */
    .streamlit-expander {
        margin-bottom: 0.3rem !important;
        border-radius: 5px;
    }

    /* ===== 标签页样式 ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1px;
        margin-bottom: 0.5rem;
    }

    .stTabs [data-baseweb="tab"] {
        height: 40px;
        padding-top: 8px;
        border-radius: 8px 8px 0 0;
        font-size: 13px;
    }

    /* ===== 数据框样式 ===== */
    div[data-testid="stDataFrame"] {
        border-radius: 5px;
        border: 1px solid #e0e0e0;
        font-size: 12px;
    }

    /* ===== 指标卡样式 ===== */
    .metric-container {
        background-color: #f8fafc;
        padding: 0.5rem !important;
        border-radius: 5px;
        border-left: 3px solid var(--wr-brand);
        margin-bottom: 0.3rem;
    }

    /* ===== 信息框样式 ===== */
    .stats-box, .success-box, .warning-box, .error-box {
        padding: 0.6rem !important;
        margin: 0.3rem 0 !important;
        border-radius: 5px;
        border-left-width: 3px;
        font-size: 13px;
    }

    .stats-box {
        background-color: #f0fdfa;
        border-left-color: #14b8a6;
    }

    .success-box {
        background-color: #f0fdf4;
        border-left-color: #22c55e;
    }

    .warning-box {
        background-color: #fffbeb;
        border-left-color: #f59e0b;
    }

    .error-box {
        background-color: #fef2f2;
        border-left-color: #ef4444;
    }

    /* ===== 分隔线 ===== */
    hr {
        margin: 0.5rem 0 !important;
        border: none;
        border-top: 1px solid #e0e0e0;
    }

    /* ===== 无线电和复选框样式 ===== */
    .stRadio, .stCheckbox {
        margin-bottom: 0.2rem !important;
    }

    .stRadio > label, .stCheckbox > label {
        margin-bottom: 0.2rem !important;
        font-size: 13px;
    }

    /* ===== 文件上传器样式 ===== */
    .stFileUploader {
        margin-bottom: 0.3rem !important;
    }

    /* ===== 进度条样式 ===== */
    .stProgress {
        margin-bottom: 0.3rem !important;
    }

    /* ===== 表格样式 ===== */
    table {
        font-size: 12px !important;
    }

    td, th {
        padding: 0.4rem !important;
    }

    /* ===== 悬浮提示样式 ===== */
    .help-icon {
        display: inline-block;
        margin-left: 6px;
        color: #0ea5e9;
        font-weight: bold;
        cursor: help;
        position: relative;
    }

    .help-icon:hover {
        color: #0284c7;
    }

    /* 工具提示样式 */
    .tooltip {
        position: relative;
        display: inline-block;
        cursor: help;
    }

    .tooltip .tooltiptext {
        visibility: hidden;
        width: 220px;
        background-color: #0f172a;
        color: #fff;
        text-align: left;
        border-radius: 6px;
        padding: 10px;
        position: absolute;
        z-index: 1000;
        bottom: 125%;
        left: 50%;
        margin-left: -110px;
        opacity: 0;
        transition: opacity 0.3s;
        font-size: 12px;
        line-height: 1.5;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        border: 1px solid #334155;
    }

    .tooltip .tooltiptext::after {
        content: "";
        position: absolute;
        top: 100%;
        left: 50%;
        margin-left: -5px;
        border-width: 5px;
        border-style: solid;
        border-color: #0f172a transparent transparent transparent;
    }

    .tooltip:hover .tooltiptext {
        visibility: visible;
        opacity: 1;
    }

    /* Excel预览滚动容器 */
    .excel-preview-container {
        height: 350px;
        overflow-y: auto;
        border: 1px solid #e0e0e0;
        border-radius: 6px;
        background-color: #f9f9f9;
    }

    /* 顶部流程与上传卡片 */
    .wr-subtitle {
        margin: 0.25rem 0 0.4rem 0;
        color: var(--wr-muted);
        font-size: 0.95rem;
    }

    .wr-step {
        display: block;
        width: 100%;
        text-align: center;
        padding: 0.42rem 0.25rem;
        border-radius: 999px;
        border: 1px solid var(--wr-line);
        background: #f8fafc;
        color: #475569;
        font-size: 0.78rem;
        font-weight: 600;
        white-space: nowrap;
    }

    .wr-step.active {
        border-color: var(--wr-brand);
        color: #0f766e;
        background: var(--wr-brand-soft);
    }

    .wr-upload-card {
        border: 1px solid var(--wr-line);
        border-radius: 10px;
        padding: 0.55rem 0.65rem;
        background: #fcfefe;
        min-height: 130px;
    }

    .wr-upload-help {
        color: var(--wr-muted);
        font-size: 0.78rem;
        margin-top: 0.2rem;
        margin-bottom: 0.35rem;
    }

    .wr-rule-toolbar {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 0.4rem;
        padding: 0.35rem 0.45rem;
        border: 1px solid var(--wr-line);
        border-radius: 8px;
        background: #f8fafc;
        font-size: 0.8rem;
        color: #475569;
    }

    .wr-rule-item {
        display: block;
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
        font-size: 0.82rem;
        color: #1f2937;
    }

    /* 响应式设计 */
    @media (max-width: 768px) {
        .main {
            padding: 0.3rem 0.5rem;
        }
        .stContainer {
            padding: 0.5rem;
        }
        .tooltip .tooltiptext {
            width: 180px;
            margin-left: -90px;
        }
    }
</style>
""", unsafe_allow_html=True)

# ==================== 帮助提示文本定义 ====================

HELP_TEXTS = {
    "word_upload": "上传包含要替换内容的Word文件(.docx格式，不支持.doc)",
    "excel_upload": "上传包含替换数据的Excel文件(.xlsx格式，默认读取第一个工作表)",
    "replace_scope": "选择替换模式：完整关键词直接替换，括号内容只替换括号里的文字",
    "file_name_col": "选择Excel中的列用于生成文件名，通常选择唯一标识符列",
    "start_row": "从第几行开始处理替换",
    "end_row": "处理到第几行（包括该行），默认到最后一行",
    "file_prefix": "为生成的文件名添加前缀，如'2024-'会生成'2024-文件名.docx'",
    "new_keyword": "从Word预览中复制要替换的关键字，如【姓名】、（部门）",
    "new_column": "选择Excel中对应的列，这一列的数据会替换关键字",
    "add_rule": "点击添加规则，规则添加成功后即可开始替换",
    "start_replace": "开始执行批量替换操作，需要：1.选择文件 2.添加规则 3.设置行范围",
    "export_zip": "将所有替换后的文件保存为一个ZIP压缩包，便于统一下载",
    "export_merge": "将所有替换后的文件合并为一个Word文档，每个文件占一页",
    "export_stats": "导出替换统计数据为CSV格式，包含文件名、行号、替换次数等",
    "export_log": "导出详细的替换操作日志为TXT文件，记录每一行的替换情况",
    "rule_list": "显示已添加的所有替换规则，可以删除不需要的规则或撤销操作",
    "rule_import": "从之前导出的JSON文件中导入替换规则",
    "rule_export": "将当前规则导出为JSON文件，可以在其他电脑导入使用",
    "rule_cache": "快速保存规则到本地缓存，下次可以快速加载使用",
    "undo": "撤销最后一次规则操作（添加、删除等）",
    "clear_rules": "清空所有已添加的替换规则",
    "single_download": "下载单个文件到本地",
    "single_log": "查看该文件的详细替换日志",
}


# ==================== 帮助提示组件 ====================

def create_tooltip(text: str, help_key: str = "") -> str:
    """
    创建带有悬浮提示的HTML组件

    Args:
        text: 显示的标签文本
        help_key: 帮助文本的键

    Returns:
        HTML字符串
    """
    help_text = HELP_TEXTS.get(help_key, "")
    if not help_text:
        return text

    # 清理帮助文本中的引号，避免HTML冲突
    help_text = help_text.replace('"', '&quot;').replace("'", "&#39;")

    html = f"""
    <span class="tooltip" style="display: inline-flex; align-items: center; gap: 4px;">
        <span>{text}</span>
        <span class="help-icon" title="点击查看帮助">ℹ️</span>
        <span class="tooltiptext">{help_text}</span>
    </span>
    """
    return html


def format_file_size(size_bytes: int) -> str:
    """
    格式化文件大小为可读的字符串

    Args:
        size_bytes: 文件大小（字节）

    Returns:
        格式化后的文件大小字符串（如 1.23MB）
    """
    if size_bytes == 0:
        return "0B"

    size_names = ("B", "KB", "MB", "GB")
    i = int(0)
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1

    return f"{size_bytes:.2f}{size_names[i]}"


def get_cache_info() -> Dict:
    """
    获取缓存目录信息（包括大小和文件数）

    Returns:
        缓存信息字典
    """
    info = {
        "rules_count": 0,
        "history_count": 0,
        "total_size": 0,
        "rules_dir": CACHE_RULES_DIR,
        "history_file": HISTORY_FILE,
    }

    # 统计规则缓存
    try:
        if os.path.exists(CACHE_RULES_DIR):
            files = [f for f in os.listdir(CACHE_RULES_DIR) if f.endswith('.json')]
            info["rules_count"] = len(files)
            for f in files:
                file_path = os.path.join(CACHE_RULES_DIR, f)
                info["total_size"] += os.path.getsize(file_path)
    except Exception as e:
        logger.warning(f"读取规则缓存统计失败: {e}")

    # 统计历史记录
    try:
        if os.path.exists(HISTORY_FILE):
            info["history_count"] = 1
            info["total_size"] += os.path.getsize(HISTORY_FILE)
    except Exception as e:
        logger.warning(f"读取历史缓存统计失败: {e}")

    return info


# ==================== 数据结构定义 ====================

@dataclass
class ReplacedFile:
    """存储替换后的文件数据结构"""
    filename: str
    data: io.BytesIO
    row_idx: int
    log: str
    replace_count: int = 0


@dataclass
class HistoryRecord:
    """历史记录数据结构"""
    timestamp: str
    word_file: str
    excel_file: str
    rules_count: int
    files_generated: int
    status: str


# ==================== 缓存管理器 ====================

class CacheManager:
    """
    管理替换规则的缓存

    缓存文件结构：
    ~/.cache/batch_replacer/
    ├── rules/                    # 规则缓存目录
    │   ├── rule_20240115_1430.json
    │   ├── rule_20240115_1435.json
    │   └── ...
    ├── history/                  # 历史记录目录
    │   └── operation_history.json
    └── temp/                      # 临时文件目录
    """

    def __init__(self):
        """初始化缓存管理器，确保缓存目录存在"""
        self.rules_dir = CACHE_RULES_DIR
        self.history_dir = CACHE_HISTORY_DIR
        self.temp_dir = CACHE_TEMP_DIR

        # 确保目录存在
        for directory in [self.rules_dir, self.history_dir, self.temp_dir]:
            if not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)

    def save_rules(self, rules: List[Tuple[str, str]], filename: str = None) -> bool:
        """
        保存规则到缓存文件

        Args:
            rules: 规则列表
            filename: 自定义文件名（不包含扩展名）。如果为None，使用时间戳

        Returns:
            是否保存成功
        """
        try:
            # 生成规范的文件名
            if filename is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = sanitize_cache_key(f"rule_{timestamp}")

            rules_data = [{"keyword": old, "excel_column": col} for old, col in rules]
            cache_file = os.path.join(self.rules_dir, f"{sanitize_cache_key(filename)}.json")

            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(rules_data, f, ensure_ascii=False, indent=2)

            return True
        except Exception as e:
            st.warning(f"⚠️ 保存规则缓存失败：{str(e)[:50]}", icon="⚠️")
            return False

    def load_rules(self, filename: str) -> List[Tuple[str, str]]:
        """
        从缓存文件加载规则

        Args:
            filename: 缓存文件名（不包含扩展名）

        Returns:
            规则列表
        """
        try:
            cache_file = os.path.join(self.rules_dir, f"{sanitize_cache_key(filename)}.json")
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    rules_data = json.load(f)
                    return [(r["keyword"], r["excel_column"]) for r in rules_data]
        except Exception as e:
            logger.warning(f"加载规则缓存失败: {e}")
        return []

    def get_cached_rules_list(self) -> List[str]:
        """
        获取所有缓存的规则文件列表（按时间降序排列）

        Returns:
            规则文件名列表（最近30个）
        """
        try:
            if os.path.exists(self.rules_dir):
                files = [f.replace('.json', '') for f in os.listdir(self.rules_dir) if f.endswith('.json')]
                # 按修改时间降序排列
                files_with_time = []
                for f in files:
                    full_path = os.path.join(self.rules_dir, f"{f}.json")
                    mtime = os.path.getmtime(full_path)
                    files_with_time.append((f, mtime))

                files_with_time.sort(key=lambda x: x[1], reverse=True)
                return [f[0] for f in files_with_time[:30]]
        except Exception as e:
            logger.warning(f"读取规则缓存列表失败: {e}")
        return []

    def delete_rule(self, filename: str) -> bool:
        """
        删除缓存的规则文件

        Args:
            filename: 规则文件名（不包含扩展名）

        Returns:
            是否删除成功
        """
        try:
            cache_file = os.path.join(self.rules_dir, f"{sanitize_cache_key(filename)}.json")
            if os.path.exists(cache_file):
                os.remove(cache_file)
                return True
        except Exception as e:
            logger.warning(f"删除规则缓存失败: {e}")
        return False

    def get_rule_info(self, filename: str) -> Dict:
        """
        获取规则文件信息

        Args:
            filename: 规则文件名（不包含扩展名）

        Returns:
            规则文件信息字典（创建时间、大小、规则数）
        """
        try:
            cache_file = os.path.join(self.rules_dir, f"{sanitize_cache_key(filename)}.json")
            if os.path.exists(cache_file):
                stat = os.stat(cache_file)
                with open(cache_file, 'r', encoding='utf-8') as f:
                    rules_data = json.load(f)

                return {
                    "filename": filename,
                    "size": format_file_size(stat.st_size),
                    "rules_count": len(rules_data),
                    "mtime": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                }
        except Exception as e:
            logger.warning(f"读取规则缓存信息失败: {e}")
        return None

    def clear_all_cache(self) -> bool:
        """
        清除所有缓存文件

        Returns:
            是否清除成功
        """
        try:
            for directory in [self.rules_dir, self.history_dir, self.temp_dir]:
                if os.path.exists(directory):
                    for file in os.listdir(directory):
                        file_path = os.path.join(directory, file)
                        if os.path.isfile(file_path):
                            os.remove(file_path)
            return True
        except Exception as e:
            logger.warning(f"清理缓存失败: {e}")
        return False


# ==================== 历史记录管理器 ====================

class HistoryManager:
    """管理操作历史记录（保存到缓存目录）"""

    def __init__(self):
        """初始化历史记录管理器"""
        self.history_file = HISTORY_FILE
        # 确保目录存在
        os.makedirs(os.path.dirname(self.history_file), exist_ok=True)

    def add_record(self, record: HistoryRecord):
        """添加操作记录到历史"""
        try:
            history = self.load_history()
            history.insert(0, {
                "timestamp": record.timestamp,
                "word_file": record.word_file,
                "excel_file": record.excel_file,
                "rules_count": record.rules_count,
                "files_generated": record.files_generated,
                "status": record.status
            })
            history = history[:MAX_HISTORY_ITEMS]
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"写入历史记录失败: {e}")

    def load_history(self) -> List[Dict]:
        """加载所有历史记录"""
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            logger.warning(f"读取历史记录失败: {e}")
        return []

    def clear_history(self):
        """清除所有历史记录"""
        try:
            if os.path.exists(self.history_file):
                os.remove(self.history_file)
                st.success("✅ 历史已清除", icon="✅")
        except Exception as e:
            logger.warning(f"清除历史记录失败: {e}")


# ==================== 会话状态初始化 ====================

def init_session_state():
    """初始化Streamlit会话状态"""
    required_states = {
        "replace_rules": [],
        "replaced_files": [],
        "replace_log": [],
        "is_replacing": False,
        "replace_params": {},
        "replace_scope": "替换完整关键词",
        "export_mode_radio": "独立文件（ZIP）",
        "undo_stack": [],
        "rule_filter": "",
        "show_advanced": False,
        "excel_cache": None,
        "current_page": 1,
    }

    for key, default in required_states.items():
        if key not in st.session_state:
            st.session_state[key] = default


init_session_state()


# ==================== 核心工具函数 ====================

def safe_load_json_bytes(raw: bytes, max_bytes: int = MAX_RULE_IMPORT_SIZE):
    """安全加载JSON字节，限制体积并校验格式"""
    if not raw:
        raise ValueError("空文件")
    if len(raw) > max_bytes:
        raise ValueError("文件过大")
    return json.loads(raw.decode('utf-8'))



# ==================== 创建管理器实例 ====================
cache_manager = CacheManager()
history_manager = HistoryManager()

# ==================== 侧栏 ====================
with st.sidebar:
    st.title("📚 快速导航")

    # 统计信息
    if st.session_state.replaced_files:
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 文件数", len(st.session_state.replaced_files), delta=None)
        with col2:
            st.metric("📋 规则数", len(st.session_state.replace_rules), delta=None)

    st.markdown("---")

    # 缓存信息显示
    cache_info = get_cache_info()
    with st.expander("💾 缓存信息", expanded=False):
        col_info1, col_info2 = st.columns(2)
        with col_info1:
            st.caption(f"**规则缓存**: {cache_info['rules_count']} 个")
        with col_info2:
            st.caption(f"**总大小**: {format_file_size(cache_info['total_size'])}")

        st.caption(f"📂 位置: {cache_info['rules_dir']}")

        # 清除缓存按钮
        if st.button("🗑️ 清除所有缓存", key="clear_cache_all", use_container_width=True):
            if cache_manager.clear_all_cache():
                st.success("✅ 缓存已清除", icon="✅")
                st.rerun()

    st.markdown("---")

    # 快速功能
    st.subheader("⚡ 快速功能")

    # 快速加载缓存规则
    cached = cache_manager.get_cached_rules_list()
    if cached:
        st.markdown("**📂 加载规则缓存**")
        selected = st.selectbox(
            "选择规则",
            options=["选择..."] + cached,
            key="sidebar_cache",
            label_visibility="collapsed"
        )

        if selected and selected != "选择...":
            # 显示规则信息
            rule_info = cache_manager.get_rule_info(selected)
            if rule_info:
                st.caption(f"📋 {rule_info['rules_count']} 个规则 | 📅 {rule_info['mtime']}")

            col_load, col_del = st.columns(2, gap="small")
            with col_load:
                if st.button("✅ 加载", key="sidebar_load", use_container_width=True):
                    loaded = cache_manager.load_rules(selected)
                    if loaded:
                        st.session_state.replace_rules = loaded
                        st.success(f"✅ 加载{len(loaded)}条规则", icon="✅")
                        st.rerun()
            with col_del:
                if st.button("🗑️ 删除", key="sidebar_del_cache", use_container_width=True):
                    if cache_manager.delete_rule(selected):
                        st.success("✅ 规则已删除", icon="✅")
                        st.rerun()
    else:
        st.info("📁 暂无缓存规则", icon="ℹ️")

    st.markdown("---")

    # 历史记录显示
    history = history_manager.load_history()
    if history:
        st.subheader("📜 最近操作")
        for h in history[:3]:
            status = "✅" if h["status"] == "success" else "❌"
            st.caption(f"{status} {h['timestamp']}\n{h['word_file'][:15]}...")

    st.markdown("---")

    # 工具操作
    st.subheader("🔧 工具")

    if st.button("🗑️ 清空当前规则", key="sidebar_clear", use_container_width=True):
        st.session_state.replace_rules = []
        st.session_state.replaced_files = []
        st.session_state.current_page = 1
        st.success("✅ 已清空", icon="✅")
        st.rerun()

    if history:
        if st.button("📜 清除历史记录", key="sidebar_clear_hist", use_container_width=True):
            history_manager.clear_history()
            st.rerun()

# ==================== 主页面 - 标题 ====================
st.markdown(
    f"""
    <div style="display:flex;justify-content:space-between;align-items:flex-end;gap:0.75rem;flex-wrap:wrap;">
      <h1 style="margin:0;line-height:1.25;word-break:break-word;">📋 Word+Excel批量替换工具</h1>
      <small style="color:#64748b;white-space:nowrap;">{VERSION}</small>
    </div>
    """,
    unsafe_allow_html=True,
)
st.markdown(
    "<p class='wr-subtitle'>上传模板与数据表，配置规则后即可批量生成目标文档。</p>",
    unsafe_allow_html=True,
)

step_cols = st.columns(5)
step_labels = ["1 上传文件", "2 预览数据", "3 配置规则", "4 执行替换", "5 下载结果"]
for idx, label in enumerate(step_labels):
    with step_cols[idx]:
        step_cls = "wr-step active" if idx == 0 else "wr-step"
        st.markdown(f"<span class='{step_cls}'>{label}</span>", unsafe_allow_html=True)

# 进度显示
if st.session_state.replaced_files and st.session_state.replace_params:
    progress_col, status_col = st.columns([3, 1])
    with progress_col:
        success_count = len([f for f in st.session_state.replaced_files
                             if f.data and len(f.data.getvalue()) > 0])
        total_count = len(st.session_state.replaced_files)
        st.progress(success_count / total_count if total_count > 0 else 0)
    with status_col:
        st.metric("成功率", f"{int(success_count / total_count * 100) if total_count > 0 else 0}%")

st.markdown("---")

# ==================== 主工作区 ====================
col_main_left, col_main_right = st.columns([2, 1], gap="medium")

# ==================== 左侧：文件上传和预览 ====================
with col_main_left:
    st.subheader("📤 文件上传")

    # 上传区域
    col_upload1, col_upload2 = st.columns(2, gap="small")

    with col_upload1:
        st.markdown("<div class='wr-upload-card'>", unsafe_allow_html=True)
        st.markdown(create_tooltip("**Word模板**", "word_upload"), unsafe_allow_html=True)
        st.markdown("<p class='wr-upload-help'>支持 .docx，建议使用标准占位符格式。</p>", unsafe_allow_html=True)

        word_file = st.file_uploader(
            "选择文件",
            type=["docx"],
            key="word",
            label_visibility="collapsed",
            help="仅支持.docx格式"
        )
        if word_file:
            file_size_bytes = len(word_file.getvalue())
            file_size_str = format_file_size(file_size_bytes)

            if file_size_bytes > MAX_WORD_FILE_SIZE:
                st.error(f"❌ 文件过大：{file_size_str}", icon="❌")
                word_file = None
            else:
                st.success(f"已加载：{word_file.name}（{file_size_str}）", icon="✅")
        st.markdown("</div>", unsafe_allow_html=True)

    with col_upload2:
        st.markdown("<div class='wr-upload-card'>", unsafe_allow_html=True)
        st.markdown(create_tooltip("**Excel数据**", "excel_upload"), unsafe_allow_html=True)
        st.markdown("<p class='wr-upload-help'>支持 .xlsx，首行建议为字段名。</p>", unsafe_allow_html=True)

        excel_file = st.file_uploader(
            "选择文件",
            type=["xlsx"],
            key="excel",
            label_visibility="collapsed",
            help="仅支持.xlsx格式"
        )
        if excel_file:
            file_size_bytes = len(excel_file.getvalue())
            file_size_str = format_file_size(file_size_bytes)

            if file_size_bytes > MAX_EXCEL_FILE_SIZE:
                st.error(f"❌ 文件过大：{file_size_str}", icon="❌")
                excel_file = None
            else:
                st.success(f"已加载：{excel_file.name}（{file_size_str}）", icon="✅")
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")

    # 文件预览
    with st.expander("👀 文件预览 - 点击查看/复制内容", expanded=False):
        col_prev1, col_prev2 = st.columns(2, gap="small")

        excel_df = None
        excel_cols = []

        with col_prev1:
            st.markdown("**Word文档内容**")
            if word_file:
                try:
                    doc = Document(io.BytesIO(word_file.getvalue()))

                    html_content = ""

                    for para in doc.paragraphs[:15]:
                        if para.text.strip():
                            text = para.text.replace("<", "&lt;").replace(">", "&gt;")
                            html_content += f"<p style='margin: 4px 0; word-break: break-all;'>{text}</p>"

                    for table_idx, table in enumerate(doc.tables[:2]):
                        html_content += f"<p style='margin-top: 8px; font-weight: bold; color: #1f77b4;'>📊 表格{table_idx + 1}：</p>"
                        html_content += "<table style='border-collapse: collapse; width: 100%; font-size: 12px;'>"

                        for row_idx, row in enumerate(table.rows[:10]):
                            html_content += "<tr>"
                            for cell in row.cells:
                                cell_text = cell.text.replace("<", "&lt;").replace(">", "&gt;")[:30]
                                html_content += f"<td style='border: 1px solid #ccc; padding: 4px;'>{cell_text}</td>"
                            html_content += "</tr>"

                        html_content += "</table>"

                    st.components.v1.html(f"""
                    <div style='height: 280px; overflow-y: auto; padding: 12px; border: 1px solid #e0e0e0; 
                                border-radius: 6px; font-size: 13px; line-height: 1.6; background-color: #f9f9f9;
                                font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif; word-wrap: break-word;
                                user-select: text;'>
                        {html_content}
                    </div>
                    """, height=300)

                    st.caption(f"📄 {len(doc.paragraphs)}段落，{len(doc.tables)}表格")
                    st.info("💡 可以在上方选中内容按Ctrl+C复制，粘贴到下方关键字输入框中", icon="ℹ️")

                except Exception as e:
                    st.error(f"❌ 预览失败", icon="❌")
            else:
                st.info("请上传Word文件", icon="ℹ️")

        with col_prev2:
            st.markdown("**Excel数据预览**")
            if excel_file:
                try:
                    excel_df = pd.read_excel(
                        io.BytesIO(excel_file.getvalue()),
                        dtype=str,
                        keep_default_na=False,
                        na_values=[]
                    )

                    if excel_df.empty:
                        st.warning("⚠️ 表格为空", icon="⚠️")
                    else:
                        excel_df = clean_excel_types(excel_df)
                        excel_cols = excel_df.columns.tolist()

                        preview_df = excel_df.head(PREVIEW_ROWS)

                        st.dataframe(
                            preview_df,
                            use_container_width=True,
                            hide_index=True,
                            height=280
                        )

                        col_s1, col_s2 = st.columns(2)
                        with col_s1:
                            st.metric("行数", len(excel_df))
                        with col_s2:
                            st.metric("列数", len(excel_cols))

                except Exception as e:
                    logger.warning(f"读取Excel失败: {e}")
                    st.error("❌ 读取失败：请确认是标准.xlsx文件", icon="❌")
            else:
                st.info("请上传Excel文件", icon="ℹ️")

# ==================== 右侧：规则管理 ====================
with col_main_right:
    st.subheader("📋 规则管理")

    # 替换范围
    st.markdown(create_tooltip("**替换范围**", "replace_scope"), unsafe_allow_html=True)

    replace_scope = st.radio(
        "模式",
        options=["完整关键词", "括号内容"],
        key="replace_scope_compact",
        horizontal=True,
        label_visibility="collapsed"
    )
    st.session_state.replace_scope = ["替换完整关键词", "仅替换括号内内容"][
        ["完整关键词", "括号内容"].index(replace_scope)]

    st.markdown("---")

    # 规则列表
    st.markdown(create_tooltip("**规则列表**", "rule_list"), unsafe_allow_html=True)
    st.markdown(
        f"<div class='wr-rule-toolbar'><span>共 {len(st.session_state.replace_rules)} 条规则</span><span>支持撤销与批量导入</span></div>",
        unsafe_allow_html=True,
    )

    if st.session_state.replace_rules:
        with st.container(border=True):
            for idx, (old, col) in enumerate(st.session_state.replace_rules):
                col_rule, col_del = st.columns([3, 0.6], gap="small")
                rule_preview = f"{old} → {col}"
                with col_del:
                    if st.button("删除", key=f"del_{idx}", use_container_width=True, help="删除此规则"):
                        st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                        st.session_state.replace_rules.pop(idx)
                        st.session_state.replaced_files = []
                        st.rerun()
                with col_rule:
                    st.markdown(
                        f"<span class='wr-rule-item' title='{rule_preview}'><strong>{old}</strong> → {col}</span>",
                        unsafe_allow_html=True,
                    )

        # 规则操作按钮
        col_undo, col_clear = st.columns(2, gap="small")
        with col_undo:
            if st.session_state.undo_stack:
                if st.button("↶ 撤销", key="undo", use_container_width=True, help=HELP_TEXTS["undo"]):
                    st.session_state.replace_rules = st.session_state.undo_stack.pop()
                    st.success("✅ 已撤销", icon="✅")
                    st.rerun()
        with col_clear:
            if st.button("🗑️ 清空", key="clear_rules", use_container_width=True, help=HELP_TEXTS["clear_rules"]):
                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                st.session_state.replace_rules.clear()
                st.session_state.replaced_files = []
                st.rerun()
    else:
        st.info("📁 暂无规则", icon="ℹ️")

    st.markdown("---")

    # 添加规则
    st.markdown(create_tooltip("**新增规则**", "add_rule"), unsafe_allow_html=True)

    new_keyword = st.text_input(
        "关键字",
        placeholder="如：【姓名】",
        key="new_keyword",
        label_visibility="collapsed",
        help=HELP_TEXTS["new_keyword"]
    )

    if excel_cols:
        new_column = st.selectbox(
            "列",
            options=excel_cols,
            key="new_column",
            label_visibility="collapsed",
            help=HELP_TEXTS["new_column"]
        )
    else:
        new_column = None

    if st.button(
            "➕ 添加规则",
            key="add_rule",
            type="primary",
            disabled=not (new_keyword and new_column),
            use_container_width=True,
            help=HELP_TEXTS["add_rule"]
    ):
        rule = (new_keyword.strip(), new_column)
        if rule in st.session_state.replace_rules:
            st.warning("⚠️ 规则已存在", icon="⚠️")
        else:
            st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
            st.session_state.replace_rules.append(rule)
            st.success("✅ 已添加", icon="✅")
            st.rerun()

    st.markdown("---")

    # 规则导入导出
    with st.expander("💾 导入/导出/缓存", expanded=False):
        # 导入
        import_file = st.file_uploader(
            "导入JSON",
            type=["json"],
            key="import_rules",
            label_visibility="collapsed",
            help=HELP_TEXTS["rule_import"]
        )

        if import_file:
            try:
                rules_data = safe_load_json_bytes(import_file.getvalue())
                valid_rules = []
                if not isinstance(rules_data, list):
                    raise ValueError("规则文件格式错误")
                if len(rules_data) > MAX_RULE_CACHE_ITEMS:
                    raise ValueError("规则条数过多")
                for rule in rules_data:
                    if isinstance(rule, dict) and "keyword" in rule and "excel_column" in rule:
                        keyword = str(rule["keyword"]).strip()
                        excel_col = str(rule["excel_column"]).strip()
                        if keyword and excel_col:
                            valid_rules.append((keyword, excel_col))

                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                for rule in valid_rules:
                    if rule not in st.session_state.replace_rules:
                        st.session_state.replace_rules.append(rule)

                st.success(f"✅ 导入{len(valid_rules)}条", icon="✅")
                st.rerun()
            except Exception as e:
                logger.warning(f"导入规则失败: {e}")
                st.error("❌ 格式错误", icon="❌")

        # 导出
        if st.session_state.replace_rules:
            rules_data = [
                {"keyword": old, "excel_column": col}
                for old, col in st.session_state.replace_rules
            ]
            rules_json = json.dumps(rules_data, ensure_ascii=False, indent=2)

            col_exp1, col_exp2 = st.columns(2, gap="small")
            with col_exp1:
                st.download_button(
                    label="📥 导出JSON",
                    data=rules_json,
                    file_name="rules.json",
                    mime="application/json",
                    key="export_rules",
                    use_container_width=True,
                    help=HELP_TEXTS["rule_export"]
                )
            with col_exp2:
                if st.button("💾 保存缓存", key="save_cache", use_container_width=True,
                             help=HELP_TEXTS["rule_cache"]):
                    if cache_manager.save_rules(st.session_state.replace_rules):
                        st.success("✅ 已保存到缓存", icon="✅")
                        st.rerun()

st.markdown("---")

# ==================== 底部：执行替换和参数配置 ====================
st.subheader("⚙️ 替换参数配置")

col_config1, col_config2, col_config3, col_config4 = st.columns(4, gap="small")

with col_config1:
    st.markdown(create_tooltip("**核心字段**", "file_name_col"), unsafe_allow_html=True)
    file_name_col = st.selectbox(
        "用于文件名",
        options=excel_cols if excel_cols else ["未选择"],
        key="file_name_col",
        disabled=not excel_cols,
        label_visibility="collapsed",
        help=HELP_TEXTS["file_name_col"]
    )

with col_config2:
    st.markdown(create_tooltip("**起始行**", "start_row"), unsafe_allow_html=True)
    start_row = st.number_input(
        "开始",
        min_value=1,
        max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        value=1,
        key="start_row",
        disabled=excel_df is None or len(excel_df) == 0,
        label_visibility="collapsed",
        help=HELP_TEXTS["start_row"]
    )

with col_config3:
    st.markdown(create_tooltip("**结束行**", "end_row"), unsafe_allow_html=True)
    end_row = st.number_input(
        "结束",
        min_value=1,
        max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        key="end_row",
        disabled=excel_df is None or len(excel_df) == 0,
        label_visibility="collapsed",
        help=HELP_TEXTS["end_row"]
    )

with col_config4:
    st.markdown(create_tooltip("**文件前缀**", "file_prefix"), unsafe_allow_html=True)
    file_prefix = st.text_input(
        "前缀",
        value="",
        key="file_prefix",
        placeholder="可选",
        max_chars=15,
        label_visibility="collapsed",
        help=HELP_TEXTS["file_prefix"]
    ).strip()

if start_row > end_row:
    st.error("❌ 起始行不能大于结束行", icon="❌")

st.markdown("---")

# ==================== 执行替换 ====================
blockers = get_replace_blockers(
    word_file,
    excel_df,
    st.session_state.replace_rules,
    int(start_row),
    int(end_row)
)
can_replace = len(blockers) == 0
if blockers:
    st.warning("；".join(blockers), icon="⚠️")

current_params = get_replace_params(
    word_file,
    excel_df,
    int(start_row),
    int(end_row),
    file_name_col,
    file_prefix,
    "",
    st.session_state.replace_rules
)

need_replace = (
        len(st.session_state.replaced_files) == 0 or
        st.session_state.replace_params != current_params
)

# 执行前摘要卡
planned_total = 0
if excel_df is not None and len(excel_df) > 0 and start_row <= end_row:
    planned_total = max(0, min(int(end_row), len(excel_df)) - int(start_row) + 1)
summary_col1, summary_col2, summary_col3 = st.columns(3, gap="small")
with summary_col1:
    st.metric("🧾 预计生成", planned_total)
with summary_col2:
    st.metric("📋 当前规则", len(st.session_state.replace_rules))
with summary_col3:
    st.metric("📄 行范围", f"{int(start_row)}-{int(end_row)}")

col_exec1, col_exec2, col_exec3, col_exec4 = st.columns([2, 1.5, 1.5, 1], gap="small")

with col_exec1:
    replace_btn = st.button(
        "▶️ 开始替换",
        key="replace",
        disabled=not can_replace or st.session_state.is_replacing or start_row > end_row,
        type="primary",
        use_container_width=True,
        help=(HELP_TEXTS["start_replace"] if can_replace else "；".join(blockers))
    )

with col_exec2:
    if st.session_state.is_replacing:
        st.info("🔄 进行中", icon="🔄")
    elif len(st.session_state.replaced_files) > 0 and not need_replace:
        st.success(f"✅ {len(st.session_state.replaced_files)}个", icon="✅")

# 执行替换逻辑
if replace_btn and not st.session_state.is_replacing:
    st.session_state.is_replacing = True
    st.session_state.replaced_files = []
    st.session_state.replace_log = []
    st.session_state.current_page = 1

    progress_bar = st.progress(0)
    progress_text = st.empty()

    try:
        actual_end_row = min(end_row, len(excel_df))
        if start_row > actual_end_row:
            st.error("❌ 行号超出范围", icon="❌")
        else:
            total_rows = actual_end_row - start_row + 1
            used_output_names = set()

            for idx, row_idx in enumerate(range(start_row - 1, actual_end_row)):
                try:
                    excel_row = excel_df.iloc[row_idx]

                    replaced_file, replace_log, replace_cnt = replace_word_with_format(
                        word_file,
                        excel_row,
                        st.session_state.replace_rules,
                        st.session_state.replace_scope
                    )

                    filename = generate_safe_filename(
                        excel_row,
                        file_name_col if file_name_col != "未选择" else "",
                        file_prefix,
                        "",
                        row_idx
                    )
                    filename = dedupe_filename(filename, used_output_names)

                    st.session_state.replaced_files.append(ReplacedFile(
                        filename=filename,
                        data=replaced_file,
                        row_idx=row_idx,
                        log=replace_log,
                        replace_count=replace_cnt
                    ))

                    st.session_state.replace_log.append(f"【{row_idx + 1}】{replace_log}")

                    progress = (idx + 1) / total_rows
                    progress_bar.progress(progress)
                    progress_text.text(f"{idx + 1}/{total_rows}")

                except Exception as e:
                    logger.warning(f"处理第{row_idx + 1}行失败: {e}")
                    st.session_state.replace_log.append(f"【{row_idx + 1}】❌ 失败")
                    continue

            st.session_state.replace_params = current_params
            st.success(f"🎉 完成！{len(st.session_state.replaced_files)} 个文件", icon="✅")

            history_record = HistoryRecord(
                timestamp=datetime.now().strftime("%m-%d %H:%M"),
                word_file=word_file.name[:20],
                excel_file=excel_file.name[:20],
                rules_count=len(st.session_state.replace_rules),
                files_generated=len(st.session_state.replaced_files),
                status="success"
            )
            history_manager.add_record(history_record)

    except Exception as e:
        logger.warning(f"批量替换流程失败: {e}")
        st.error(f"❌ 出错", icon="❌")
    finally:
        st.session_state.is_replacing = False
        progress_bar.empty()
        progress_text.empty()

st.markdown("---")

# ==================== 下载结果区 ====================
if len(st.session_state.replaced_files) > 0:
    st.subheader("💾 下载结果")

    # 摘要优先展示
    success_count = len([f for f in st.session_state.replaced_files if f.data and len(f.data.getvalue()) > 0])
    total_count = len(st.session_state.replaced_files)
    fail_count = max(0, total_count - success_count)
    total_replace = sum(f.replace_count for f in st.session_state.replaced_files)
    sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4, gap="small")
    with sum_col1:
        st.metric("📄 总数", total_count)
    with sum_col2:
        st.metric("✅ 成功", success_count)
    with sum_col3:
        st.metric("❌ 失败", fail_count)
    with sum_col4:
        st.metric("🔄 替换次", total_replace)

    col_export_opt1, col_export_opt2 = st.columns([2, 2])

    with col_export_opt1:
        st.markdown("**导出方式**")

    export_mode = st.radio(
        "方式",
        options=["独立文件（ZIP）", "合并为单个文档"],
        key="export_mode_radio",
        horizontal=True,
        label_visibility="collapsed"
    )

    st.markdown("---")

    # 统计信息（保留规则数，其他核心摘要已上移）
    col_stat1, col_stat2 = st.columns(2, gap="small")
    with col_stat1:
        st.metric("📋 规则数", len(st.session_state.replace_rules))
    with col_stat2:
        st.metric("📈 成功率", f"{int(success_count / total_count * 100) if total_count > 0 else 0}%")

    st.markdown("---")

    # 导出按钮
    col_down1, col_down2, col_down3 = st.columns(3, gap="small")

    with col_down1:
        if export_mode == "独立文件（ZIP）":
            try:
                valid_files = [f for f in st.session_state.replaced_files
                               if f.data and len(f.data.getvalue()) > 0]

                if valid_files:
                    zip_buffer = io.BytesIO()
                    zip_ready = False
                    if len(valid_files) > MAX_EXPORT_FILES:
                        st.error(f"❌ 文件数量过多（>{MAX_EXPORT_FILES}）", icon="❌")
                    else:
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                            used_names = set()
                            for file in valid_files:
                                safe_name = clean_filename(os.path.basename(file.filename))
                                base, ext = os.path.splitext(safe_name)
                                candidate = safe_name
                                idx = 1
                                while candidate in used_names:
                                    candidate = f"{base}_{idx}{ext}"
                                    idx += 1
                                used_names.add(candidate)
                                zipf.writestr(candidate, file.data.getvalue())
                        zip_ready = True

                    if zip_ready:
                        zip_buffer.seek(0)
                        zip_filename = f"批量替换_{len(valid_files)}个.zip"

                        st.download_button(
                            label=f"📦 下载ZIP（{len(valid_files)}个）",
                            data=zip_buffer,
                            file_name=zip_filename,
                            mime="application/zip",
                            key="download_all_zip",
                            on_click="ignore",
                            use_container_width=True,
                            type="primary",
                            help=HELP_TEXTS["export_zip"]
                        )
            except Exception as e:
                logger.warning(f"创建ZIP失败: {e}")
                st.error("❌ 创建ZIP失败", icon="❌")
        else:
            valid_files = [f for f in st.session_state.replaced_files
                           if f.data and len(f.data.getvalue()) > 0]

            if valid_files:
                try:
                    merged_data = merge_word_documents(valid_files)

                    st.download_button(
                        label=f"📋 下载合并文档（{len(valid_files)}个）",
                        data=merged_data,
                        file_name="合并结果.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_merged",
                        on_click="ignore",
                        use_container_width=True,
                        type="primary",
                        help=HELP_TEXTS["export_merge"]
                    )
                except Exception as e:
                    logger.warning(f"合并导出失败: {e}")
                    st.error("❌ 合并失败", icon="❌")

    with col_down2:
        csv_data = export_statistics_to_csv(st.session_state.replaced_files)
        st.download_button(
            label="📊 下载CSV统计",
            data=csv_data,
            file_name="统计.csv",
            mime="text/csv",
            key="download_stats",
            on_click="ignore",
            use_container_width=True,
            help=HELP_TEXTS["export_stats"]
        )

    with col_down3:
        if st.session_state.replace_log:
            log_text = "\n".join(st.session_state.replace_log)
            st.download_button(
                label="📝 导出日志",
                data=log_text,
                file_name="替换日志.txt",
                mime="text/plain",
                key="download_log",
                on_click="ignore",
                use_container_width=True,
                help=HELP_TEXTS["export_log"]
            )

    st.markdown("---")

    # 文件列表
    st.markdown(create_tooltip(f"**文件列表** ({len(st.session_state.replaced_files)})", "rule_list"),
                unsafe_allow_html=True)

    # 分页
    total_pages = (len(st.session_state.replaced_files) + PAGE_SIZE - 1) // PAGE_SIZE

    col_page1, col_page2, col_page3 = st.columns([2, 1, 2])

    with col_page2:
        st.session_state.current_page = min(max(1, int(st.session_state.current_page)), total_pages)
        current_page = st.number_input(
            "页",
            min_value=1,
            max_value=total_pages,
            value=st.session_state.current_page,
            key="current_page",
            label_visibility="collapsed"
        )

    start_idx = (current_page - 1) * PAGE_SIZE
    end_idx = min(start_idx + PAGE_SIZE, len(st.session_state.replaced_files))
    current_files = st.session_state.replaced_files[start_idx:end_idx]

    st.caption(f"第 {current_page}/{total_pages} 页")

    # 文件表格
    file_data = []
    for idx, file in enumerate(current_files, start=start_idx + 1):
        is_valid = file.data and len(file.data.getvalue()) > 0
        status = "✅" if is_valid else "❌"
        file_data.append({
            "状态": status,
            "序号": idx,
            "文件名": file.filename[:25] + "..." if len(file.filename) > 25 else file.filename,
            "行号": file.row_idx + 1,
            "替换": file.replace_count
        })

    if file_data:
        file_df = pd.DataFrame(file_data)
        st.dataframe(file_df, use_container_width=True, hide_index=True)

    # 单个文件下载
    st.markdown("**单个文件下载**")

    for idx, file in enumerate(current_files, start=start_idx + 1):
        is_valid = file.data and len(file.data.getvalue()) > 0

        col_name, col_log, col_download = st.columns([2, 1, 1], gap="small")

        with col_name:
            st.caption(f"#{idx} {file.filename}")

        with col_log:
            if st.button("📋 日志", key=f"log_{idx}", use_container_width=True,
                         help=HELP_TEXTS["single_log"]):
                st.write(file.log)

        with col_download:
            st.download_button(
                label="⬇️ 下载",
                data=file.data,
                file_name=file.filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{idx}",
                on_click="ignore",
                disabled=not is_valid,
                use_container_width=True,
                help=HELP_TEXTS["single_download"]
            )

else:
    st.info("💡 上传文件、添加规则后点击'开始替换'", icon="ℹ️")

# ==================== 底部帮助 ====================
st.markdown("---")

with st.expander("❓ 帮助指南", expanded=False):
    col_help1, col_help2 = st.columns(2, gap="medium")

    with col_help1:
        st.markdown("""
        **快速开始**
        1. 📤 上传Word和Excel文件
        2. 📋 添加替换规则
        3. ▶️ 点击"开始替换"
        4. 💾 下载结果

        **支持格式**
        • Word：.docx（不支持.doc）
        • Excel：.xlsx
        • 括号：【】（）()〔〕

        **文件限制**
        • Word最大50MB
        • Excel最大50MB
        • 建议行数<1000
        """)

    with col_help2:
        st.markdown("""
        **常见问题**

        ❓ **Word文件不支持.doc？**
        用Word打开文件 → 另存为.docx格式

        ❓ **怎样保留格式？**
        所有格式自动保留：字体、颜色、表格等

        ❓ **如何合并文档？**
        选择"合并为单个文档"导出方式

        ❓ **能否撤销操作？**
        点击"↶ 撤销"按钮撤销最后一次规则操作

        ❓ **缓存文件保存在哪？**
        Windows：`%APPDATA%/BatchReplacer`
        Mac/Linux：`~/.cache/batch_replacer`
        """)

st.caption(f"Word+Excel批量替换工具 {VERSION} © 2026")
