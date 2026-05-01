"""Business services for replacement, merge, and statistics."""

import copy
import io
import re
import hashlib
import logging
from collections import defaultdict
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from core_utils import clean_text
except ImportError:
    from app.core_utils import clean_text

logger = logging.getLogger("wordreplace")

MAX_WORD_FILE_SIZE = 50 * 1024 * 1024


def precompute_replace_patterns(
    replace_rules: List[Tuple[str, str]],
    excel_row: pd.Series,
    replace_scope: str,
) -> List[Tuple[str, str, str, str]]:
    """预计算所有需要替换的模式。"""
    replace_patterns = []

    for old_text, col_name in replace_rules:
        replacement = str(excel_row[col_name]).strip() if col_name in excel_row.index else ""
        cleaned_text = clean_text(old_text)
        if not cleaned_text:
            continue

        if replace_scope == "仅替换括号内内容":
            if cleaned_text.startswith("【") and cleaned_text.endswith("】"):
                new_format = f"【{replacement}】"
            elif cleaned_text.startswith("（") and cleaned_text.endswith("）"):
                new_format = f"（{replacement}）"
            elif cleaned_text.startswith("(") and cleaned_text.endswith(")"):
                new_format = f"({replacement})"
            elif cleaned_text.startswith("〔") and cleaned_text.endswith("〕"):
                new_format = f"〔{replacement}〕"
            else:
                new_format = replacement
            replace_patterns.append((old_text, col_name, cleaned_text, new_format))
        else:
            replace_patterns.append((old_text, col_name, cleaned_text, replacement))

    return replace_patterns


def process_paragraph(paragraph, replace_patterns: List[Tuple[str, str, str, str]], cleaned_para: str = None) -> Dict:
    """处理单个段落的关键字替换。"""
    para_text = paragraph.text
    if cleaned_para is None:
        cleaned_para = clean_text(para_text)
    replace_count = defaultdict(int)

    if not para_text or not replace_patterns:
        return replace_count

    has_keyword = any(format_keyword and format_keyword in cleaned_para for _, _, format_keyword, _ in replace_patterns)
    if has_keyword:
        new_text = para_text
        for old_text, col_name, format_keyword, replacement in replace_patterns:
            if format_keyword and format_keyword in cleaned_para:
                count = new_text.count(format_keyword)
                if count > 0:
                    new_text = new_text.replace(format_keyword, replacement)
                    replace_count[(old_text, col_name)] += count

        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""

    return replace_count


def replace_word_with_format(word_file, excel_row: pd.Series, replace_rules: List[Tuple[str, str]], replace_scope: str):
    """替换Word文件中的关键字，保留格式。"""
    replace_count = defaultdict(int)
    total_replace = 0

    try:
        if len(word_file.getvalue()) > MAX_WORD_FILE_SIZE:
            raise ValueError("文件过大")

        doc = Document(io.BytesIO(word_file.getvalue()))
        replace_patterns = precompute_replace_patterns(replace_rules, excel_row, replace_scope)

        if not replace_patterns:
            output_file = io.BytesIO()
            doc.save(output_file)
            output_file.seek(0)
            return output_file, "⚠ 未找到匹配规则", 0

        for paragraph in doc.paragraphs:
            para_count = process_paragraph(paragraph, replace_patterns)
            for key, count in para_count.items():
                replace_count[key] += count
                total_replace += count

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_count = process_paragraph(paragraph, replace_patterns)
                        for key, count in para_count.items():
                            replace_count[key] += count
                            total_replace += count

        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)

        if replace_count:
            log_lines = [f"✓ {old}({count}次)" for (old, _), count in replace_count.items()]
            replace_log = ", ".join(log_lines[:3])
            if len(replace_count) > 3:
                replace_log += f" 等{len(replace_count) - 3}个"
        else:
            replace_log = "⚠ 无替换"

        return output_file, replace_log, total_replace
    except Exception as e:
        logger.warning(f"替换Word失败: {e}")
        return io.BytesIO(), "❌ 失败", 0


def merge_word_documents(replaced_files) -> io.BytesIO:
    """合并多个Word文档（保留所有格式和结构）。"""
    if not replaced_files:
        raise ValueError("没有文件")

    main_doc = Document(io.BytesIO(replaced_files[0].data.getvalue()))
    main_body = main_doc._body._element

    for idx in range(1, len(replaced_files)):
        try:
            file = replaced_files[idx]
            if not file.data or len(file.data.getvalue()) == 0:
                continue

            sub_doc = Document(io.BytesIO(file.data.getvalue()))
            sub_body = sub_doc._body._element

            page_break_para = OxmlElement("w:p")
            page_break_pPr = OxmlElement("w:pPr")
            page_break_element = OxmlElement("w:pageBreakBefore")
            page_break_element.set(qn("w:val"), "1")
            page_break_pPr.append(page_break_element)
            page_break_para.append(page_break_pPr)
            main_body.append(page_break_para)

            for element in sub_body:
                main_body.append(copy.deepcopy(element))
        except Exception as e:
            logger.warning(f"合并文档时跳过一个子文档: {e}")
            continue

    output = io.BytesIO()
    main_doc.save(output)
    output.seek(0)
    return output


def get_replace_params(word_file, excel_df: Optional[pd.DataFrame], start_row: int, end_row: int, file_name_col: str, file_prefix: str, file_suffix: str, replace_rules: List[Tuple[str, str]]) -> Dict:
    """获取替换参数，用于判断是否需要重新替换。"""
    return {
        "word_filename": word_file.name if word_file else "",
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "file_prefix": file_prefix,
        "file_suffix": file_suffix,
        "rule_count": len(replace_rules),
        "rule_hash": hash(tuple(replace_rules)),
    }


def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """清理Excel数据类型，避免混合类型导致的问题。"""
    df_clean = df.copy()
    for col in df_clean.columns:
        try:
            col_name = str(col)
            if col_name != col:
                df_clean = df_clean.rename(columns={col: col_name})
                col = col_name
            df_clean[col] = df_clean[col].fillna("")
            df_clean[col] = df_clean[col].astype(str).str.strip()
        except Exception as e:
            logger.warning(f"清理Excel列类型失败({col}): {e}")
            try:
                df_clean[col] = df_clean[col].astype(str).str.strip()
            except Exception as inner_e:
                logger.warning(f"二次清理Excel列类型失败({col}): {inner_e}")
    return df_clean


def get_file_hash(file_data: bytes) -> str:
    return hashlib.md5(file_data).hexdigest()[:6]


def export_statistics_to_csv(replaced_files) -> str:
    """导出替换统计数据到CSV格式。"""
    try:
        data = []
        for idx, file in enumerate(replaced_files, 1):
            data.append({
                "序号": idx,
                "文件名": file.filename,
                "行号": file.row_idx + 1,
                "替换次数": file.replace_count,
                "状态": "✅" if file.data and len(file.data.getvalue()) > 0 else "❌",
            })
        df = pd.DataFrame(data)
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding="utf-8-sig")
        return csv_buffer.getvalue()
    except Exception as e:
        logger.warning(f"导出统计CSV失败: {e}")
        return ""


def get_keyword_statistics(replace_rules: List[Tuple[str, str]], replaced_files) -> Dict:
    """获取关键字替换统计。"""
    stats = {keyword: 0 for keyword, _ in replace_rules}
    for file in replaced_files:
        for keyword, _ in replace_rules:
            if f"✓ {keyword}" in file.log:
                pattern = f"✓ {re.escape(keyword)}\\((\\d+)次\\)"
                matches = re.findall(pattern, file.log)
                if matches:
                    stats[keyword] += sum(int(m) for m in matches)
    return stats
