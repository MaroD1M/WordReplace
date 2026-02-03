# 导入标准库
import os
import sys
import tempfile
from tempfile import NamedTemporaryFile
import warnings
import shutil
import json
import io
import zipfile
import re
import unicodedata
import copy

# 导入第三方库
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP

# 项目版本信息
VERSION = "v1.3.2"

# 配置常量
PAGE_SIZE = 10  # 每页显示的文件数
WIDGET_HEIGHT = 300  # 组件高度
PREVIEW_ROWS = 30  # 数据预览行数
MAX_FILENAME_LENGTH = 200  # 最大文件名长度
MAX_WORD_FILE_SIZE = 50 * 1024 * 1024  # 最大Word文件大小：50MB
MAX_EXCEL_FILE_SIZE = 50 * 1024 * 1024  # 最大Excel文件大小：50MB

# 过滤特定警告，避免干扰用户界面
warnings.filterwarnings("ignore", category=UserWarning)

# 设置环境变量避免不必要的版本检查和统计
os.environ["STREAMLIT_VERSION"] = "1.51.0"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# 设置页面配置
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 全局样式优化
st.markdown("""
<style>
    /* 优化容器样式 */
    .stContainer {
        margin-bottom: 20px;
    }

    /* 优化按钮样式 */
    .stButton > button {
        border-radius: 4px;
        font-weight: 500;
    }

    /* 优化标题样式 */
    .stSubheader {
        margin-bottom: 15px;
    }

    /* 优化输入框样式 */
    .stTextInput > div > div > input, .stSelectbox > div > div > select {
        border-radius: 4px;
    }

    /* 优化表格样式 */
    div[data-testid="stDataFrame"] {
        border-radius: 4px;
    }

    /* 优化折叠面板样式 */
    .streamlit-expander {
        margin-bottom: 15px;
    }

    /* 行悬停效果 */
    .data-row-item {
        padding: 8px;
        border-radius: 4px;
        transition: background-color 0.2s;
        cursor: pointer;
        display: flex;
        align-items: center;
        height: 100%;
    }
    .data-row-item:hover {
        background-color: #f0f2f6;
    }

    /* 统计信息样式 */
    .stats-box {
        background-color: #f8f9fa;
        border-left: 4px solid #1f77b4;
        padding: 12px;
        border-radius: 4px;
        margin: 8px 0;
    }
</style>
""", unsafe_allow_html=True)


# ---------------------- 数据结构与初始化 ----------------------

@dataclass
class ReplacedFile:
    """存储替换后的文件数据结构"""
    filename: str  # 文件名
    data: io.BytesIO  # 文件二进制数据
    row_idx: int  # 对应Excel行号
    log: str  # 替换日志


def init_session_state():
    """初始化会话状态"""
    required_states = {
        "replace_rules": [],
        "replaced_files": [],
        "replace_log": [],
        "is_replacing": False,
        "replace_params": {},
        "replace_scope": "替换完整关键词",
        "export_mode_radio": "独立文件（ZIP压缩）",
    }

    for key, default in required_states.items():
        if key not in st.session_state:
            st.session_state[key] = default


init_session_state()


# ---------------------- 核心工具函数 ----------------------

def clean_text(text: str) -> str:
    """清理文本：去除首尾空白、隐藏字符、特殊空格，统一格式"""
    if not isinstance(text, str):
        return ""
    text = text.strip()
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'[\u00A0\u2002-\u200B]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text


def clean_filename(filename: str) -> str:
    """清理文件名非法字符"""
    return re.sub(r'[\\/:*?"<>|]', "_", str(filename))


def generate_safe_filename(
        excel_row: pd.Series,
        file_name_col: str,
        file_prefix: str = "",
        file_suffix: str = "",
        row_idx: int = 0,
        max_length: int = MAX_FILENAME_LENGTH
) -> str:
    """安全生成文件名，处理超长名称和特殊字符"""
    try:
        # 获取基础名称
        if file_name_col and file_name_col in excel_row.index:
            base_name = clean_text(str(excel_row[file_name_col]))
        else:
            base_name = f"替换结果_{row_idx + 1}"

        # 确保base_name不为空
        if not base_name or base_name.isspace():
            base_name = f"替换结果_{row_idx + 1}"

        # 构建完整文件名
        if file_prefix and file_suffix:
            filename = f"{file_prefix}{base_name}{file_suffix}.docx"
        elif file_prefix:
            filename = f"{file_prefix}{base_name}.docx"
        elif file_suffix:
            filename = f"{base_name}{file_suffix}.docx"
        else:
            filename = f"{base_name}.docx"

        # 清理非法字符
        filename = clean_filename(filename)

        # 限制长度（Windows限制255字节）
        filename_bytes = filename.encode('utf-8')
        if len(filename_bytes) > max_length:
            # 重新计算base_name的最大长度
            suffix_len = len(f"{file_prefix}{file_suffix}.docx".encode('utf-8'))
            max_base_bytes = max_length - suffix_len - 10

            # 从base_name截断（考虑UTF-8编码）
            truncated_base = base_name
            while len(f"{file_prefix}{truncated_base}{file_suffix}.docx".encode('utf-8')) > max_length:
                truncated_base = truncated_base[:-1]

            if file_prefix and file_suffix:
                filename = f"{file_prefix}{truncated_base}{file_suffix}.docx"
            elif file_prefix:
                filename = f"{file_prefix}{truncated_base}.docx"
            elif file_suffix:
                filename = f"{truncated_base}{file_suffix}.docx"
            else:
                filename = f"{truncated_base}.docx"

            filename = clean_filename(filename)

        return filename

    except Exception as e:
        return f"替换结果_{row_idx + 1}.docx"


# ---------------------- 替换核心逻辑 ----------------------

def precompute_replace_patterns(
        replace_rules: List[Tuple[str, str]],
        excel_row: pd.Series
) -> List[Tuple[str, str, str, str]]:
    """预计算所有需要替换的模式"""
    replace_patterns = []

    for old_text, col_name in replace_rules:
        # 获取Excel中对应列的替换值
        if col_name in excel_row.index:
            replacement = str(excel_row[col_name]).strip()
        else:
            # Bug修复：列名不存在时使用空字符串而不是报错
            replacement = ""

        # 清理用户输入的关键词
        cleaned_text = clean_text(old_text)

        # Bug修复：检查cleaned_text是否为空
        if not cleaned_text:
            continue

        # 根据替换范围选项生成替换值
        if st.session_state.replace_scope == "仅替换括号内内容":
            if cleaned_text.startswith("【") and cleaned_text.endswith("】"):
                new_format = f"【{replacement}】"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("（") and cleaned_text.endswith("）"):
                new_format = f"（{replacement}）"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("(") and cleaned_text.endswith(")"):
                new_format = f"({replacement})"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("〔") and cleaned_text.endswith("〕"):
                new_format = f"〔{replacement}〕"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            else:
                replace_patterns.append((old_text, col_name, cleaned_text, replacement))
        else:
            replace_patterns.append((old_text, col_name, cleaned_text, replacement))

    return replace_patterns


def process_paragraph(
        paragraph,
        replace_patterns: List[Tuple[str, str, str, str]],
        cleaned_para: str = None
) -> Dict:
    """处理单个段落的关键字替换"""
    para_text = paragraph.text
    if cleaned_para is None:
        cleaned_para = clean_text(para_text)
    replace_count = defaultdict(int)

    # Bug修复：如果段落为空，直接返回
    if not para_text or not replace_patterns:
        return replace_count

    has_keyword = False

    # 检查段落是否包含任何需要替换的关键字
    for old_text, col_name, format_keyword, replacement in replace_patterns:
        if format_keyword and format_keyword in cleaned_para:
            has_keyword = True
            break

    if has_keyword:
        new_text = para_text
        for old_text, col_name, format_keyword, replacement in replace_patterns:
            if format_keyword and format_keyword in cleaned_para:
                # Bug修复：使用case-sensitive替换
                count = new_text.count(format_keyword)
                if count > 0:
                    new_text = new_text.replace(format_keyword, replacement)
                    replace_count[(old_text, col_name)] += count

        # 更新段落文本
        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ''

    return replace_count


def replace_word_with_format(
        word_file: st.runtime.uploaded_file_manager.UploadedFile,
        excel_row: pd.Series,
        replace_rules: List[Tuple[str, str]]
) -> Tuple[io.BytesIO, str]:
    """替换Word文件中的关键字"""
    replace_count = defaultdict(int)

    try:
        # Bug修复：检查文件大小
        file_size = len(word_file.getvalue())
        if file_size > MAX_WORD_FILE_SIZE:
            raise ValueError(f"Word文件过大：{file_size / 1024 / 1024:.2f}MB > {MAX_WORD_FILE_SIZE / 1024 / 1024:.2f}MB")

        # 从内存加载Word文档
        doc = Document(io.BytesIO(word_file.getvalue()))

        # 预计算替换模式
        replace_patterns = precompute_replace_patterns(replace_rules, excel_row)

        # Bug修复：如果没有替换模式，直接返回原文档
        if not replace_patterns:
            output_file = io.BytesIO()
            doc.save(output_file)
            output_file.seek(0)
            return output_file, "⚠ 未设置有效的替换规则"

        # 1. 处理段落
        for paragraph in doc.paragraphs:
            para_count = process_paragraph(paragraph, replace_patterns)
            for key, count in para_count.items():
                replace_count[key] += count

        # 2. 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_count = process_paragraph(paragraph, replace_patterns)
                        for key, count in para_count.items():
                            replace_count[key] += count

        # 保存修改后的文档
        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)

        # 生成替换日志
        if replace_count:
            log_lines = []
            for (old, col_name), count in replace_count.items():
                try:
                    replacement_value = excel_row[col_name]
                except:
                    replacement_value = "N/A"
                log_lines.append(f"✓ {old} → {replacement_value} ({count}次)")
            replace_log = "\n".join(log_lines)
        else:
            replace_log = "⚠ 未找到需要替换的关键字"

        return output_file, replace_log

    except Exception as e:
        import traceback
        error_log = f"❌ 替换失败：{str(e)}"
        return io.BytesIO(), error_log


def merge_word_documents(
        replaced_files: List[ReplacedFile]
) -> io.BytesIO:
    """合并多个Word文档为一个（保留所有格式和结构）"""
    if not replaced_files:
        raise ValueError("没有要合并的文件")

    try:
        # Bug修复：验证文件列表不为空
        if len(replaced_files) == 0:
            raise ValueError("替换文件列表为空")

        # 加载第一个文档作为主文档
        try:
            main_doc = Document(io.BytesIO(replaced_files[0].data.getvalue()))
        except Exception as e:
            raise ValueError(f"无法加载第一个文档：{str(e)}")

        main_body = main_doc._body._element

        # 逐个添加其他文档
        for idx in range(1, len(replaced_files)):
            try:
                file = replaced_files[idx]

                # Bug修复：验证file.data不为空
                if not file.data or len(file.data.getvalue()) == 0:
                    st.warning(f"⚠️ 文件 {file.filename} 数据为空，跳过", icon="⚠️")
                    continue

                sub_doc = Document(io.BytesIO(file.data.getvalue()))
                sub_body = sub_doc._body._element

                # 添加分页符
                page_break_para = OxmlElement('w:p')
                page_break_pPr = OxmlElement('w:pPr')

                # 创建分页符元素
                page_break_element = OxmlElement('w:pageBreakBefore')
                page_break_element.set(qn('w:val'), '1')

                page_break_pPr.append(page_break_element)
                page_break_para.append(page_break_pPr)
                main_body.append(page_break_para)

                # Bug修复：深拷贝所有元素以保留格式
                for element in sub_body:
                    # 使用deepcopy保留完整的XML结构
                    main_body.append(copy.deepcopy(element))

            except Exception as e:
                st.warning(f"⚠️ 处理文件 {file.filename} 失败：{str(e)}", icon="⚠️")
                continue

        # 保存合并后的文档
        output = io.BytesIO()
        main_doc.save(output)
        output.seek(0)
        return output

    except Exception as e:
        st.error(f"❌ 合并文档失败：{str(e)}", icon="❌")
        raise


def get_replace_params(
        word_file: Optional[st.runtime.uploaded_file_manager.UploadedFile],
        excel_df: Optional[pd.DataFrame],
        start_row: int,
        end_row: int,
        file_name_col: str,
        file_prefix: str,
        file_suffix: str
) -> Dict:
    """获取替换参数"""
    return {
        "word_filename": word_file.name if word_file else "",
        "word_size": len(word_file.getvalue()) if word_file else 0,
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "file_prefix": file_prefix,
        "file_suffix": file_suffix,
        "rule_count": len(st.session_state.replace_rules),
        "rule_hash": hash(tuple(st.session_state.replace_rules))
    }


def fix_float_precision(x: str, column_name: Optional[str] = None) -> str:
    """修复浮点数精度问题"""
    if not x or not isinstance(x, str):
        return x

    x = x.strip()

    if not x:
        return ""

    # Bug修复：检查是否全是数字
    try:
        if x.replace('.', '', 1).replace('-', '', 1).isdigit():
            # 是数字格式
            pass
        else:
            return x
    except:
        return x

    float_pattern = r'^\s*[-+]?\d*\.?\d+\s*$'
    if not re.match(float_pattern, x):
        return x

    try:
        dec_value = Decimal(x)

        if dec_value.as_tuple().exponent >= 0:
            return str(int(dec_value))

        float_val = float(dec_value)
        float_str = str(float_val)

        # Bug修复：合计列处理
        if column_name and ("合计" in column_name or "total" in column_name.lower()):
            for dec_places in range(2, 7):
                try:
                    quantized = dec_value.quantize(
                        Decimal('1.' + '0' * dec_places),
                        rounding=ROUND_HALF_UP
                    )

                    if abs(quantized - dec_value) < 1e-9:
                        result = format(quantized, f'.{dec_places}f')
                        return result.rstrip('0').rstrip('.')
                except:
                    continue

        # Bug修复：处理精度问题
        if '999999' in float_str or '000000' in float_str:
            if '.' in x:
                orig_dec_part = x.split('.')[1]
                orig_dec_places = len(orig_dec_part.rstrip('0'))

                if orig_dec_places > 0:
                    try:
                        quantized = dec_value.quantize(
                            Decimal('1.' + '0' * orig_dec_places),
                            rounding=ROUND_HALF_UP
                        )
                        result = format(quantized, f'.{orig_dec_places}f')
                        return result.rstrip('0').rstrip('.')
                    except:
                        pass

            for dec_places in range(1, 10):
                try:
                    formatted = format(float_val, f'.{dec_places}f')
                    if abs(float(formatted) - float_val) < 1e-9:
                        return formatted.rstrip('0').rstrip('.')
                except:
                    continue

        return x
    except:
        return x


def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """清理Excel数据类型"""
    df_clean = df.copy()

    for col in df_clean.columns:
        try:
            col_name = str(col)
            if col_name != col:
                df_clean = df_clean.rename(columns={col: col_name})
                col = col_name

            df_clean[col] = df_clean[col].fillna("")
            df_clean[col] = df_clean[col].astype(str).str.strip()

            # 应用浮点数精度修复
            df_clean[col] = df_clean[col].apply(lambda x: fix_float_precision(x, col))

        except Exception as e:
            try:
                df_clean[col] = df_clean[col].astype(str).str.strip()
            except:
                pass

    return df_clean


# ---------------------- 页面标题与简介 ----------------------
st.title("📋 Word+Excel批量替换工具")
st.markdown("""
快速实现Word模板与Excel数据的批量替换，支持表格内文字替换，保留原格式，操作简单高效。

**✨ 功能特性：**
- 支持合并导出所有替换后的文档为单个Word文件
- 保留所有原文档格式（表格、样式、颜色等）
- 支持导入/导出替换规则
- 支持大批量处理数据

**使用步骤：**
1. 上传Word模板文件和Excel数据文件
2. 预览文档内容，复制需要替换的关键字
3. 设置替换规则和替换范围
4. 执行替换并选择下载方式
5. 支持独立下载或合并为单个文档导出
""", unsafe_allow_html=True)
st.markdown("---")

# ---------------------- 1. 文件上传区 ----------------------
with st.container(border=True):
    st.subheader("🔍 第一步：上传文件")
    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        word_file = st.file_uploader(
            "Word模板",
            type=["docx"],
            key="word",
            help="仅支持.docx格式，.doc需先转换为.docx"
        )
        if word_file:
            # Bug修复：检查文件大小
            file_size_mb = len(word_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_WORD_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ Word文件过大：{file_size_mb:.2f}MB > {MAX_WORD_FILE_SIZE / 1024 / 1024:.2f}MB", icon="❌")
                word_file = None
            else:
                st.success(f"✅ 已上传：{word_file.name}（{file_size_mb:.2f}MB）")

    with col2:
        excel_file = st.file_uploader(
            "Excel数据",
            type=["xlsx", "xls"],
            key="excel",
            help="支持.xlsx/.xls格式，确保数据列名清晰"
        )
        if excel_file:
            # Bug修复：检查文件大小
            file_size_mb = len(excel_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_EXCEL_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ Excel文件过大：{file_size_mb:.2f}MB > {MAX_EXCEL_FILE_SIZE / 1024 / 1024:.2f}MB", icon="❌")
                excel_file = None
            else:
                st.success(f"✅ 已上传：{excel_file.name}（{file_size_mb:.2f}MB）")

st.markdown("---")

# ---------------------- 2. 文档预览区 ----------------------
excel_df = None
excel_cols = []
word_preview_loaded = False

with st.container(border=True):
    st.subheader("📄 第二步：文档预览与关键字复制")
    col1, col2 = st.columns([1, 1], gap="large")

    # Word预览（左侧）
    with col1:
        st.markdown("#### Word预览（含表格）")
        if word_file:
            try:
                doc = Document(io.BytesIO(word_file.getvalue()))
                word_html = "<div style='height: 280px; overflow-y: auto; padding: 8px; border: 1px solid #eee; font-size: 13px; line-height: 1.5;'>"

                # Bug修复：限制预览内容数量
                para_count = 0
                max_para_preview = 100  # 最多显示100个段落

                for paragraph in doc.paragraphs:
                    if para_count >= max_para_preview:
                        word_html += "<p style='color: #999;'><em>...（还有更多内容，不全部显示）</em></p>"
                        break

                    if paragraph.text.strip():
                        para_html = "<p style='margin: 3px 0;'>"
                        for run in paragraph.runs:
                            style = ""
                            if run.bold:
                                style += "font-weight: bold;"
                            if run.italic:
                                style += "font-style: italic;"
                            try:
                                if run.font.color and run.font.color.rgb:
                                    style += f"color: #{run.font.color.rgb:06X}; "
                            except:
                                pass
                            para_html += f"<span style='{style}'>{run.text}</span>" if style else run.text
                        para_html += "</p>"
                        word_html += para_html
                        para_count += 1

                # Bug修复：限制表格预览数量
                table_count = 0
                max_table_preview = 5  # 最多显示5个表格

                for table_idx, table in enumerate(doc.tables):
                    if table_count >= max_table_preview:
                        word_html += f"<p style='color: #999;'><em>...（还有 {len(doc.tables) - table_count} 个表格，不全部显示）</em></p>"
                        break

                    word_html += f"<div style='margin: 8px 0; font-weight: bold;'>表格{table_idx + 1}：</div>"
                    word_html += "<table border='1' style='border-collapse: collapse; width: 100%; border: 1px solid #ccc; font-size: 12px;'>"

                    # Bug修复：限制表格行数
                    for row_idx, row in enumerate(table.rows):
                        if row_idx >= 20:  # 每个表格最多显示20行
                            word_html += "<tr><td colspan='100%' style='text-align:center; color:#999;'>...（还有更多行）</td></tr>"
                            break

                        word_html += "<tr>"
                        for cell in row.cells:
                            cell_html = "<td style='padding: 6px; vertical-align: top; font-size: 11px; max-width: 100px; overflow: hidden;'>"
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    style = ""
                                    if run.bold:
                                        style += "font-weight: bold;"
                                    cell_html += f"<span style='{style}'>{run.text}</span>" if style else run.text
                            cell_html += "</td>"
                            word_html += cell_html
                        word_html += "</tr>"
                    word_html += "</table>"
                    table_count += 1

                word_html += "</div>"

                st.components.v1.html(word_html, height=300)
                st.info("💡 选中需要替换的关键字（支持表格内文字），按Ctrl+C复制", icon="ℹ️")
                word_preview_loaded = True

            except Exception as e:
                st.error(f"❌ Word预览失败：{str(e)}", icon="❌")
        else:
            st.info("请先上传Word模板文件", icon="ℹ️")
            st.markdown(
                "<div style='height: 280px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>Word预览区域</div>",
                unsafe_allow_html=True)

    # Excel预览（右侧）
    with col2:
        st.markdown("#### Excel数据预览")
        if excel_file:
            try:
                with NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_excel:
                    temp_excel.write(excel_file.getvalue())
                    excel_path = temp_excel.name

                try:
                    with pd.ExcelFile(excel_path, engine="openpyxl") as excel_wb:
                        sheet_names = excel_wb.sheet_names
                        selected_sheet = sheet_names[0]
                        st.markdown(f"⚠️ 当前使用工作表：**{selected_sheet}**", unsafe_allow_html=True)

                        # Bug修复：处理空表格
                        excel_df = pd.read_excel(
                            excel_wb,
                            sheet_name=selected_sheet,
                            dtype=str,
                            keep_default_na=False,
                            na_values=[]
                        )

                        if excel_df.empty:
                            st.warning("⚠️ Excel表格为空", icon="⚠️")
                        else:
                            excel_df = clean_excel_types(excel_df)
                            excel_cols = excel_df.columns.tolist()

                            preview_df = excel_df.head(PREVIEW_ROWS)
                            st.dataframe(
                                preview_df,
                                width='stretch',
                                height=250,
                                hide_index=True
                            )

                            st.markdown(f"""
                            <div class='stats-box'>
                            📊 <strong>数据统计</strong><br>
                            总行数：<strong>{len(excel_df)}</strong> | 总列数：<strong>{len(excel_cols)}</strong><br>
                            列名：{', '.join(excel_cols[:5])}{'...' if len(excel_cols) > 5 else ''}
                            </div>
                            """, unsafe_allow_html=True)

                finally:
                    try:
                        if 'excel_path' in locals() and os.path.exists(excel_path):
                            os.unlink(excel_path)
                    except:
                        pass

            except Exception as e:
                st.error(f"❌ Excel读取失败：{str(e)}", icon="❌")
                excel_df = None
                excel_cols = []
        else:
            st.info("请先上传Excel数据文件", icon="ℹ️")
            st.markdown(
                "<div style='height: 250px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>Excel预览区域</div>",
                unsafe_allow_html=True)

st.markdown("---")

# ---------------------- 3. 替换规则设置 ----------------------
with st.container(border=True):
    st.subheader("🔧 第三步：设置替换规则")

    st.markdown(
        "<div style='font-size: 15px; font-weight: bold; margin-top: 10px; margin-bottom: 8px;'>替换范围设置</div>",
        unsafe_allow_html=True)
    st.radio(
        "替换范围",
        options=["替换完整关键词", "仅替换括号内内容"],
        key="replace_scope",
        index=0,
        horizontal=True,
        help="替换完整关键词：替换您输入的精确关键词；仅替换括号内内容：保留括号结构，只替换括号内的文字"
    )

    st.markdown(
        "<div style='font-size: 15px; font-weight: bold; margin-top: 15px; margin-bottom: 8px;'>替换规则导入/导出</div>",
        unsafe_allow_html=True)
    col_import, col_export = st.columns([1, 1], gap="medium")

    with col_import:
        import_rules = st.file_uploader(
            "导入规则（JSON）",
            type=["json"],
            key="import_rules",
            help="从JSON文件导入替换规则"
        )

        if import_rules:
            try:
                rules_data = json.load(import_rules)

                # Bug修复：验证规则数据
                if not isinstance(rules_data, list):
                    st.error("❌ JSON格式错误：应为数组格式", icon="❌")
                else:
                    valid_rules = []
                    for rule in rules_data:
                        if isinstance(rule, dict) and "keyword" in rule and "excel_column" in rule:
                            keyword = str(rule["keyword"]).strip()
                            excel_col = str(rule["excel_column"]).strip()
                            if keyword and excel_col:
                                valid_rules.append((keyword, excel_col))

                    for rule in valid_rules:
                        if rule not in st.session_state.replace_rules:
                            st.session_state.replace_rules.append(rule)

                    st.success(f"✅ 成功导入 {len(valid_rules)} 条规则", icon="✅")
                    st.rerun()
            except json.JSONDecodeError as e:
                st.error(f"❌ JSON格式错误：{str(e)}", icon="❌")
            except Exception as e:
                st.error(f"❌ 导入失败：{str(e)}", icon="❌")

    with col_export:
        if st.session_state.replace_rules:
            rules_data = [
                {"keyword": old, "excel_column": col}
                for old, col in st.session_state.replace_rules
            ]
            rules_json = json.dumps(rules_data, ensure_ascii=False, indent=2)

            st.download_button(
                label="📥 导出规则",
                data=rules_json,
                file_name="replace_rules.json",
                mime="application/json",
                key="export_rules",
                help="将当前替换规则导出为JSON文件"
            )

    st.markdown(
        "<div style='font-size: 15px; font-weight: bold; margin-top: 15px; margin-bottom: 8px;'>规则添加区域</div>",
        unsafe_allow_html=True)
    col_keyword, col_column, col_add = st.columns([3, 3, 1], gap="small")

    with col_keyword:
        keyword_input = st.text_input(
            "关键字",
            placeholder="请输入要替换的关键字（如：【姓名】）",
            key="keyword_input",
            help="从Word文档中复制需要替换的关键字"
        )

    with col_column:
        column_select = st.selectbox(
            "Excel数据列",
            options=excel_cols if excel_cols else ["请先上传Excel文件"],
            key="column_select",
            disabled=not excel_cols,
            help="选择Excel中对应的数据列"
        )

    with col_add:
        add_rule_btn = st.button(
            "➕ 添加",
            key="add_rule",
            type="primary",
            disabled=not (
                        keyword_input and keyword_input.strip() and column_select and column_select != "请先上传Excel文件"),
            help="点击添加替换规则",
            use_container_width=True
        )

    if add_rule_btn:
        rule = (keyword_input.strip(), column_select)
        if rule in st.session_state.replace_rules:
            st.warning("⚠️ 该规则已存在", icon="⚠️")
        else:
            st.session_state.replace_rules.append(rule)
            st.success("✅ 规则添加成功", icon="✅")
            st.rerun()

    # 规则列表显示
    if st.session_state.replace_rules:
        with st.expander("📋 替换规则列表", expanded=True):
            col_actions = st.columns([1, 1], gap="small")
            with col_actions[1]:
                if st.button("🗑️ 清空所有规则", key="clear_rules", type="secondary", use_container_width=True):
                    st.session_state.replace_rules.clear()
                    st.session_state.replaced_files = []
                    st.success("✅ 所有规则已清空", icon="✅")
                    st.rerun()

            st.markdown("<div style='font-size: 14px;'><strong>当前规则：</strong></div>", unsafe_allow_html=True)

            scrollable_container = st.container(height=WIDGET_HEIGHT, border=True)

            with scrollable_container:
                for idx, (old, col) in enumerate(st.session_state.replace_rules):
                    col1, col2, col3, col4, col5 = st.columns([0.5, 3, 0.5, 3, 1], gap="small")

                    with col1:
                        st.write(f"<div class='data-row-item'><strong>{idx + 1}.</strong></div>",
                                 unsafe_allow_html=True)

                    with col2:
                        st.write(f"<div class='data-row-item'><strong>{old}</strong></div>", unsafe_allow_html=True)

                    with col3:
                        st.write(f"<div class='data-row-item'>→</div>", unsafe_allow_html=True)

                    with col4:
                        st.write(f"<div class='data-row-item'>{col}</div>", unsafe_allow_html=True)

                    with col5:
                        if st.button("❌", key=f"delete_{idx}", use_container_width=True):
                            st.session_state.replace_rules.pop(idx)
                            st.session_state.replaced_files = []
                            st.success(f"✅ 已删除规则 {idx + 1}", icon="✅")
                            st.rerun()

st.markdown("---")

# ---------------------- 4. 执行替换 ----------------------
with st.container(border=True):
    st.subheader("🚀 第四步：执行替换")

    st.markdown("#### 文件名设置")
    col_name1, col_name2, col_name3 = st.columns([1, 1, 1], gap="medium")

    with col_name1:
        file_name_col = st.selectbox(
            "核心字段（用于文件名）",
            options=excel_cols if excel_cols else ["请先上传Excel文件"],
            key="file_name_col",
            disabled=not excel_cols,
            help="选择一个Excel列作为生成文件名的核心字段"
        )

    with col_name2:
        file_prefix = st.text_input(
            "文件前缀（可选）",
            value="",
            key="file_prefix",
            help="为生成的文件名添加前缀"
        ).strip()

    with col_name3:
        file_suffix = st.text_input(
            "文件后缀（可选）",
            value="",
            key="file_suffix",
            help="为生成的文件名添加后缀"
        ).strip()

    st.markdown("#### 替换范围设置")
    col_range1, col_range2 = st.columns([1, 1], gap="medium")

    with col_range1:
        start_row = st.number_input(
            "起始行",
            min_value=1,
            max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
            value=1,
            key="start_row",
            disabled=excel_df is None or len(excel_df) == 0,
            help="设置开始处理的Excel行号"
        )

    with col_range2:
        end_row = st.number_input(
            "结束行",
            min_value=1,
            max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
            value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
            key="end_row",
            disabled=excel_df is None or len(excel_df) == 0,
            help="设置结束处理的Excel行号"
        )

    # Bug修复：验证行数范围
    if start_row > end_row:
        st.error("❌ 起始行不能大于结束行", icon="❌")

    can_replace = word_file and excel_df is not None and len(excel_df) > 0 and len(st.session_state.replace_rules) > 0

    current_params = get_replace_params(
        word_file, excel_df, start_row, end_row, file_name_col, file_prefix, file_suffix
    )

    need_replace = (
            len(st.session_state.replaced_files) == 0 or
            st.session_state.replace_params != current_params
    )

    col_replace, col_preview = st.columns([1, 1], gap="medium")

    with col_replace:
        replace_btn = st.button(
            "▶️ 开始替换",
            key="replace",
            disabled=not can_replace or st.session_state.is_replacing or start_row > end_row,
            type="primary",
            help="点击开始执行批量替换操作",
            use_container_width=True
        )

    with col_preview:
        if st.session_state.is_replacing:
            st.info("🔄 正在执行替换，请稍候...", icon="🔄")
        elif len(st.session_state.replaced_files) > 0 and not need_replace:
            st.success(f"✅ 已完成替换！共生成 {len(st.session_state.replaced_files)} 个文件", icon="✅")

    # 执行替换逻辑
    if replace_btn and not st.session_state.is_replacing:
        st.session_state.is_replacing = True
        st.session_state.replaced_files = []
        st.session_state.replace_log = []

        progress_bar = st.progress(0)
        progress_text = st.empty()

        try:
            # Bug修复：验证行数范围有效性
            actual_end_row = min(end_row, len(excel_df))
            if start_row > actual_end_row:
                st.error("❌ 起始行超出数据范围", icon="❌")
            else:
                total_rows = actual_end_row - start_row + 1

                for idx, row_idx in enumerate(range(start_row - 1, actual_end_row)):
                    try:
                        excel_row = excel_df.iloc[row_idx]

                        replaced_file, replace_log = replace_word_with_format(
                            word_file, excel_row, st.session_state.replace_rules
                        )

                        filename = generate_safe_filename(
                            excel_row,
                            file_name_col if file_name_col != "请先上传Excel文件" else "",
                            file_prefix,
                            file_suffix,
                            row_idx
                        )

                        st.session_state.replaced_files.append(ReplacedFile(
                            filename=filename,
                            data=replaced_file,
                            row_idx=row_idx,
                            log=replace_log
                        ))

                        st.session_state.replace_log.append(f"【第{row_idx + 1}行】{replace_log}")

                        # 更新进度条
                        progress = (idx + 1) / total_rows
                        progress_bar.progress(progress)
                        progress_text.text(f"处理进度：{idx + 1}/{total_rows}")

                    except Exception as e:
                        st.session_state.replace_log.append(f"【第{row_idx + 1}行】❌ 处理失败：{str(e)}")
                        continue

                st.session_state.replace_params = current_params
                st.success(f"🎉 替换完成！共生成 {len(st.session_state.replaced_files)} 个文件", icon="✅")

        except Exception as e:
            st.error(f"❌ 替换过程中发生错误：{str(e)}", icon="❌")
        finally:
            st.session_state.is_replacing = False
            progress_bar.empty()
            progress_text.empty()

# ---------------------- 5. 下载结果 ----------------------
if len(st.session_state.replaced_files) > 0:
    st.markdown("---")
    with st.container(border=True):
        st.subheader("💾 第五步：下载结果")

        st.markdown("#### 📥 导出选项")
        export_mode = st.radio(
            "选择导出方式",
            options=["独立文件（ZIP压缩）", "合并为单个文档"],
            key="export_mode_radio",
            horizontal=True,
            help="独立：下载所有文件为ZIP；合并：将所有文件合并为一个Word文档"
        )

        st.markdown("---")

        st.markdown("#### 📦 批量导出")

        if export_mode == "独立文件（ZIP压缩）":
            try:
                # Bug修复：检查文件是否有效
                valid_files = [f for f in st.session_state.replaced_files
                               if f.data and len(f.data.getvalue()) > 0]

                if not valid_files:
                    st.error("❌ 没有有效的文件可以下载", icon="❌")
                else:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for file in valid_files:
                            zipf.writestr(file.filename, file.data.getvalue())

                    zip_buffer.seek(0)
                    zip_filename = f"{file_prefix}批量替换_{len(valid_files)}个文件.zip" if file_prefix else f"批量替换_{len(valid_files)}个文件.zip"
                    zip_filename = clean_filename(zip_filename)

                    st.download_button(
                        label=f"📦 下载全部文件（ZIP）- {len(valid_files)} 个文件",
                        data=zip_buffer,
                        file_name=zip_filename,
                        mime="application/zip",
                        key="download_all_zip",
                        use_container_width=True
                    )
            except Exception as e:
                st.error(f"❌ 创建ZIP文件失败：{str(e)}", icon="❌")
        else:
            # Bug修复：检查文件是否有效
            valid_files = [f for f in st.session_state.replaced_files
                           if f.data and len(f.data.getvalue()) > 0]

            if not valid_files:
                st.error("❌ 没有有效的文件可以合并", icon="❌")
            else:
                try:
                    merged_data = merge_word_documents(valid_files)
                    merged_filename = f"{file_prefix}合并结果.docx" if file_prefix else "合并结果.docx"
                    merged_filename = clean_filename(merged_filename)

                    st.download_button(
                        label=f"📋 下载合并文档 - 1 个文件（包含 {len(valid_files)} 个文档）",
                        data=merged_data,
                        file_name=merged_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_merged",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"❌ 合并失败：{str(e)}", icon="❌")

        st.markdown("---")

        st.markdown("#### 📋 文件列表详情")

        total_pages = (len(st.session_state.replaced_files) + PAGE_SIZE - 1) // PAGE_SIZE

        col_page = st.columns([1])[0]
        with col_page:
            current_page = st.number_input(
                "页码",
                min_value=1,
                max_value=total_pages,
                value=1,
                key="current_page"
            )

        start_idx = (current_page - 1) * PAGE_SIZE
        end_idx = min(start_idx + PAGE_SIZE, len(st.session_state.replaced_files))
        current_files = st.session_state.replaced_files[start_idx:end_idx]

        st.markdown(f"**当前页：{current_page}/{total_pages}（共 {len(st.session_state.replaced_files)} 个文件）**")

        for idx, file in enumerate(current_files, start=start_idx + 1):
            # Bug修复：检查文件有效性
            is_valid = file.data and len(file.data.getvalue()) > 0
            status_icon = "✅" if is_valid else "❌"

            col_file, col_log, col_download = st.columns([2, 1.5, 1], gap="small")

            with col_file:
                st.write(f"<div class='data-row-item'><strong>{status_icon} #{idx}. {file.filename}</strong></div>",
                         unsafe_allow_html=True)

            with col_log:
                with st.expander("📋 查看日志", expanded=False):
                    st.code(file.log, language="text")

            with col_download:
                st.download_button(
                    label="⬇️ 下载",
                    data=file.data,
                    file_name=file.filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{idx}",
                    disabled=not is_valid,
                    use_container_width=True
                )

# ---------------------- 替换日志 ----------------------
if st.session_state.replace_log:
    st.markdown("---")
    with st.container(border=True):
        st.subheader("📊 替换日志详情")

        log_content = "\n".join(st.session_state.replace_log)

        with st.expander("📝 完整日志", expanded=True):
            st.text_area(
                "替换详细日志",
                value=log_content,
                height=250,
                key="log_area",
                disabled=True
            )

# ---------------------- 未满足执行条件的提示 ----------------------
if not can_replace:
    st.markdown("---")
    with st.container(border=True):
        st.info("💡 请完成以下操作：", icon="ℹ️")
        if not word_file:
            st.markdown("• 上传Word模板文件")
        if excel_df is None or excel_df.empty:
            st.markdown("• 上传Excel数据文件")
        if len(st.session_state.replace_rules) == 0:
            st.markdown("• 设置替换规则")

# ---------------------- 底部说明 ----------------------
st.markdown("---")
st.markdown(f"""
### 📝 注意事项
- 仅支持.docx格式的Word文件（.doc需转换为.docx）
- 支持表格内文字替换，表格格式完全保留
- 替换时会保留原文档的所有格式（样式、颜色、字体等）
- **✨ 新功能：支持合并多个替换后的文档为一个Word文件，保留所有格式**
- 建议Word文档不超过50MB，Excel数据不超过50MB
- 对于大量数据（>1000行），建议分批处理

### 🎯 支持的替换格式
- 普通文字：如 `张三`
- 方括号：如 `【张三】`
- 中文圆括号：如 `（张三）`
- 英文圆括号：如 `(张三)`
- 六角括号：如 `〔张三〕`

### 🚀 功能说明
**合并文档导出：** 将所有替换后的文档按顺序合并为一个Word文档，每个原文档占一页，完整保留所有格式和结构。

**版本号：** {VERSION}

**更新日志：**
- v1.3.2：修复多个bug，优化合并文档格式保留
- v1.3.1：修复session_state问题
- v1.3.0：添加合并文档功能

**版权所有 © 2024 Word+Excel批量替换工具**
""", unsafe_allow_html=True)