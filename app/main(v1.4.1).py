# 导入标准库
import os
import sys
import tempfile
from tempfile import NamedTemporaryFile
import warnings
import shutil
import json
import io
import zipfile
import re
import unicodedata
import copy
from datetime import datetime
import hashlib
import base64

# 导入第三方库
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple, Set
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP

# 项目版本信息
VERSION = "v1.4.1"

# 配置常量
PAGE_SIZE = 10
WIDGET_HEIGHT = 300
PREVIEW_ROWS = 30
MAX_FILENAME_LENGTH = 200
MAX_WORD_FILE_SIZE = 50 * 1024 * 1024
MAX_EXCEL_FILE_SIZE = 50 * 1024 * 1024
CACHE_DIR = ".replace_cache"
HISTORY_FILE = ".replace_history.json"
MAX_HISTORY_ITEMS = 20

# 过滤警告
warnings.filterwarnings("ignore", category=UserWarning)

# 环境变量
os.environ["STREAMLIT_VERSION"] = "1.51.0"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# 设置页面配置
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 全局样式优化
st.markdown("""
<style>
    /* 主容器优化 */
    .main {
        padding: 0rem 1rem;
    }

    .stContainer {
        padding: 1rem;
        margin-bottom: 1rem;
        border-radius: 8px;
        background-color: #ffffff;
    }

    /* 按钮样式 */
    .stButton > button {
        border-radius: 6px;
        font-weight: 500;
        padding: 0.5rem 1rem;
        transition: all 0.3s ease;
    }

    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }

    /* 输入框样式 */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select,
    .stNumberInput > div > div > input {
        border-radius: 6px;
        border: 1px solid #e0e0e0;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus,
    .stSelectbox > div > div > select:focus,
    .stNumberInput > div > div > input:focus {
        border-color: #1f77b4;
        box-shadow: 0 0 0 3px rgba(31, 119, 180, 0.1);
    }

    /* 标题样式 */
    h1 {
        padding-bottom: 1rem;
        border-bottom: 3px solid #1f77b4;
    }

    h2 {
        margin-top: 1.5rem;
        margin-bottom: 1rem;
        color: #1f77b4;
    }

    h3 {
        margin-top: 1rem;
        margin-bottom: 0.5rem;
        color: #333;
    }

    .stSubheader {
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e0e0e0;
    }

    /* 展开器样式 */
    .streamlit-expander {
        margin-bottom: 1rem;
        border-radius: 6px;
    }

    /* 数据框样式 */
    div[data-testid="stDataFrame"] {
        border-radius: 6px;
        border: 1px solid #e0e0e0;
    }

    /* 指标卡样式 */
    .stMetric {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 6px;
        border-left: 4px solid #1f77b4;
    }

    /* 行项目样式 */
    .data-row-item {
        padding: 10px;
        border-radius: 6px;
        transition: background-color 0.2s;
        cursor: pointer;
    }

    .data-row-item:hover {
        background-color: #f0f2f6;
    }

    /* 信息框样式 */
    .stats-box {
        background-color: #f0f9ff;
        border-left: 4px solid #0ea5e9;
        padding: 12px;
        border-radius: 6px;
        margin: 8px 0;
    }

    .success-box {
        background-color: #f0fdf4;
        border-left: 4px solid #22c55e;
        padding: 12px;
        border-radius: 6px;
        margin: 8px 0;
    }

    .warning-box {
        background-color: #fffbeb;
        border-left: 4px solid #f59e0b;
        padding: 12px;
        border-radius: 6px;
        margin: 8px 0;
    }

    .error-box {
        background-color: #fef2f2;
        border-left: 4px solid #ef4444;
        padding: 12px;
        border-radius: 6px;
        margin: 8px 0;
    }

    /* 分隔线优化 */
    hr {
        margin: 2rem 0 !important;
        border: none;
        border-top: 2px solid #e0e0e0;
    }

    /* 标签页样式 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }

    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding-top: 10px;
        border-radius: 6px 6px 0 0;
    }

    /* 无线电按钮和复选框 */
    .stRadio > label,
    .stCheckbox > label {
        margin-bottom: 0.5rem;
    }

    /* 文件上传器样式 */
    .stFileUploader {
        border-radius: 6px;
    }

    /* 对齐优化 */
    .element-container {
        margin-bottom: 1rem;
    }

    /* 列间距优化 */
    .stColumn {
        gap: 1rem;
    }

    /* 响应式设计 */
    @media (max-width: 768px) {
        .main {
            padding: 0 0.5rem;
        }

        .stContainer {
            padding: 0.5rem;
        }
    }
</style>
""", unsafe_allow_html=True)


# ---------------------- 数据结构与初始化 ----------------------

@dataclass
class ReplacedFile:
    """存储替换后的文件数据结构"""
    filename: str
    data: io.BytesIO
    row_idx: int
    log: str


@dataclass
class HistoryRecord:
    """历史记录数据结构"""
    timestamp: str
    word_file: str
    excel_file: str
    rules_count: int
    files_generated: int
    status: str


class CacheManager:
    """缓存管理器"""

    def __init__(self):
        self.cache_dir = CACHE_DIR
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)

    def save_rules(self, rules: List[Tuple[str, str]], filename: str):
        """保存规则到缓存"""
        try:
            rules_data = [{"keyword": old, "excel_column": col} for old, col in rules]
            cache_file = os.path.join(self.cache_dir, f"{filename}.json")
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(rules_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"⚠️ 保存缓存失败：{str(e)}", icon="⚠️")

    def load_rules(self, filename: str) -> List[Tuple[str, str]]:
        """加载缓存的规则"""
        try:
            cache_file = os.path.join(self.cache_dir, f"{filename}.json")
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    rules_data = json.load(f)
                    return [(r["keyword"], r["excel_column"]) for r in rules_data]
        except Exception as e:
            st.warning(f"⚠️ 加载缓存失败：{str(e)}", icon="⚠️")
        return []

    def get_cached_rules_list(self) -> List[str]:
        """获取所有缓存的规则文件"""
        try:
            if os.path.exists(self.cache_dir):
                files = [f.replace('.json', '') for f in os.listdir(self.cache_dir) if f.endswith('.json')]
                return sorted(files, reverse=True)
        except:
            pass
        return []


class HistoryManager:
    """历史记录管理器"""

    def __init__(self):
        self.history_file = HISTORY_FILE

    def add_record(self, record: HistoryRecord):
        """添加历史记录"""
        try:
            history = self.load_history()
            history.insert(0, {
                "timestamp": record.timestamp,
                "word_file": record.word_file,
                "excel_file": record.excel_file,
                "rules_count": record.rules_count,
                "files_generated": record.files_generated,
                "status": record.status
            })
            history = history[:MAX_HISTORY_ITEMS]
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"⚠️ 保存历史记录失败：{str(e)}", icon="⚠️")

    def load_history(self) -> List[Dict]:
        """加载历史记录"""
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return []

    def clear_history(self):
        """清除历史记录"""
        try:
            if os.path.exists(self.history_file):
                os.remove(self.history_file)
                st.success("✅ 历史记录已清除", icon="✅")
        except Exception as e:
            st.error(f"❌ 清除历史记录失败：{str(e)}", icon="❌")


def init_session_state():
    """初始化会话状态"""
    required_states = {
        "replace_rules": [],
        "replaced_files": [],
        "replace_log": [],
        "is_replacing": False,
        "replace_params": {},
        "replace_scope": "替换完整关键词",
        "export_mode_radio": "独立文件（ZIP压缩）",
        "show_statistics": False,
        "undo_stack": [],
    }

    for key, default in required_states.items():
        if key not in st.session_state:
            st.session_state[key] = default


init_session_state()


# ---------------------- 核心工具函数 ----------------------

def clean_text(text: str) -> str:
    """清理文本"""
    if not isinstance(text, str):
        return ""
    text = text.strip()
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'[\u00A0\u2002-\u200B]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text


def clean_filename(filename: str) -> str:
    """清理文件名非法字符"""
    return re.sub(r'[\\/:*?"<>|]', "_", str(filename))


def generate_safe_filename(
        excel_row: pd.Series,
        file_name_col: str,
        file_prefix: str = "",
        file_suffix: str = "",
        row_idx: int = 0,
        max_length: int = MAX_FILENAME_LENGTH
) -> str:
    """安全生成文件名"""
    try:
        if file_name_col and file_name_col in excel_row.index:
            base_name = clean_text(str(excel_row[file_name_col]))
        else:
            base_name = f"替换结果_{row_idx + 1}"

        if not base_name or base_name.isspace():
            base_name = f"替换结果_{row_idx + 1}"

        if file_prefix and file_suffix:
            filename = f"{file_prefix}{base_name}{file_suffix}.docx"
        elif file_prefix:
            filename = f"{file_prefix}{base_name}.docx"
        elif file_suffix:
            filename = f"{base_name}{file_suffix}.docx"
        else:
            filename = f"{base_name}.docx"

        filename = clean_filename(filename)

        filename_bytes = filename.encode('utf-8')
        if len(filename_bytes) > max_length:
            suffix_len = len(f"{file_prefix}{file_suffix}.docx".encode('utf-8'))
            max_base_bytes = max_length - suffix_len - 10

            truncated_base = base_name
            while len(f"{file_prefix}{truncated_base}{file_suffix}.docx".encode('utf-8')) > max_length:
                truncated_base = truncated_base[:-1]

            if file_prefix and file_suffix:
                filename = f"{file_prefix}{truncated_base}{file_suffix}.docx"
            elif file_prefix:
                filename = f"{file_prefix}{truncated_base}.docx"
            elif file_suffix:
                filename = f"{truncated_base}{file_suffix}.docx"
            else:
                filename = f"{truncated_base}.docx"

            filename = clean_filename(filename)

        return filename

    except Exception as e:
        return f"替换结果_{row_idx + 1}.docx"


def precompute_replace_patterns(
        replace_rules: List[Tuple[str, str]],
        excel_row: pd.Series
) -> List[Tuple[str, str, str, str]]:
    """预计算替换模式"""
    replace_patterns = []

    for old_text, col_name in replace_rules:
        if col_name in excel_row.index:
            replacement = str(excel_row[col_name]).strip()
        else:
            replacement = ""

        cleaned_text = clean_text(old_text)

        if not cleaned_text:
            continue

        if st.session_state.replace_scope == "仅替换括号内内容":
            if cleaned_text.startswith("【") and cleaned_text.endswith("】"):
                new_format = f"【{replacement}】"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("（") and cleaned_text.endswith("）"):
                new_format = f"（{replacement}）"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("(") and cleaned_text.endswith(")"):
                new_format = f"({replacement})"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("〔") and cleaned_text.endswith("〕"):
                new_format = f"〔{replacement}〕"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            else:
                replace_patterns.append((old_text, col_name, cleaned_text, replacement))
        else:
            replace_patterns.append((old_text, col_name, cleaned_text, replacement))

    return replace_patterns


def process_paragraph(
        paragraph,
        replace_patterns: List[Tuple[str, str, str, str]],
        cleaned_para: str = None
) -> Dict:
    """处理段落替换"""
    para_text = paragraph.text
    if cleaned_para is None:
        cleaned_para = clean_text(para_text)
    replace_count = defaultdict(int)

    if not para_text or not replace_patterns:
        return replace_count

    has_keyword = False

    for old_text, col_name, format_keyword, replacement in replace_patterns:
        if format_keyword and format_keyword in cleaned_para:
            has_keyword = True
            break

    if has_keyword:
        new_text = para_text
        for old_text, col_name, format_keyword, replacement in replace_patterns:
            if format_keyword and format_keyword in cleaned_para:
                count = new_text.count(format_keyword)
                if count > 0:
                    new_text = new_text.replace(format_keyword, replacement)
                    replace_count[(old_text, col_name)] += count

        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ''

    return replace_count


def replace_word_with_format(
        word_file: st.runtime.uploaded_file_manager.UploadedFile,
        excel_row: pd.Series,
        replace_rules: List[Tuple[str, str]]
) -> Tuple[io.BytesIO, str]:
    """替换Word文件"""
    replace_count = defaultdict(int)

    try:
        file_size = len(word_file.getvalue())
        if file_size > MAX_WORD_FILE_SIZE:
            raise ValueError(f"Word文件过大：{file_size / 1024 / 1024:.2f}MB > {MAX_WORD_FILE_SIZE / 1024 / 1024:.2f}MB")

        doc = Document(io.BytesIO(word_file.getvalue()))

        replace_patterns = precompute_replace_patterns(replace_rules, excel_row)

        if not replace_patterns:
            output_file = io.BytesIO()
            doc.save(output_file)
            output_file.seek(0)
            return output_file, "⚠ 未设置有效的替换规则"

        for paragraph in doc.paragraphs:
            para_count = process_paragraph(paragraph, replace_patterns)
            for key, count in para_count.items():
                replace_count[key] += count

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_count = process_paragraph(paragraph, replace_patterns)
                        for key, count in para_count.items():
                            replace_count[key] += count

        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)

        if replace_count:
            log_lines = []
            for (old, col_name), count in replace_count.items():
                try:
                    replacement_value = excel_row[col_name]
                except:
                    replacement_value = "N/A"
                log_lines.append(f"✓ {old} → {replacement_value} ({count}次)")
            replace_log = "\n".join(log_lines)
        else:
            replace_log = "⚠ 未找到需要替换的关键字"

        return output_file, replace_log

    except Exception as e:
        import traceback
        error_log = f"❌ 替换失败：{str(e)}"
        return io.BytesIO(), error_log


def merge_word_documents(
        replaced_files: List[ReplacedFile]
) -> io.BytesIO:
    """合并Word文档（保留格式）"""
    if not replaced_files:
        raise ValueError("没有要合并的文件")

    try:
        if len(replaced_files) == 0:
            raise ValueError("替换文件列表为空")

        try:
            main_doc = Document(io.BytesIO(replaced_files[0].data.getvalue()))
        except Exception as e:
            raise ValueError(f"无法加载第一个文档：{str(e)}")

        main_body = main_doc._body._element

        for idx in range(1, len(replaced_files)):
            try:
                file = replaced_files[idx]

                if not file.data or len(file.data.getvalue()) == 0:
                    st.warning(f"⚠️ 文件 {file.filename} 数据为空，跳过", icon="⚠️")
                    continue

                sub_doc = Document(io.BytesIO(file.data.getvalue()))
                sub_body = sub_doc._body._element

                page_break_para = OxmlElement('w:p')
                page_break_pPr = OxmlElement('w:pPr')

                page_break_element = OxmlElement('w:pageBreakBefore')
                page_break_element.set(qn('w:val'), '1')

                page_break_pPr.append(page_break_element)
                page_break_para.append(page_break_pPr)
                main_body.append(page_break_para)

                for element in sub_body:
                    main_body.append(copy.deepcopy(element))

            except Exception as e:
                st.warning(f"⚠️ 处理文件 {file.filename} 失败：{str(e)}", icon="⚠️")
                continue

        output = io.BytesIO()
        main_doc.save(output)
        output.seek(0)
        return output

    except Exception as e:
        st.error(f"❌ 合并文档失败：{str(e)}", icon="❌")
        raise


def get_replace_params(
        word_file: Optional[st.runtime.uploaded_file_manager.UploadedFile],
        excel_df: Optional[pd.DataFrame],
        start_row: int,
        end_row: int,
        file_name_col: str,
        file_prefix: str,
        file_suffix: str
) -> Dict:
    """获取替换参数"""
    return {
        "word_filename": word_file.name if word_file else "",
        "word_size": len(word_file.getvalue()) if word_file else 0,
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "file_prefix": file_prefix,
        "file_suffix": file_suffix,
        "rule_count": len(st.session_state.replace_rules),
        "rule_hash": hash(tuple(st.session_state.replace_rules))
    }


def fix_float_precision(x: str, column_name: Optional[str] = None) -> str:
    """修复浮点数精度"""
    if not x or not isinstance(x, str):
        return x

    x = x.strip()

    if not x:
        return ""

    try:
        if x.replace('.', '', 1).replace('-', '', 1).isdigit():
            pass
        else:
            return x
    except:
        return x

    float_pattern = r'^\s*[-+]?\d*\.?\d+\s*$'
    if not re.match(float_pattern, x):
        return x

    try:
        dec_value = Decimal(x)

        if dec_value.as_tuple().exponent >= 0:
            return str(int(dec_value))

        float_val = float(dec_value)
        float_str = str(float_val)

        if column_name and ("合计" in column_name or "total" in column_name.lower()):
            for dec_places in range(2, 7):
                try:
                    quantized = dec_value.quantize(
                        Decimal('1.' + '0' * dec_places),
                        rounding=ROUND_HALF_UP
                    )

                    if abs(quantized - dec_value) < 1e-9:
                        result = format(quantized, f'.{dec_places}f')
                        return result.rstrip('0').rstrip('.')
                except:
                    continue

        if '999999' in float_str or '000000' in float_str:
            if '.' in x:
                orig_dec_part = x.split('.')[1]
                orig_dec_places = len(orig_dec_part.rstrip('0'))

                if orig_dec_places > 0:
                    try:
                        quantized = dec_value.quantize(
                            Decimal('1.' + '0' * orig_dec_places),
                            rounding=ROUND_HALF_UP
                        )
                        result = format(quantized, f'.{orig_dec_places}f')
                        return result.rstrip('0').rstrip('.')
                    except:
                        pass

            for dec_places in range(1, 10):
                try:
                    formatted = format(float_val, f'.{dec_places}f')
                    if abs(float(formatted) - float_val) < 1e-9:
                        return formatted.rstrip('0').rstrip('.')
                except:
                    continue

        return x
    except:
        return x


def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """清理Excel数据类型"""
    df_clean = df.copy()

    for col in df_clean.columns:
        try:
            col_name = str(col)
            if col_name != col:
                df_clean = df_clean.rename(columns={col: col_name})
                col = col_name

            df_clean[col] = df_clean[col].fillna("")
            df_clean[col] = df_clean[col].astype(str).str.strip()

            df_clean[col] = df_clean[col].apply(lambda x: fix_float_precision(x, col))

        except Exception as e:
            try:
                df_clean[col] = df_clean[col].astype(str).str.strip()
            except:
                pass

    return df_clean


def get_file_hash(file_data: bytes) -> str:
    """获取文件哈希值"""
    return hashlib.md5(file_data).hexdigest()[:8]


def export_statistics_to_csv(replaced_files: List[ReplacedFile]) -> str:
    """导出统计数据到CSV"""
    try:
        data = []
        for idx, file in enumerate(replaced_files, 1):
            data.append({
                "序号": idx,
                "文件名": file.filename,
                "Excel行号": file.row_idx + 1,
                "替换日志": file.log.replace("\n", "; ")
            })

        df = pd.DataFrame(data)
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
        return csv_buffer.getvalue()
    except Exception as e:
        st.error(f"❌ 导出统计失败：{str(e)}", icon="❌")
        return ""


def get_keyword_statistics(replace_rules: List[Tuple[str, str]],
                           replaced_files: List[ReplacedFile]) -> Dict:
    """获取关键字替换统计"""
    stats = {}
    for keyword, _ in replace_rules:
        stats[keyword] = 0

    for file in replaced_files:
        for keyword, _ in replace_rules:
            if f"✓ {keyword}" in file.log:
                pattern = f"✓ {re.escape(keyword)}.*?\((\d+)次\)"
                matches = re.findall(pattern, file.log)
                if matches:
                    stats[keyword] += int(matches[0])

    return stats


# 创建缓存和历史管理器
cache_manager = CacheManager()
history_manager = HistoryManager()

# ---------------------- 页面标题与简介 ----------------------
col_title1, col_title2 = st.columns([8, 2])
with col_title1:
    st.title("📋 Word+Excel批量替换工具")
with col_title2:
    st.markdown(f"<div style='text-align: right; padding-top: 10px;'><small>v{VERSION}</small></div>",
                unsafe_allow_html=True)

st.markdown("""
快速实现Word模板与Excel数据的批量替换，支持表格内文字替换，保留原格式，操作简单高效。

**✨ 核心功能：** 批量替换 | 格式保留 | 文档合并 | 规则管理 | 统计分析
""", unsafe_allow_html=True)

# 创建标签页
tab1, tab2, tab3, tab4 = st.tabs(["🚀 快速开始", "📚 规则管理", "💾 下载结果", "⚙️ 工具设置"])

with tab1:
    st.markdown("### 替换工作流程")

    # ==================== 第一步：文件上传 ====================
    st.subheader("📤 步骤1：上传文件")

    col_upload1, col_upload2 = st.columns([1, 1], gap="medium")

    with col_upload1:
        st.markdown("**Word 模板文件**")
        word_file = st.file_uploader(
            "选择Word文件",
            type=["docx"],
            key="word",
            help="仅支持.docx格式"
        )
        if word_file:
            file_size_mb = len(word_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_WORD_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ 文件过大：{file_size_mb:.2f}MB", icon="❌")
                word_file = None
            else:
                file_hash = get_file_hash(word_file.getvalue())
                st.markdown(f"""
                <div class='success-box'>
                <strong>✅ 文件已上传</strong><br>
                📄 {word_file.name}<br>
                📊 大小：{file_size_mb:.2f}MB<br>
                🔐 哈希：{file_hash}
                </div>
                """, unsafe_allow_html=True)

    with col_upload2:
        st.markdown("**Excel 数据文件**")
        excel_file = st.file_uploader(
            "选择Excel文件",
            type=["xlsx", "xls"],
            key="excel",
            help="支持.xlsx/.xls格式"
        )
        if excel_file:
            file_size_mb = len(excel_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_EXCEL_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ 文件过大：{file_size_mb:.2f}MB", icon="❌")
                excel_file = None
            else:
                file_hash = get_file_hash(excel_file.getvalue())
                st.markdown(f"""
                <div class='success-box'>
                <strong>✅ 文件已上传</strong><br>
                📄 {excel_file.name}<br>
                📊 大小：{file_size_mb:.2f}MB<br>
                🔐 哈希：{file_hash}
                </div>
                """, unsafe_allow_html=True)

    st.markdown("---")

    # ==================== 第二步：文档预览 ====================
    st.subheader("👀 步骤2：预览文档内容")

    col_preview1, col_preview2 = st.columns([1, 1], gap="medium")

    excel_df = None
    excel_cols = []

    with col_preview1:
        st.markdown("**Word 文档预览**")
        if word_file:
            try:
                doc = Document(io.BytesIO(word_file.getvalue()))
                word_html = "<div style='height: 250px; overflow-y: auto; padding: 12px; border: 1px solid #e0e0e0; border-radius: 6px; font-size: 13px; line-height: 1.6; background-color: #f9f9f9;'>"

                para_count = 0
                max_para_preview = 100

                for paragraph in doc.paragraphs:
                    if para_count >= max_para_preview:
                        word_html += "<p style='color: #999;'><em>...（还有更多内容）</em></p>"
                        break

                    if paragraph.text.strip():
                        para_html = "<p style='margin: 4px 0;'>"
                        for run in paragraph.runs:
                            style = ""
                            if run.bold:
                                style += "font-weight: bold;"
                            if run.italic:
                                style += "font-style: italic;"
                            try:
                                if run.font.color and run.font.color.rgb:
                                    style += f"color: #{run.font.color.rgb:06X}; "
                            except:
                                pass
                            para_html += f"<span style='{style}'>{run.text}</span>" if style else run.text
                        para_html += "</p>"
                        word_html += para_html
                        para_count += 1

                table_count = 0
                max_table_preview = 3

                for table_idx, table in enumerate(doc.tables):
                    if table_count >= max_table_preview:
                        word_html += f"<p style='color: #999;'><em>...（还有{len(doc.tables) - table_count}个表格）</em></p>"
                        break

                    word_html += f"<div style='margin: 8px 0; font-weight: bold; color: #1f77b4;'>📊 表格{table_idx + 1}：</div>"
                    word_html += "<table border='1' style='border-collapse: collapse; width: 100%; border: 1px solid #ccc; font-size: 12px;'>"

                    for row_idx, row in enumerate(table.rows):
                        if row_idx >= 15:
                            word_html += "<tr><td colspan='100%' style='text-align:center; color:#999;'>...</td></tr>"
                            break

                        word_html += "<tr>"
                        for cell in row.cells:
                            cell_html = "<td style='padding: 6px; vertical-align: top; font-size: 11px;'>"
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    cell_html += run.text
                            cell_html += "</td>"
                            word_html += cell_html
                        word_html += "</tr>"
                    word_html += "</table>"
                    table_count += 1

                word_html += "</div>"

                st.components.v1.html(word_html, height=280)
                st.caption("💡 按Ctrl+C复制需要替换的关键字")

            except Exception as e:
                st.error(f"❌ 预览失败：{str(e)}", icon="❌")
        else:
            st.info("📁 请先上传Word文件", icon="ℹ️")

    with col_preview2:
        st.markdown("**Excel 数据预览**")
        if excel_file:
            try:
                with NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_excel:
                    temp_excel.write(excel_file.getvalue())
                    excel_path = temp_excel.name

                try:
                    with pd.ExcelFile(excel_path, engine="openpyxl") as excel_wb:
                        sheet_names = excel_wb.sheet_names
                        selected_sheet = sheet_names[0]
                        st.caption(f"📋 工作表：**{selected_sheet}**")

                        excel_df = pd.read_excel(
                            excel_wb,
                            sheet_name=selected_sheet,
                            dtype=str,
                            keep_default_na=False,
                            na_values=[]
                        )

                        if excel_df.empty:
                            st.warning("⚠️ Excel表格为空", icon="⚠️")
                        else:
                            excel_df = clean_excel_types(excel_df)
                            excel_cols = excel_df.columns.tolist()

                            preview_df = excel_df.head(PREVIEW_ROWS)
                            st.dataframe(
                                preview_df,
                                width='stretch',
                                height=250,
                                use_container_width=True,
                                hide_index=True
                            )

                            col_stat1, col_stat2, col_stat3 = st.columns(3)
                            with col_stat1:
                                st.metric("📊 行数", len(excel_df))
                            with col_stat2:
                                st.metric("📋 列数", len(excel_cols))
                            with col_stat3:
                                st.metric("💾 文件大小", f"{len(excel_file.getvalue()) / 1024:.1f}KB")

                finally:
                    try:
                        if os.path.exists(excel_path):
                            os.unlink(excel_path)
                    except:
                        pass

            except Exception as e:
                st.error(f"❌ 读取失败：{str(e)}", icon="❌")
        else:
            st.info("📁 请先上传Excel文件", icon="ℹ️")

    st.markdown("---")

    # ==================== 第三步：替换范围 ====================
    st.subheader("⚙️ 步骤3：配置替换参数")

    col_config1, col_config2, col_config3 = st.columns([1, 1, 1], gap="medium")

    with col_config1:
        st.markdown("**文件命名设置**")
        file_name_col = st.selectbox(
            "核心字段",
            options=excel_cols if excel_cols else ["请先上传Excel文件"],
            key="file_name_col",
            disabled=not excel_cols,
            help="用于生成文件名"
        )

    with col_config2:
        st.markdown("**文件前后缀**")
        col_prefix, col_suffix = st.columns(2, gap="small")
        with col_prefix:
            file_prefix = st.text_input(
                "前缀",
                value="",
                key="file_prefix",
                placeholder="可选",
                max_chars=20
            ).strip()
        with col_suffix:
            file_suffix = st.text_input(
                "后缀",
                value="",
                key="file_suffix",
                placeholder="可选",
                max_chars=20
            ).strip()

    with col_config3:
        st.markdown("**替换数据范围**")
        col_start, col_end = st.columns(2, gap="small")
        with col_start:
            start_row = st.number_input(
                "起始行",
                min_value=1,
                max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
                value=1,
                key="start_row",
                disabled=excel_df is None or len(excel_df) == 0
            )
        with col_end:
            end_row = st.number_input(
                "结束行",
                min_value=1,
                max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
                value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
                key="end_row",
                disabled=excel_df is None or len(excel_df) == 0
            )

    if start_row > end_row:
        st.error("❌ 起始行不能大于结束行", icon="❌")

    st.markdown("---")

    # ==================== 第四步：执行替换 ====================
    st.subheader("🚀 步骤4：执行批量替换")

    can_replace = word_file and excel_df is not None and len(excel_df) > 0 and len(st.session_state.replace_rules) > 0

    current_params = get_replace_params(
        word_file, excel_df, start_row, end_row, file_name_col, file_prefix, file_suffix
    )

    need_replace = (
            len(st.session_state.replaced_files) == 0 or
            st.session_state.replace_params != current_params
    )

    col_execute1, col_execute2, col_execute3 = st.columns([2, 2, 2], gap="medium")

    with col_execute1:
        replace_btn = st.button(
            "▶️ 开始替换",
            key="replace",
            disabled=not can_replace or st.session_state.is_replacing or start_row > end_row,
            type="primary",
            use_container_width=True
        )

    with col_execute2:
        if st.session_state.is_replacing:
            st.info("🔄 替换进行中...", icon="🔄")
        elif len(st.session_state.replaced_files) > 0 and not need_replace:
            st.success(f"✅ {len(st.session_state.replaced_files)} 个文件已生成", icon="✅")

    with col_execute3:
        pass

    if replace_btn and not st.session_state.is_replacing:
        st.session_state.is_replacing = True
        st.session_state.replaced_files = []
        st.session_state.replace_log = []

        progress_bar = st.progress(0)
        progress_text = st.empty()

        try:
            actual_end_row = min(end_row, len(excel_df))
            if start_row > actual_end_row:
                st.error("❌ 起始行超出数据范围", icon="❌")
            else:
                total_rows = actual_end_row - start_row + 1

                for idx, row_idx in enumerate(range(start_row - 1, actual_end_row)):
                    try:
                        excel_row = excel_df.iloc[row_idx]

                        replaced_file, replace_log = replace_word_with_format(
                            word_file, excel_row, st.session_state.replace_rules
                        )

                        filename = generate_safe_filename(
                            excel_row,
                            file_name_col if file_name_col != "请先上传Excel文件" else "",
                            file_prefix,
                            file_suffix,
                            row_idx
                        )

                        st.session_state.replaced_files.append(ReplacedFile(
                            filename=filename,
                            data=replaced_file,
                            row_idx=row_idx,
                            log=replace_log
                        ))

                        st.session_state.replace_log.append(f"【第{row_idx + 1}行】{replace_log}")

                        progress = (idx + 1) / total_rows
                        progress_bar.progress(progress)
                        progress_text.text(f"处理进度：{idx + 1}/{total_rows}")

                    except Exception as e:
                        st.session_state.replace_log.append(f"【第{row_idx + 1}行】❌ 失败：{str(e)}")
                        continue

                st.session_state.replace_params = current_params
                st.success(f"🎉 完成！共生成 {len(st.session_state.replaced_files)} 个文件", icon="✅")

                history_record = HistoryRecord(
                    timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    word_file=word_file.name,
                    excel_file=excel_file.name,
                    rules_count=len(st.session_state.replace_rules),
                    files_generated=len(st.session_state.replaced_files),
                    status="success"
                )
                history_manager.add_record(history_record)

        except Exception as e:
            st.error(f"❌ 错误：{str(e)}", icon="❌")
            history_record = HistoryRecord(
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                word_file=word_file.name if word_file else "N/A",
                excel_file=excel_file.name if excel_file else "N/A",
                rules_count=len(st.session_state.replace_rules),
                files_generated=0,
                status="failed"
            )
            history_manager.add_record(history_record)
        finally:
            st.session_state.is_replacing = False
            progress_bar.empty()
            progress_text.empty()

    # 显示替换日志
    if st.session_state.replace_log:
        st.markdown("---")
        with st.expander("📝 替换日志", expanded=False):
            log_content = "\n".join(st.session_state.replace_log)
            st.text_area(
                "日志内容",
                value=log_content,
                height=200,
                key="log_area",
                disabled=True,
                label_visibility="collapsed"
            )

    # 未满足条件提示
    if not can_replace:
        st.markdown("---")
        st.info("💡 需要：1️⃣ Word文件 2️⃣ Excel文件 3️⃣ 替换规则", icon="ℹ️")

# ==================== 标签页2：规则管理 ====================
with tab2:
    st.subheader("📋 替换规则管理")

    # 替换范围选择
    st.markdown("### 替换范围选择")
    col_scope1, col_scope2 = st.columns(2, gap="medium")
    with col_scope1:
        st.radio(
            "替换模式",
            options=["替换完整关键词", "仅替换括号内内容"],
            key="replace_scope",
            index=0,
            horizontal=False,
            help="完整关键词：精确替换\n括号内容：保留括号结构"
        )

    with col_scope2:
        st.markdown("**模式说明**")
        if st.session_state.replace_scope == "替换完整关键词":
            st.markdown("""
            ✓ 直接替换整个关键词

            **示例：**
            - 【张三】→ 【李四】
            - （2024年）→ （2025年）
            """)
        else:
            st.markdown("""
            ✓ 保留括号，只替换内容

            **示例：**
            - 【张三】→ 【李四】
            - （张三）→ （李四）
            """)

    st.markdown("---")

    # 规则导入导出缓存
    st.markdown("### 规则导入/导出/缓存")

    col_imp1, col_imp2, col_imp3 = st.columns([1, 1, 1], gap="medium")

    with col_imp1:
        st.markdown("**导入规则**")
        import_rules = st.file_uploader(
            "上传JSON文件",
            type=["json"],
            key="import_rules",
            help="导入保存的规则"
        )

        if import_rules:
            try:
                rules_data = json.load(import_rules)

                if not isinstance(rules_data, list):
                    st.error("❌ JSON格式错误：应为数组", icon="❌")
                else:
                    valid_rules = []
                    for rule in rules_data:
                        if isinstance(rule, dict) and "keyword" in rule and "excel_column" in rule:
                            keyword = str(rule["keyword"]).strip()
                            excel_col = str(rule["excel_column"]).strip()
                            if keyword and excel_col:
                                valid_rules.append((keyword, excel_col))

                    st.session_state.undo_stack.append(st.session_state.replace_rules.copy())

                    for rule in valid_rules:
                        if rule not in st.session_state.replace_rules:
                            st.session_state.replace_rules.append(rule)

                    st.success(f"✅ 导入 {len(valid_rules)} 条规则", icon="✅")
                    st.rerun()
            except json.JSONDecodeError as e:
                st.error(f"❌ JSON错误：{str(e)}", icon="❌")
            except Exception as e:
                st.error(f"❌ 导入失败：{str(e)}", icon="❌")

    with col_imp2:
        st.markdown("**导出规则**")
        if st.session_state.replace_rules:
            rules_data = [
                {"keyword": old, "excel_column": col}
                for old, col in st.session_state.replace_rules
            ]
            rules_json = json.dumps(rules_data, ensure_ascii=False, indent=2)

            st.download_button(
                label="📥 导出JSON",
                data=rules_json,
                file_name="rules.json",
                mime="application/json",
                key="export_rules",
                use_container_width=True
            )
        else:
            st.info("📁 无规则可导出", icon="ℹ️")

    with col_imp3:
        st.markdown("**缓存规则**")
        if st.session_state.replace_rules:
            if st.button("💾 保存到缓存", key="save_cache", use_container_width=True):
                cache_name = f"rules_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                cache_manager.save_rules(st.session_state.replace_rules, cache_name)
                st.success("✅ 已保存", icon="✅")
        else:
            st.info("📁 无规则可保存", icon="ℹ️")

    st.markdown("---")

    # 快速加载缓存
    st.markdown("### 快速加载缓存")
    cached_rules = cache_manager.get_cached_rules_list()

    if cached_rules:
        col_load1, col_load2, col_load3 = st.columns([2, 1, 1], gap="medium")
        with col_load1:
            selected_cache = st.selectbox(
                "选择缓存",
                options=cached_rules,
                key="select_cache",
                label_visibility="collapsed"
            )
        with col_load2:
            if st.button("📂 加载", key="load_cache", use_container_width=True):
                loaded_rules = cache_manager.load_rules(selected_cache)
                if loaded_rules:
                    st.session_state.replace_rules = loaded_rules
                    st.success(f"✅ 加载 {len(loaded_rules)} 条规则", icon="✅")
                    st.rerun()
        with col_load3:
            if st.button("🗑️ 删除", key="delete_cache", use_container_width=True):
                try:
                    cache_file = os.path.join(cache_manager.cache_dir, f"{selected_cache}.json")
                    if os.path.exists(cache_file):
                        os.remove(cache_file)
                        st.success("✅ 已删除", icon="✅")
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ 删除失败：{str(e)}", icon="❌")
    else:
        st.info("📁 暂无缓存规则", icon="ℹ️")

    st.markdown("---")

    # 规则添加
    st.markdown("### 添加新规则")
    col_add1, col_add2, col_add3 = st.columns([2, 2, 1], gap="medium")

    with col_add1:
        keyword_input = st.text_input(
            "关键字",
            placeholder="如：【姓名】",
            key="keyword_input",
            help="从Word文档复制"
        )

    with col_add2:
        if excel_cols:
            column_select = st.selectbox(
                "Excel列",
                options=excel_cols,
                key="column_select"
            )
        else:
            st.info("请先上传Excel文件", icon="ℹ️")
            column_select = None

    with col_add3:
        add_rule_btn = st.button(
            "➕ 添加",
            key="add_rule",
            type="primary",
            disabled=not (keyword_input and keyword_input.strip() and column_select),
            use_container_width=True
        )

    if add_rule_btn and column_select:
        rule = (keyword_input.strip(), column_select)
        if rule in st.session_state.replace_rules:
            st.warning("⚠️ 规则已存在", icon="⚠️")
        else:
            st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
            st.session_state.replace_rules.append(rule)
            st.success("✅ 规则已添加", icon="✅")
            st.rerun()

    st.markdown("---")

    # 规则列表
    st.markdown("### 当前规则列表")

    if st.session_state.replace_rules:
        col_action1, col_action2, col_action3 = st.columns([1, 1, 1], gap="medium")

        with col_action1:
            if st.session_state.undo_stack:
                if st.button("↶ 撤销", key="undo", use_container_width=True):
                    st.session_state.replace_rules = st.session_state.undo_stack.pop()
                    st.success("✅ 已撤销", icon="✅")
                    st.rerun()

        with col_action2:
            pass

        with col_action3:
            if st.button("🗑️ 清空所有", key="clear_rules", type="secondary", use_container_width=True):
                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                st.session_state.replace_rules.clear()
                st.session_state.replaced_files = []
                st.success("✅ 已清空", icon="✅")
                st.rerun()

        # 规则表格显示
        rule_data = []
        for idx, (old, col) in enumerate(st.session_state.replace_rules, 1):
            rule_data.append({
                "序号": idx,
                "关键字": old,
                "Excel列": col,
                "操作": f"❌ {idx}"  # 占位符
            })

        rule_df = pd.DataFrame(rule_data)
        st.dataframe(rule_df, use_container_width=True, hide_index=True)

        # 删除按钮（单独放在下方）
        st.markdown("**删除规则**")
        col_del1, col_del2, col_del3 = st.columns(3, gap="small")

        rules_to_delete = len(st.session_state.replace_rules)
        if rules_to_delete <= 3:
            for idx in range(rules_to_delete):
                with st.columns([1, 1, 1])[idx]:
                    if st.button(f"删除规则 {idx + 1}", key=f"delete_{idx}", use_container_width=True):
                        st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                        st.session_state.replace_rules.pop(idx)
                        st.session_state.replaced_files = []
                        st.success(f"✅ 已删除规则 {idx + 1}", icon="✅")
                        st.rerun()
        else:
            # 超过3个规则，用可滚动的容器
            with st.container(height=200, border=True):
                for idx in range(rules_to_delete):
                    if st.button(f"删除规则 {idx + 1}", key=f"delete_{idx}", use_container_width=True):
                        st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                        st.session_state.replace_rules.pop(idx)
                        st.session_state.replaced_files = []
                        st.success(f"✅ 已删除规则 {idx + 1}", icon="✅")
                        st.rerun()
    else:
        st.info("📁 暂无规则，请添加规则后开始替换", icon="ℹ️")

# ==================== 标签页3：下载结果 ====================
with tab3:
    st.subheader("💾 下载替换结果")

    if len(st.session_state.replaced_files) > 0:

        # 导出选项
        st.markdown("### 导出方式选择")
        export_mode = st.radio(
            "选择导出方式",
            options=["独立文件（ZIP压缩）", "合并为单个文档"],
            key="export_mode_radio",
            horizontal=True,
            help="ZIP：下载所有文件 | 合并：一个文档包含所有内容"
        )

        st.markdown("---")

        # 统计信息
        st.markdown("### 替换统计")

        col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4, gap="medium")

        with col_stat1:
            st.metric("📄 生成文件数", len(st.session_state.replaced_files))

        with col_stat2:
            st.metric("📋 替换规则数", len(st.session_state.replace_rules))

        with col_stat3:
            success_count = len([f for f in st.session_state.replaced_files
                                 if f.data and len(f.data.getvalue()) > 0])
            st.metric("✅ 成功文件数", success_count)

        with col_stat4:
            st.metric("⏱️ 生成时间", datetime.now().strftime("%H:%M:%S"))

        # 关键字统计
        st.markdown("---")
        st.markdown("### 关键字替换统计")

        keyword_stats = get_keyword_statistics(st.session_state.replace_rules,
                                               st.session_state.replaced_files)
        if keyword_stats and any(v > 0 for v in keyword_stats.values()):
            stat_data = [
                {"关键字": k, "总替换次数": v}
                for k, v in keyword_stats.items() if v > 0
            ]
            if stat_data:
                stat_df = pd.DataFrame(stat_data)
                st.dataframe(stat_df, use_container_width=True, hide_index=True)

                # 导出统计
                if st.button("📊 导出统计到CSV", key="export_stats", use_container_width=True):
                    csv_data = export_statistics_to_csv(st.session_state.replaced_files)
                    st.download_button(
                        label="📥 下载统计CSV",
                        data=csv_data,
                        file_name="统计数据.csv",
                        mime="text/csv",
                        key="download_stats",
                        use_container_width=True
                    )

        st.markdown("---")
        st.markdown("### 批量导出")

        # 导出按钮
        if export_mode == "独立文件（ZIP压缩）":
            try:
                valid_files = [f for f in st.session_state.replaced_files
                               if f.data and len(f.data.getvalue()) > 0]

                if not valid_files:
                    st.error("❌ 没有有效的文件", icon="❌")
                else:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for file in valid_files:
                            zipf.writestr(file.filename, file.data.getvalue())

                    zip_buffer.seek(0)
                    zip_filename = f"批量替换_{len(valid_files)}个文件.zip"

                    st.download_button(
                        label=f"📦 下载ZIP - {len(valid_files)} 个文件",
                        data=zip_buffer,
                        file_name=zip_filename,
                        mime="application/zip",
                        key="download_all_zip",
                        use_container_width=True,
                        type="primary"
                    )
            except Exception as e:
                st.error(f"❌ 创建ZIP失败：{str(e)}", icon="❌")
        else:
            valid_files = [f for f in st.session_state.replaced_files
                           if f.data and len(f.data.getvalue()) > 0]

            if not valid_files:
                st.error("❌ 没有有效的文件", icon="❌")
            else:
                try:
                    merged_data = merge_word_documents(valid_files)
                    merged_filename = "合并结果.docx"

                    st.download_button(
                        label=f"📋 下载合并文档 - {len(valid_files)} 个文件",
                        data=merged_data,
                        file_name=merged_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_merged",
                        use_container_width=True,
                        type="primary"
                    )
                except Exception as e:
                    st.error(f"❌ 合并失败：{str(e)}", icon="❌")

        st.markdown("---")
        st.markdown("### 文件列表")

        # 分页显示
        total_pages = (len(st.session_state.replaced_files) + PAGE_SIZE - 1) // PAGE_SIZE

        col_page_col = st.columns([4])[0]
        with col_page_col:
            current_page = st.number_input(
                "页码",
                min_value=1,
                max_value=total_pages,
                value=1,
                key="current_page"
            )

        start_idx = (current_page - 1) * PAGE_SIZE
        end_idx = min(start_idx + PAGE_SIZE, len(st.session_state.replaced_files))
        current_files = st.session_state.replaced_files[start_idx:end_idx]

        st.caption(f"第 {current_page}/{total_pages} 页（共 {len(st.session_state.replaced_files)} 个文件）")

        # 文件列表
        for idx, file in enumerate(current_files, start=start_idx + 1):
            is_valid = file.data and len(file.data.getvalue()) > 0
            status_icon = "✅" if is_valid else "❌"

            col_info, col_log, col_download = st.columns([3, 2, 1], gap="medium")

            with col_info:
                st.markdown(f"**{status_icon} {idx}. {file.filename}**")
                st.caption(f"Excel行号：#{file.row_idx + 1}")

            with col_log:
                with st.expander("📋 查看日志", expanded=False):
                    st.code(file.log, language="text")

            with col_download:
                st.download_button(
                    label="⬇️ 下载",
                    data=file.data,
                    file_name=file.filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{idx}",
                    disabled=not is_valid,
                    use_container_width=True
                )

            st.divider()

    else:
        st.info("📁 暂无生成的文件，请先执行替换", icon="ℹ️")

# ==================== 标签页4：工具设置 ====================
with tab4:
    st.subheader("⚙️ 工具设置与帮助")

    # 操作历史
    st.markdown("### 📜 操作历史记录")
    history = history_manager.load_history()

    if history:
        col_hist1, col_hist2 = st.columns([4, 1], gap="medium")
        with col_hist1:
            st.markdown(f"**最近 {len(history)} 次操作**")
        with col_hist2:
            if st.button("🗑️ 清除历史", key="clear_history", use_container_width=True):
                history_manager.clear_history()
                st.rerun()

        # 历史记录表格
        history_data = []
        for i, record in enumerate(history[:20], 1):
            status_emoji = "✅" if record["status"] == "success" else "❌"
            history_data.append({
                "序号": i,
                "时间": record["timestamp"],
                "Word文件": record["word_file"][:20] + "..." if len(record["word_file"]) > 20 else record["word_file"],
                "Excel文件": record["excel_file"][:20] + "..." if len(record["excel_file"]) > 20 else record[
                    "excel_file"],
                "规则数": record["rules_count"],
                "生成文件": record["files_generated"],
                "状态": status_emoji
            })

        history_df = pd.DataFrame(history_data)
        st.dataframe(history_df, use_container_width=True, hide_index=True)
    else:
        st.info("📁 暂无操作历史", icon="ℹ️")

    st.markdown("---")

    # 关于工具
    st.markdown("### ℹ️ 关于此工具")

    col_about1, col_about2 = st.columns([2, 2], gap="medium")

    with col_about1:
        st.markdown(f"""
        **Word+Excel批量替换工具**

        版本：{VERSION}

        **功能特性：**
        ✅ 批量替换
        ✅ 格式保留
        ✅ 文档合并
        ✅ 规则管理
        ✅ 统计分析
        ✅ 历史记录
        """)

    with col_about2:
        st.markdown("""
        **快速指南：**

        1. 上传Word和Excel文件
        2. 预览内容，复制关键字
        3. 添加替换规则
        4. 执行批量替换
        5. 下载结果文件

        **支持的格式：**
        • Word：.docx
        • Excel：.xlsx/.xls
        • 括号：【】（）()〔〕
        """)

    st.markdown("---")

    # 常见问题
    st.markdown("### ❓ 常见问题")

    with st.expander("1️⃣ 支持.doc格式吗？"):
        st.markdown("""
        不支持.doc格式，需要转换为.docx。

        **转换方法：**
        1. 用Word打开.doc文件
        2. 另存为 → Word文档(.docx)
        3. 重新上传
        """)

    with st.expander("2️⃣ 怎样保持原文档格式？"):
        st.markdown("""
        本工具自动保留：
        • 段落格式
        • 字体样式
        • 表格结构
        • 颜色等

        只替换文本内容，不影响其他格式。
        """)

    with st.expander("3️⃣ 如何合并多个文档？"):
        st.markdown("""
        1. 设置替换规则并执行替换
        2. 在"下载结果"选择"合并为单个文档"
        3. 点击"下载合并文档"

        会自动在每个文档间插入分页符。
        """)

    with st.expander("4️⃣ 能否处理大数据？"):
        st.markdown("""
        **限制说明：**
        • Word文件：最大50MB
        • Excel文件：最大50MB
        • 行数：建议<1000行

        大数据建议分批处理。
        """)

    with st.expander("5️⃣ 规则如何保存？"):
        st.markdown("""
        **两种保存方式：**

        1. **导出JSON**
           - 下载规则文件
           - 可在其他电脑导入

        2. **保存缓存**
           - 快速保存
           - 本地快速加载
        """)

    st.markdown("---")

    # 更新日志
    st.markdown("### 📝 更新日志")

    st.markdown("""
    **v1.4.1** ⭐ 最新
    - 完全重构布局，改善用户体验
    - 整合功能到标签页
    - 优化侧栏信息展示
    - 改进响应式设计

    **v1.4.0**
    - 新增快速加载缓存规则
    - 新增操作历史记录
    - 新增关键字替换统计
    - 新增导出统计数据到CSV
    - 新增规则撤销功能

    **v1.3.2**
    - 修复多个bug
    - 优化合并文档格式保留

    **v1.3.0**
    - 添加合并文档功能

    **v1.0.0**
    - 初始版本
    """)

    st.markdown("---")

    st.markdown("""
    <div style='text-align: center; padding: 20px; color: #666;'>
    <p>© 2024 Word+Excel批量替换工具</p>
    <p>让批量替换变得简单高效</p>
    </div>
    """, unsafe_allow_html=True)