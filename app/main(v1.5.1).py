import os
import sys
import tempfile
from tempfile import NamedTemporaryFile
import warnings
import shutil
import json
import io
import zipfile
import re
import unicodedata
import copy
from datetime import datetime
import hashlib

# 数据处理库
import streamlit as st
import pandas as pd

# Word处理库
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 数据结构和类型提示
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple, Set
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP

# ==================== 配置和常量 ====================

# 项目版本信息
VERSION = "v1.5.1"

# 页面配置常量
PAGE_SIZE = 10  # 每页显示的文件数
WIDGET_HEIGHT = 250  # 组件高度
PREVIEW_ROWS = 20  # 数据预览行数
MAX_FILENAME_LENGTH = 200  # 最大文件名长度
MAX_WORD_FILE_SIZE = 50 * 1024 * 1024  # 最大Word文件大小：50MB
MAX_EXCEL_FILE_SIZE = 50 * 1024 * 1024  # 最大Excel文件大小：50MB
CACHE_DIR = ".replace_cache"  # 缓存目录
HISTORY_FILE = ".replace_history.json"  # 历史记录文件
MAX_HISTORY_ITEMS = 30  # 最多保存历史记录数

# 过滤警告消息
warnings.filterwarnings("ignore", category=UserWarning)

# 环境变量配置
os.environ["STREAMLIT_VERSION"] = "1.51.0"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# ==================== Streamlit页面配置 ====================
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 全局样式 ====================
st.markdown("""
<style>
    /* ===== 全局间距优化 ===== */
    .main {
        padding: 0.5rem 1rem !important;
    }

    [data-testid="stMainBlockContainer"] {
        padding-top: 0.5rem !important;
        padding-bottom: 0.5rem !important;
    }

    /* 块容器紧凑 */
    .stContainer {
        padding: 0.75rem !important;
        margin-bottom: 0.5rem !important;
        border-radius: 6px;
        background-color: #ffffff;
    }

    /* 删除元素间多余间距 */
    .element-container {
        margin-bottom: 0.3rem !important;
    }

    .stColumn {
        gap: 0.5rem !important;
    }

    /* ===== 按钮样式 ===== */
    .stButton > button {
        border-radius: 5px;
        font-weight: 500;
        padding: 0.4rem 0.8rem !important;
        font-size: 13px !important;
        margin-bottom: 0.2rem !important;
    }

    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 2px 8px rgba(0,0,0,0.12);
    }

    /* ===== 输入框样式 ===== */
    .stTextInput, .stTextArea, .stSelectbox, .stNumberInput {
        margin-bottom: 0.3rem !important;
    }

    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select,
    .stNumberInput > div > div > input {
        border-radius: 5px;
        border: 1px solid #e0e0e0;
        font-size: 13px;
        padding: 0.5rem !important;
    }

    /* ===== 标题和文字样式 ===== */
    h1 {
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #1f77b4;
        margin-bottom: 0.5rem;
        line-height: 1.2;
    }

    h2 {
        margin-top: 0.5rem;
        margin-bottom: 0.3rem;
        color: #1f77b4;
        font-size: 1.2rem;
    }

    h3 {
        margin-top: 0.3rem;
        margin-bottom: 0.2rem;
        color: #333;
        font-size: 1.05rem;
    }

    .stSubheader {
        margin-bottom: 0.5rem !important;
        padding-bottom: 0.3rem;
        border-bottom: 1.5px solid #e0e0e0;
        font-size: 1.1rem !important;
    }

    /* ===== 展开器样式 ===== */
    .streamlit-expander {
        margin-bottom: 0.3rem !important;
        border-radius: 5px;
    }

    /* ===== 标签页样式 ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1px;
        margin-bottom: 0.5rem;
    }

    .stTabs [data-baseweb="tab"] {
        height: 40px;
        padding-top: 8px;
        border-radius: 5px 5px 0 0;
        font-size: 13px;
    }

    /* ===== 数据框样式 ===== */
    div[data-testid="stDataFrame"] {
        border-radius: 5px;
        border: 1px solid #e0e0e0;
        font-size: 12px;
    }

    /* ===== 指标卡样式 ===== */
    .metric-container {
        background-color: #f8f9fa;
        padding: 0.5rem !important;
        border-radius: 5px;
        border-left: 3px solid #1f77b4;
        margin-bottom: 0.3rem;
    }

    /* ===== 信息框样式 ===== */
    .stats-box, .success-box, .warning-box, .error-box {
        padding: 0.6rem !important;
        margin: 0.3rem 0 !important;
        border-radius: 5px;
        border-left-width: 3px;
        font-size: 13px;
    }

    .stats-box {
        background-color: #f0f9ff;
        border-left-color: #0ea5e9;
    }

    .success-box {
        background-color: #f0fdf4;
        border-left-color: #22c55e;
    }

    .warning-box {
        background-color: #fffbeb;
        border-left-color: #f59e0b;
    }

    .error-box {
        background-color: #fef2f2;
        border-left-color: #ef4444;
    }

    /* ===== 分隔线 ===== */
    hr {
        margin: 0.5rem 0 !important;
        border: none;
        border-top: 1px solid #e0e0e0;
    }

    /* ===== 无线电和复选框样式 ===== */
    .stRadio, .stCheckbox {
        margin-bottom: 0.2rem !important;
    }

    .stRadio > label, .stCheckbox > label {
        margin-bottom: 0.2rem !important;
        font-size: 13px;
    }

    /* ===== 文件上传器样式 ===== */
    .stFileUploader {
        margin-bottom: 0.3rem !important;
    }

    /* ===== 进度条样式 ===== */
    .stProgress {
        margin-bottom: 0.3rem !important;
    }

    /* ===== 表格样式 ===== */
    table {
        font-size: 12px !important;
    }

    td, th {
        padding: 0.4rem !important;
    }

    /* ===== Word预览区域样式（修复高度问题） ===== */
    .word-preview-area {
        height: 280px;
        overflow-y: auto;
        padding: 12px;
        border: 1px solid #e0e0e0;
        border-radius: 6px;
        font-size: 13px;
        line-height: 1.6;
        background-color: #f9f9f9;
        word-break: break-word;
        white-space: pre-wrap;
    }

    /* ===== 规则列表容器样式 ===== */
    .rule-list-container {
        border: 1px solid #e0e0e0;
        border-radius: 6px;
        padding: 8px;
        background-color: #f9f9f9;
        height: 280px;
        overflow-y: auto;
    }

    /* ===== 响应式设计 ===== */
    @media (max-width: 768px) {
        .main {
            padding: 0.3rem 0.5rem;
        }
        .stContainer {
            padding: 0.5rem;
        }
    }

    /* ===== 帮助提示样式 ===== */
    .help-text {
        color: #0ea5e9;
        cursor: help;
        font-weight: bold;
    }

    .help-tooltip {
        background-color: #f0f9ff;
        border-left: 3px solid #0ea5e9;
        padding: 8px;
        border-radius: 4px;
        margin-top: 4px;
        font-size: 12px;
        color: #0c4a6e;
    }
</style>
""", unsafe_allow_html=True)


# ==================== 数据结构定义 ====================

@dataclass
class ReplacedFile:
    """
    存储替换后的文件数据结构

    Attributes:
        filename (str): 生成的文件名
        data (io.BytesIO): 文件二进制数据
        row_idx (int): 对应Excel行号
        log (str): 替换日志信息
        replace_count (int): 替换次数
    """
    filename: str
    data: io.BytesIO
    row_idx: int
    log: str
    replace_count: int = 0


@dataclass
class HistoryRecord:
    """
    历史记录数据结构

    Attributes:
        timestamp (str): 操作时间
        word_file (str): Word文件名
        excel_file (str): Excel文件名
        rules_count (int): 规则数量
        files_generated (int): 生成文件数
        status (str): 操作状态(success/failed)
    """
    timestamp: str
    word_file: str
    excel_file: str
    rules_count: int
    files_generated: int
    status: str


# ==================== 缓存管理器 ====================

class CacheManager:
    """
    管理替换规则的缓存
    负责保存、加载和列出缓存的规则
    """

    def __init__(self):
        """初始化缓存管理器，创建缓存目录"""
        self.cache_dir = CACHE_DIR
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)

    def save_rules(self, rules: List[Tuple[str, str]], filename: str):
        """
        保存规则到JSON缓存文件

        Args:
            rules: 规则列表，每个元素为(关键字, Excel列名)的元组
            filename: 缓存文件名
        """
        try:
            rules_data = [{"keyword": old, "excel_column": col} for old, col in rules]
            cache_file = os.path.join(self.cache_dir, f"{filename}.json")
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(rules_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"⚠️ 保存缓存失败", icon="⚠️")

    def load_rules(self, filename: str) -> List[Tuple[str, str]]:
        """
        从缓存文件加载规则

        Args:
            filename: 缓存文件名

        Returns:
            规则列表
        """
        try:
            cache_file = os.path.join(self.cache_dir, f"{filename}.json")
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    rules_data = json.load(f)
                    return [(r["keyword"], r["excel_column"]) for r in rules_data]
        except:
            pass
        return []

    def get_cached_rules_list(self) -> List[str]:
        """
        获取所有缓存的规则文件列表

        Returns:
            缓存文件名列表（最近10个）
        """
        try:
            if os.path.exists(self.cache_dir):
                files = [f.replace('.json', '') for f in os.listdir(self.cache_dir) if f.endswith('.json')]
                return sorted(files, reverse=True)[:10]
        except:
            pass
        return []


# ==================== 历史记录管理器 ====================

class HistoryManager:
    """
    管理操作历史记录
    记录每次替换操作的信息
    """

    def __init__(self):
        """初始化历史记录管理器"""
        self.history_file = HISTORY_FILE

    def add_record(self, record: HistoryRecord):
        """
        添加操作记录到历史

        Args:
            record: HistoryRecord对象
        """
        try:
            history = self.load_history()
            history.insert(0, {
                "timestamp": record.timestamp,
                "word_file": record.word_file,
                "excel_file": record.excel_file,
                "rules_count": record.rules_count,
                "files_generated": record.files_generated,
                "status": record.status
            })
            history = history[:MAX_HISTORY_ITEMS]
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except:
            pass

    def load_history(self) -> List[Dict]:
        """
        加载所有历史记录

        Returns:
            历史记录列表
        """
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return []

    def clear_history(self):
        """清除所有历史记录"""
        try:
            if os.path.exists(self.history_file):
                os.remove(self.history_file)
                st.success("✅ 历史已清除", icon="✅")
        except:
            pass


# ==================== 会话状态初始化 ====================

def init_session_state():
    """
    初始化Streamlit会话状态
    确保所有必需的状态变量都存在
    """
    required_states = {
        "replace_rules": [],  # 替换规则列表
        "replaced_files": [],  # 替换后的文件列表
        "replace_log": [],  # 替换日志
        "is_replacing": False,  # 是否正在替换中
        "replace_params": {},  # 替换参数缓存
        "replace_scope": "替换完整关键词",  # 替换范围模式
        "export_mode_radio": "独立文件（ZIP压缩）",  # 导出模式
        "undo_stack": [],  # 规则撤销栈
        "rule_filter": "",  # 规则筛选文本
        "show_advanced": False,  # 是否显示高级选项
    }

    for key, default in required_states.items():
        if key not in st.session_state:
            st.session_state[key] = default


# 调用初始化函数
init_session_state()


# ==================== 核心工具函数 ====================

def clean_text(text: str) -> str:
    """
    清理文本：去除首尾空白、隐藏字符、特殊空格，统一格式

    Args:
        text: 输入文本

    Returns:
        清理后的文本
    """
    if not isinstance(text, str):
        return ""
    text = text.strip()
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'[\u00A0\u2002-\u200B]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text


def clean_filename(filename: str) -> str:
    """
    清理文件名中的非法字符

    Args:
        filename: 原始文件名

    Returns:
        清理后的合法文件名
    """
    return re.sub(r'[\\/:*?"<>|]', "_", str(filename))


def generate_safe_filename(
        excel_row: pd.Series,
        file_name_col: str,
        file_prefix: str = "",
        file_suffix: str = "",
        row_idx: int = 0,
        max_length: int = MAX_FILENAME_LENGTH
) -> str:
    """
    安全生成文件名，处理超长名称和特殊字符

    Args:
        excel_row: Excel行数据
        file_name_col: 文件名所在的列名
        file_prefix: 文件名前缀
        file_suffix: 文件名后缀
        row_idx: 行号（备用）
        max_length: 最大长度

    Returns:
        清理后的文件名
    """
    try:
        # 获取基础名称
        if file_name_col and file_name_col in excel_row.index:
            base_name = clean_text(str(excel_row[file_name_col]))
        else:
            base_name = f"文件_{row_idx + 1}"

        # 确保基础名称不为空
        if not base_name or base_name.isspace():
            base_name = f"文件_{row_idx + 1}"

        # 构建完整文件名
        if file_prefix and file_suffix:
            filename = f"{file_prefix}{base_name}{file_suffix}.docx"
        elif file_prefix:
            filename = f"{file_prefix}{base_name}.docx"
        elif file_suffix:
            filename = f"{base_name}{file_suffix}.docx"
        else:
            filename = f"{base_name}.docx"

        # 清理非法字符
        filename = clean_filename(filename)

        # 限制长度（Windows限制255字节）
        filename_bytes = filename.encode('utf-8')
        if len(filename_bytes) > max_length:
            truncated_base = base_name
            while len(f"{file_prefix}{truncated_base}{file_suffix}.docx".encode('utf-8')) > max_length:
                truncated_base = truncated_base[:-1]

            if file_prefix and file_suffix:
                filename = f"{file_prefix}{truncated_base}{file_suffix}.docx"
            elif file_prefix:
                filename = f"{file_prefix}{truncated_base}.docx"
            elif file_suffix:
                filename = f"{truncated_base}{file_suffix}.docx"
            else:
                filename = f"{truncated_base}.docx"

            filename = clean_filename(filename)

        return filename

    except:
        return f"文件_{row_idx + 1}.docx"


def precompute_replace_patterns(
        replace_rules: List[Tuple[str, str]],
        excel_row: pd.Series
) -> List[Tuple[str, str, str, str]]:
    """
    预计算所有需要替换的模式，减少重复计算

    Args:
        replace_rules: 替换规则列表
        excel_row: 当前处理的Excel行数据

    Returns:
        替换模式列表：[(原始关键词, 列名, 清理后关键词, 替换值), ...]
    """
    replace_patterns = []

    for old_text, col_name in replace_rules:
        # 获取Excel中对应列的替换值
        if col_name in excel_row.index:
            replacement = str(excel_row[col_name]).strip()
        else:
            replacement = ""

        # 清理用户输入的关键词
        cleaned_text = clean_text(old_text)

        # 跳过空关键词
        if not cleaned_text:
            continue

        # 根据替换范围选项生成替换值
        if st.session_state.replace_scope == "仅替换括号内内容":
            # 检查是否是带括号的格式，只替换括号内的内容
            if cleaned_text.startswith("【") and cleaned_text.endswith("】"):
                new_format = f"【{replacement}】"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("（") and cleaned_text.endswith("）"):
                new_format = f"（{replacement}）"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("(") and cleaned_text.endswith(")"):
                new_format = f"({replacement})"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("〔") and cleaned_text.endswith("〕"):
                new_format = f"〔{replacement}〕"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            else:
                replace_patterns.append((old_text, col_name, cleaned_text, replacement))
        else:
            # 替换完整关键词
            replace_patterns.append((old_text, col_name, cleaned_text, replacement))

    return replace_patterns


def process_paragraph(
        paragraph,
        replace_patterns: List[Tuple[str, str, str, str]],
        cleaned_para: str = None
) -> Dict:
    """
    处理单个段落的关键字替换

    Args:
        paragraph: 要处理的段落对象
        replace_patterns: 替换模式列表
        cleaned_para: 预清理的段落文本（可选）

    Returns:
        替换计数字典：{(原始关键词, 列名): 替换次数, ...}
    """
    para_text = paragraph.text
    if cleaned_para is None:
        cleaned_para = clean_text(para_text)
    replace_count = defaultdict(int)

    # 如果段落为空，直接返回
    if not para_text or not replace_patterns:
        return replace_count

    has_keyword = False

    # 检查段落是否包含任何需要替换的关键字（优化性能）
    for old_text, col_name, format_keyword, replacement in replace_patterns:
        if format_keyword and format_keyword in cleaned_para:
            has_keyword = True
            break

    if has_keyword:
        # 创建新文本并替换所有关键字
        new_text = para_text
        for old_text, col_name, format_keyword, replacement in replace_patterns:
            if format_keyword and format_keyword in cleaned_para:
                count = new_text.count(format_keyword)
                if count > 0:
                    new_text = new_text.replace(format_keyword, replacement)
                    replace_count[(old_text, col_name)] += count

        # 更新段落文本（保留第一个Run的格式）
        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
            # 清空其他Run
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ''

    return replace_count


def replace_word_with_format(
        word_file: st.runtime.uploaded_file_manager.UploadedFile,
        excel_row: pd.Series,
        replace_rules: List[Tuple[str, str]]
) -> Tuple[io.BytesIO, str, int]:
    """
    替换Word文件中的关键字，保留格式

    Args:
        word_file: 上传的Word文件
        excel_row: 当前Excel行数据
        replace_rules: 替换规则列表

    Returns:
        (替换后的文件数据, 替换日志, 替换总次数)
    """
    replace_count = defaultdict(int)
    total_replace = 0

    try:
        # 检查文件大小
        file_size = len(word_file.getvalue())
        if file_size > MAX_WORD_FILE_SIZE:
            raise ValueError(f"文件过大")

        # 从内存加载Word文档
        doc = Document(io.BytesIO(word_file.getvalue()))

        # 预计算替换模式，减少重复计算
        replace_patterns = precompute_replace_patterns(replace_rules, excel_row)

        # 如果没有替换模式，直接返回原文档
        if not replace_patterns:
            output_file = io.BytesIO()
            doc.save(output_file)
            output_file.seek(0)
            return output_file, "⚠ 未找到匹配规则", 0

        # 1. 处理段落中的关键词
        for paragraph in doc.paragraphs:
            para_count = process_paragraph(paragraph, replace_patterns)
            for key, count in para_count.items():
                replace_count[key] += count
                total_replace += count

        # 2. 处理表格内的关键词（支持表格内文字替换）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_count = process_paragraph(paragraph, replace_patterns)
                        for key, count in para_count.items():
                            replace_count[key] += count
                            total_replace += count

        # 保存修改后的文档
        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)

        # 生成替换日志
        if replace_count:
            log_lines = [f"✓ {old}" for old, _ in replace_count.keys()]
            replace_log = ", ".join(log_lines[:3])
            if len(replace_count) > 3:
                replace_log += f" 等{len(replace_count) - 3}个"
        else:
            replace_log = "⚠ 无替换"

        return output_file, replace_log, total_replace

    except Exception as e:
        return io.BytesIO(), f"❌ 失败", 0


def merge_word_documents(
        replaced_files: List[ReplacedFile]
) -> io.BytesIO:
    """
    合并多个Word文档为一个（保留所有格式和结构）

    Args:
        replaced_files: 替换后的文件列表

    Returns:
        合并后的Word文档数据
    """
    if not replaced_files:
        raise ValueError("没有文件")

    try:
        # 加载第一个文档作为主文档
        main_doc = Document(io.BytesIO(replaced_files[0].data.getvalue()))
        main_body = main_doc._body._element

        # 逐个添加其他文档
        for idx in range(1, len(replaced_files)):
            try:
                file = replaced_files[idx]

                # 检查文件数据有效性
                if not file.data or len(file.data.getvalue()) == 0:
                    continue

                # 加载子文档
                sub_doc = Document(io.BytesIO(file.data.getvalue()))
                sub_body = sub_doc._body._element

                # 添加分页符
                page_break_para = OxmlElement('w:p')
                page_break_pPr = OxmlElement('w:pPr')

                page_break_element = OxmlElement('w:pageBreakBefore')
                page_break_element.set(qn('w:val'), '1')

                page_break_pPr.append(page_break_element)
                page_break_para.append(page_break_pPr)
                main_body.append(page_break_para)

                # 深拷贝所有元素以保留格式
                for element in sub_body:
                    main_body.append(copy.deepcopy(element))

            except:
                continue

        # 保存合并后的文档
        output = io.BytesIO()
        main_doc.save(output)
        output.seek(0)
        return output

    except Exception as e:
        raise


def get_replace_params(
        word_file: Optional[st.runtime.uploaded_file_manager.UploadedFile],
        excel_df: Optional[pd.DataFrame],
        start_row: int,
        end_row: int,
        file_name_col: str,
        file_prefix: str,
        file_suffix: str
) -> Dict:
    """
    获取替换参数，用于判断是否需要重新替换

    Args:
        word_file: 上传的Word文件
        excel_df: Excel数据框
        start_row: 起始行
        end_row: 结束行
        file_name_col: 文件名列
        file_prefix: 文件前缀
        file_suffix: 文件后缀

    Returns:
        替换参数字典
    """
    return {
        "word_filename": word_file.name if word_file else "",
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "rule_count": len(st.session_state.replace_rules),
        "rule_hash": hash(tuple(st.session_state.replace_rules))
    }


def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """
    清理Excel数据类型，避免混合类型导致的问题

    Args:
        df: 输入的数据框

    Returns:
        清理后的数据框
    """
    df_clean = df.copy()

    for col in df_clean.columns:
        try:
            col_name = str(col)
            if col_name != col:
                df_clean = df_clean.rename(columns={col: col_name})
                col = col_name

            df_clean[col] = df_clean[col].fillna("")
            df_clean[col] = df_clean[col].astype(str).str.strip()

        except:
            try:
                df_clean[col] = df_clean[col].astype(str).str.strip()
            except:
                pass

    return df_clean


def get_file_hash(file_data: bytes) -> str:
    """
    获取文件哈希值（用于验证文件完整性）

    Args:
        file_data: 文件二进制数据

    Returns:
        文件哈希值的前6位
    """
    return hashlib.md5(file_data).hexdigest()[:6]


def export_statistics_to_csv(replaced_files: List[ReplacedFile]) -> str:
    """
    导出替换统计数据到CSV格式

    Args:
        replaced_files: 替换后的文件列表

    Returns:
        CSV格式的字符串
    """
    try:
        data = []
        for idx, file in enumerate(replaced_files, 1):
            data.append({
                "序号": idx,
                "文件名": file.filename,
                "行号": file.row_idx + 1,
                "替换次数": file.replace_count,
                "状态": "✅" if file.data and len(file.data.getvalue()) > 0 else "❌"
            })

        df = pd.DataFrame(data)
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
        return csv_buffer.getvalue()
    except:
        return ""


def get_keyword_statistics(replace_rules: List[Tuple[str, str]],
                           replaced_files: List[ReplacedFile]) -> Dict:
    """
    获取关键字替换统计

    Args:
        replace_rules: 替换规则列表
        replaced_files: 替换后的文件列表

    Returns:
        关键字统计字典
    """
    stats = {}
    for keyword, _ in replace_rules:
        stats[keyword] = 0

    for file in replaced_files:
        for keyword, _ in replace_rules:
            if f"✓ {keyword}" in file.log:
                pattern = f"✓ {re.escape(keyword)}.*?\((\d+)次\)"
                matches = re.findall(pattern, file.log)
                if matches:
                    stats[keyword] += int(matches[0])

    return stats


# ==================== 创建管理器实例 ====================
cache_manager = CacheManager()
history_manager = HistoryManager()

# ==================== 侧栏 ====================
with st.sidebar:
    st.title("📚 快速导航")

    # 统计信息
    if st.session_state.replaced_files:
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 文件数", len(st.session_state.replaced_files), delta=None)
        with col2:
            st.metric("📋 规则数", len(st.session_state.replace_rules), delta=None)

    st.markdown("---")

    # 快速功能
    st.subheader("⚡ 快速功能")

    # 快速加载缓存规则
    cached = cache_manager.get_cached_rules_list()
    if cached:
        selected = st.selectbox("📂 加载规则", ["选择..."] + cached, key="sidebar_cache")
        if selected and selected != "选择...":
            if st.button("✅ 加载", key="sidebar_load", use_container_width=True):
                loaded = cache_manager.load_rules(selected)
                if loaded:
                    st.session_state.replace_rules = loaded
                    st.success(f"✅ 加载{len(loaded)}条", icon="✅")
                    st.rerun()

    # 历史记录显示
    history = history_manager.load_history()
    if history:
        st.subheader("📜 最近操作")
        for h in history[:3]:
            status = "✅" if h["status"] == "success" else "❌"
            st.caption(f"{status} {h['timestamp']}\n{h['word_file'][:15]}...")

    st.markdown("---")

    # 工具操作
    st.subheader("🔧 工具")

    if st.button("🗑️ 清空所有", key="sidebar_clear", use_container_width=True):
        st.session_state.replace_rules = []
        st.session_state.replaced_files = []
        st.success("✅ 已清空", icon="✅")
        st.rerun()

    if history:
        if st.button("📜 清除历史", key="sidebar_clear_hist", use_container_width=True):
            history_manager.clear_history()
            st.rerun()

# ==================== 主页面 - 标题 ====================
col_title1, col_title2 = st.columns([8, 2])
with col_title1:
    st.title("📋 Word+Excel批量替换工具")
with col_title2:
    st.markdown(
        f"<div style='text-align: right; padding-top: 5px;'><small style='color: #999;'>v{VERSION}</small></div>",
        unsafe_allow_html=True)

# 进度显示（如果有替换任务）
if st.session_state.replaced_files and st.session_state.replace_params:
    progress_col, status_col = st.columns([3, 1])
    with progress_col:
        success_count = len([f for f in st.session_state.replaced_files
                             if f.data and len(f.data.getvalue()) > 0])
        total_count = len(st.session_state.replaced_files)
        st.progress(success_count / total_count if total_count > 0 else 0)
    with status_col:
        st.metric("成功率", f"{int(success_count / total_count * 100) if total_count > 0 else 0}%")

st.markdown("---")

# ==================== 主工作区 ====================
col_main_left, col_main_right = st.columns([2, 1], gap="medium")

# ==================== 左侧：文件上传和预览 ====================
with col_main_left:
    st.subheader("📤 文件上传")

    # 上传区域
    col_upload1, col_upload2 = st.columns(2, gap="small")

    with col_upload1:
        st.markdown("**Word模板** ℹ️")
        with st.expander("", expanded=False):
            st.markdown("""
            **步骤：**
            1. 上传包含要替换内容的Word文件
            2. 格式必须是.docx（不支持.doc）
            3. 最大文件大小50MB

            **示例：**
            在Word中有「【姓名】」需要替换为Excel中的真实姓名
            """)

        word_file = st.file_uploader(
            "选择文件",
            type=["docx"],
            key="word",
            label_visibility="collapsed",
            help="仅支持.docx格式，.doc需先转换"
        )
        if word_file:
            file_size_mb = len(word_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_WORD_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ 文件过大", icon="❌")
                word_file = None
            else:
                st.caption(f"✅ {file_size_mb:.1f}MB")

    with col_upload2:
        st.markdown("**Excel数据** ℹ️")
        with st.expander("", expanded=False):
            st.markdown("""
            **步骤：**
            1. 上传包含替换数据的Excel文件
            2. 支持.xlsx和.xls格式
            3. 第一行通常为列标题

            **示例：**
            姓名列包含要替换的真实姓名
            """)

        excel_file = st.file_uploader(
            "选择文件",
            type=["xlsx", "xls"],
            key="excel",
            label_visibility="collapsed",
            help="支持.xlsx/.xls格式"
        )
        if excel_file:
            file_size_mb = len(excel_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_EXCEL_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ 文件过大", icon="❌")
                excel_file = None
            else:
                st.caption(f"✅ {file_size_mb:.1f}MB")

    st.markdown("---")

    # 文件预览 - 可折叠
    with st.expander("👀 文件预览 - 点击查看/复制内容", expanded=False):
        col_prev1, col_prev2 = st.columns(2, gap="small")

        excel_df = None
        excel_cols = []

        with col_prev1:
            st.markdown("**Word文档内容**")
            if word_file:
                try:
                    # 读取Word文件并生成HTML预览
                    doc = Document(io.BytesIO(word_file.getvalue()))

                    # 构建HTML
                    html_content = ""

                    # 添加段落
                    for para in doc.paragraphs[:15]:  # 限制显示段落数
                        if para.text.strip():
                            text = para.text.replace("<", "&lt;").replace(">", "&gt;")
                            html_content += f"<p style='margin: 4px 0; word-break: break-all;'>{text}</p>"

                    # 添加表格预览
                    for table_idx, table in enumerate(doc.tables[:2]):
                        html_content += f"<p style='margin-top: 8px; font-weight: bold; color: #1f77b4;'>📊 表格{table_idx + 1}：</p>"
                        html_content += "<table style='border-collapse: collapse; width: 100%; font-size: 12px;'>"

                        for row_idx, row in enumerate(table.rows[:10]):
                            html_content += "<tr>"
                            for cell in row.cells:
                                cell_text = cell.text.replace("<", "&lt;").replace(">", "&gt;")[:30]
                                html_content += f"<td style='border: 1px solid #ccc; padding: 4px;'>{cell_text}</td>"
                            html_content += "</tr>"

                        html_content += "</table>"

                    # 使用自定义HTML组件显示
                    st.components.v1.html(f"""
                    <div style='height: 280px; overflow-y: auto; padding: 12px; border: 1px solid #e0e0e0; 
                                border-radius: 6px; font-size: 13px; line-height: 1.6; background-color: #f9f9f9;
                                font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif; word-wrap: break-word;'>
                        {html_content}
                    </div>
                    """, height=300)

                    st.caption(f"📄 {len(doc.paragraphs)}段落，{len(doc.tables)}表格")
                    st.info("💡 可以在上方选中内容按Ctrl+C复制，粘贴到下方关键字输入框中", icon="ℹ️")

                except Exception as e:
                    st.error(f"❌ 预览失败：{str(e)[:50]}", icon="❌")
            else:
                st.info("请上传Word文件", icon="ℹ️")

        with col_prev2:
            st.markdown("**Excel数据预览**")
            if excel_file:
                try:
                    with NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_excel:
                        temp_excel.write(excel_file.getvalue())
                        excel_path = temp_excel.name

                    try:
                        with pd.ExcelFile(excel_path, engine="openpyxl") as excel_wb:
                            sheet_names = excel_wb.sheet_names
                            selected_sheet = sheet_names[0]

                            excel_df = pd.read_excel(
                                excel_wb,
                                sheet_name=selected_sheet,
                                dtype=str,
                                keep_default_na=False,
                                na_values=[]
                            )

                            if excel_df.empty:
                                st.warning("⚠️ 表格为空", icon="⚠️")
                            else:
                                excel_df = clean_excel_types(excel_df)
                                excel_cols = excel_df.columns.tolist()

                                # 显示摘要
                                preview_df = excel_df.head(5)
                                st.dataframe(
                                    preview_df,
                                    use_container_width=True,
                                    hide_index=True,
                                    height=150
                                )

                                col_s1, col_s2 = st.columns(2)
                                with col_s1:
                                    st.metric("行数", len(excel_df))
                                with col_s2:
                                    st.metric("列数", len(excel_cols))

                    finally:
                        try:
                            os.unlink(excel_path)
                        except:
                            pass

                except Exception as e:
                    st.error(f"❌ 读取失败", icon="❌")
            else:
                st.info("请上传Excel文件", icon="ℹ️")

# ==================== 右侧：规则管理 ====================
with col_main_right:
    st.subheader("📋 规则管理")

    # 替换范围 - 紧凑显示
    st.markdown("**替换范围** ℹ️")
    with st.expander("", expanded=False):
        st.markdown("""
        **完整关键词：** 直接替换整个内容
        - 【张三】→【李四】

        **括号内容：** 只替换括号内的文字
        - 保留括号结构，只改变括号里的内容
        """)

    replace_scope = st.radio(
        "模式",
        options=["完整关键词", "括号内容"],
        key="replace_scope_compact",
        horizontal=True,
        label_visibility="collapsed"
    )
    st.session_state.replace_scope = ["替换完整关键词", "仅替换括号内内容"][
        ["完整关键词", "括号内容"].index(replace_scope)]

    st.markdown("---")

    # 规则列表 - 紧凑显示
    st.markdown(f"**规则列表** ({len(st.session_state.replace_rules)}) ℹ️")
    with st.expander("", expanded=False):
        st.markdown("""
        **如何添加规则：**
        1. 从Word预览中复制关键字（如【姓名】）
        2. 粘贴到下方"关键字"输入框
        3. 选择Excel对应的列
        4. 点击"➕ 添加规则"
        """)

    if st.session_state.replace_rules:
        # 规则快速显示 - 高度统一
        with st.container(border=True):
            for idx, (old, col) in enumerate(st.session_state.replace_rules):
                col_del, col_rule = st.columns([0.5, 3], gap="small")
                with col_del:
                    if st.button("❌", key=f"del_{idx}", use_container_width=True,
                                 help="删除此规则"):
                        st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                        st.session_state.replace_rules.pop(idx)
                        st.session_state.replaced_files = []
                        st.rerun()
                with col_rule:
                    st.caption(f"**{old[:12]}** → {col[:12]}")

        # 规则操作按钮
        col_undo, col_clear = st.columns(2, gap="small")
        with col_undo:
            if st.session_state.undo_stack:
                if st.button("↶ 撤销", key="undo", use_container_width=True, help="撤销最后一次操作"):
                    st.session_state.replace_rules = st.session_state.undo_stack.pop()
                    st.success("✅ 已撤销", icon="✅")
                    st.rerun()
        with col_clear:
            if st.button("🗑️ 清空", key="clear_rules", use_container_width=True, help="清空所有规则"):
                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                st.session_state.replace_rules.clear()
                st.session_state.replaced_files = []
                st.rerun()
    else:
        st.info("📁 暂无规则", icon="ℹ️")

    st.markdown("---")

    # 添加规则 - 紧凑
    st.markdown("**新增规则** ℹ️")
    with st.expander("", expanded=False):
        st.markdown("从Word预览中复制关键字粘贴到下面")

    new_keyword = st.text_input(
        "关键字",
        placeholder="如：【姓名】",
        key="new_keyword",
        label_visibility="collapsed",
        help="从Word文档中复制，如【姓名】、（部门）等"
    )

    if excel_cols:
        new_column = st.selectbox(
            "列",
            options=excel_cols,
            key="new_column",
            label_visibility="collapsed",
            help="选择Excel中对应的数据列"
        )
    else:
        new_column = None

    if st.button(
            "➕ 添加规则",
            key="add_rule",
            type="primary",
            disabled=not (new_keyword and new_column),
            use_container_width=True,
            help="点击添加替换规则"
    ):
        rule = (new_keyword.strip(), new_column)
        if rule in st.session_state.replace_rules:
            st.warning("⚠️ 规则已存在", icon="⚠️")
        else:
            st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
            st.session_state.replace_rules.append(rule)
            st.success("✅ 已添加", icon="✅")
            st.rerun()

    st.markdown("---")

    # 规则导入导出 - 折叠显示
    with st.expander("💾 导入/导出/缓存", expanded=False):
        # 导入
        import_file = st.file_uploader(
            "导入JSON",
            type=["json"],
            key="import_rules",
            label_visibility="collapsed",
            help="导入之前导出的规则文件"
        )

        if import_file:
            try:
                rules_data = json.load(import_file)
                valid_rules = []
                for rule in rules_data:
                    if isinstance(rule, dict) and "keyword" in rule and "excel_column" in rule:
                        keyword = str(rule["keyword"]).strip()
                        excel_col = str(rule["excel_column"]).strip()
                        if keyword and excel_col:
                            valid_rules.append((keyword, excel_col))

                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                for rule in valid_rules:
                    if rule not in st.session_state.replace_rules:
                        st.session_state.replace_rules.append(rule)

                st.success(f"✅ 导入{len(valid_rules)}条", icon="✅")
                st.rerun()
            except:
                st.error("❌ 格式错误", icon="❌")

        # 导出
        if st.session_state.replace_rules:
            rules_data = [
                {"keyword": old, "excel_column": col}
                for old, col in st.session_state.replace_rules
            ]
            rules_json = json.dumps(rules_data, ensure_ascii=False, indent=2)

            col_exp1, col_exp2 = st.columns(2, gap="small")
            with col_exp1:
                st.download_button(
                    label="📥 导出JSON",
                    data=rules_json,
                    file_name="rules.json",
                    mime="application/json",
                    key="export_rules",
                    use_container_width=True,
                    help="导出为JSON文件，可以在其他电脑导入"
                )
            with col_exp2:
                if st.button("💾 保存缓存", key="save_cache", use_container_width=True,
                             help="快速保存规则到本地，下次自动加载"):
                    cache_name = f"rules_{datetime.now().strftime('%m%d_%H%M')}"
                    cache_manager.save_rules(st.session_state.replace_rules, cache_name)
                    st.success("✅ 已保存", icon="✅")

st.markdown("---")

# ==================== 底部：执行替换和参数配置 ====================
st.subheader("⚙️ 替换参数配置")

col_config1, col_config2, col_config3, col_config4 = st.columns(4, gap="small")

with col_config1:
    st.markdown("**核心字段** ℹ️")
    with st.expander("", expanded=False):
        st.markdown("选择Excel中的列用于生成文件名")
    file_name_col = st.selectbox(
        "用于文件名",
        options=excel_cols if excel_cols else ["未选择"],
        key="file_name_col",
        disabled=not excel_cols,
        label_visibility="collapsed",
        help="选择用于文件名的列"
    )

with col_config2:
    st.markdown("**起始行** ℹ️")
    with st.expander("", expanded=False):
        st.markdown("从第几行开始处理（第一行是标题）")
    start_row = st.number_input(
        "开始",
        min_value=1,
        max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        value=1,
        key="start_row",
        disabled=excel_df is None or len(excel_df) == 0,
        label_visibility="collapsed",
        help="从第1行开始（跳过标题）"
    )

with col_config3:
    st.markdown("**结束行** ℹ️")
    with st.expander("", expanded=False):
        st.markdown("处理到第几行（包括该行）")
    end_row = st.number_input(
        "结束",
        min_value=1,
        max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        key="end_row",
        disabled=excel_df is None or len(excel_df) == 0,
        label_visibility="collapsed",
        help="到最后一行"
    )

with col_config4:
    st.markdown("**文件前缀** ℹ️")
    with st.expander("", expanded=False):
        st.markdown("给生成的文件名添加前缀，如\"2024-\"")
    file_prefix = st.text_input(
        "前缀",
        value="",
        key="file_prefix",
        placeholder="可选",
        max_chars=15,
        label_visibility="collapsed",
        help="文件名前缀"
    ).strip()

if start_row > end_row:
    st.error("❌ 起始行不能大于结束行", icon="❌")

st.markdown("---")

# ==================== 执行替换 ====================
can_replace = word_file and excel_df is not None and len(excel_df) > 0 and len(st.session_state.replace_rules) > 0

current_params = get_replace_params(
    word_file, excel_df, start_row, end_row, file_name_col, file_prefix, ""
)

need_replace = (
        len(st.session_state.replaced_files) == 0 or
        st.session_state.replace_params != current_params
)

col_exec1, col_exec2, col_exec3, col_exec4 = st.columns([2, 1.5, 1.5, 1], gap="small")

with col_exec1:
    replace_btn = st.button(
        "▶️ 开始替换",
        key="replace",
        disabled=not can_replace or st.session_state.is_replacing or start_row > end_row,
        type="primary",
        use_container_width=True,
        help="点击开始批量替换（需要文件、规则和行数范围）"
    )

with col_exec2:
    if st.session_state.is_replacing:
        st.info("🔄 进行中", icon="🔄")
    elif len(st.session_state.replaced_files) > 0 and not need_replace:
        st.success(f"✅ {len(st.session_state.replaced_files)}个", icon="✅")

with col_exec3:
    pass

with col_exec4:
    pass

# 执行替换逻辑
if replace_btn and not st.session_state.is_replacing:
    st.session_state.is_replacing = True
    st.session_state.replaced_files = []
    st.session_state.replace_log = []

    progress_bar = st.progress(0)
    progress_text = st.empty()

    try:
        actual_end_row = min(end_row, len(excel_df))
        if start_row > actual_end_row:
            st.error("❌ 行号超出范围", icon="❌")
        else:
            total_rows = actual_end_row - start_row + 1

            # 遍历Excel每一行进行替换
            for idx, row_idx in enumerate(range(start_row - 1, actual_end_row)):
                try:
                    excel_row = excel_df.iloc[row_idx]

                    # 调用替换函数
                    replaced_file, replace_log, replace_cnt = replace_word_with_format(
                        word_file, excel_row, st.session_state.replace_rules
                    )

                    # 生成安全的文件名
                    filename = generate_safe_filename(
                        excel_row,
                        file_name_col if file_name_col != "未选择" else "",
                        file_prefix,
                        "",
                        row_idx
                    )

                    # 保存替换后的文件
                    st.session_state.replaced_files.append(ReplacedFile(
                        filename=filename,
                        data=replaced_file,
                        row_idx=row_idx,
                        log=replace_log,
                        replace_count=replace_cnt
                    ))

                    st.session_state.replace_log.append(f"【{row_idx + 1}】{replace_log}")

                    # 更新进度条
                    progress = (idx + 1) / total_rows
                    progress_bar.progress(progress)
                    progress_text.text(f"{idx + 1}/{total_rows}")

                except Exception as e:
                    st.session_state.replace_log.append(f"【{row_idx + 1}】❌ 失败")
                    continue

            # 替换完成，记录参数和历史
            st.session_state.replace_params = current_params
            st.success(f"🎉 完成！{len(st.session_state.replaced_files)} 个文件", icon="✅")

            history_record = HistoryRecord(
                timestamp=datetime.now().strftime("%m-%d %H:%M"),
                word_file=word_file.name[:20],
                excel_file=excel_file.name[:20],
                rules_count=len(st.session_state.replace_rules),
                files_generated=len(st.session_state.replaced_files),
                status="success"
            )
            history_manager.add_record(history_record)

    except Exception as e:
        st.error(f"❌ 出错", icon="❌")
    finally:
        st.session_state.is_replacing = False
        progress_bar.empty()
        progress_text.empty()

st.markdown("---")

# ==================== 下载结果区 ====================
if len(st.session_state.replaced_files) > 0:
    st.subheader("💾 下载结果 ℹ️")
    with st.expander("", expanded=False):
        st.markdown("""
        **两种导出方式：**
        1. **独立文件（ZIP）** - 保存所有文件到一个ZIP压缩包
        2. **合并文档** - 将所有文件合并为一个Word文档，每个文档一页
        """)

    # 选项卡：导出方式
    col_export_opt1, col_export_opt2, col_export_opt3 = st.columns([2, 2, 2], gap="small")

    with col_export_opt1:
        st.markdown("**导出方式**")

    export_mode = st.radio(
        "方式",
        options=["独立文件（ZIP）", "合并为单个文档"],
        key="export_mode_radio",
        horizontal=True,
        label_visibility="collapsed",
        help="选择导出方式"
    )

    st.markdown("---")

    # 统计信息 - 紧凑显示
    col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4, gap="small")

    with col_stat1:
        st.metric("📄 总数", len(st.session_state.replaced_files))

    with col_stat2:
        success_count = len([f for f in st.session_state.replaced_files
                             if f.data and len(f.data.getvalue()) > 0])
        st.metric("✅ 成功", success_count)

    with col_stat3:
        total_replace = sum(f.replace_count for f in st.session_state.replaced_files)
        st.metric("🔄 替换次", total_replace)

    with col_stat4:
        st.metric("📋 规则数", len(st.session_state.replace_rules))

    st.markdown("---")

    # 导出按钮
    col_down1, col_down2, col_down3 = st.columns(3, gap="small")

    with col_down1:
        if export_mode == "独立文件（ZIP）":
            try:
                valid_files = [f for f in st.session_state.replaced_files
                               if f.data and len(f.data.getvalue()) > 0]

                if valid_files:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for file in valid_files:
                            zipf.writestr(file.filename, file.data.getvalue())

                    zip_buffer.seek(0)
                    zip_filename = f"批量替换_{len(valid_files)}个.zip"

                    st.download_button(
                        label=f"📦 下载ZIP（{len(valid_files)}个）",
                        data=zip_buffer,
                        file_name=zip_filename,
                        mime="application/zip",
                        key="download_all_zip",
                        use_container_width=True,
                        type="primary"
                    )
            except:
                st.error("❌ 创建ZIP失败", icon="❌")
        else:
            valid_files = [f for f in st.session_state.replaced_files
                           if f.data and len(f.data.getvalue()) > 0]

            if valid_files:
                try:
                    merged_data = merge_word_documents(valid_files)

                    st.download_button(
                        label=f"📋 下载合并文档（{len(valid_files)}个）",
                        data=merged_data,
                        file_name="合并结果.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_merged",
                        use_container_width=True,
                        type="primary"
                    )
                except:
                    st.error("❌ 合并失败", icon="❌")

    with col_down2:
        # 导出统计
        if st.button("📊 导出统计", key="export_stats", use_container_width=True, help="导出替换统计数据为CSV"):
            csv_data = export_statistics_to_csv(st.session_state.replaced_files)
            st.download_button(
                label="📥 下载CSV统计",
                data=csv_data,
                file_name="统计.csv",
                mime="text/csv",
                key="download_stats",
                use_container_width=True
            )

    with col_down3:
        # 日志导出
        if st.session_state.replace_log:
            log_text = "\n".join(st.session_state.replace_log)
            st.download_button(
                label="📝 导出日志",
                data=log_text,
                file_name="替换日志.txt",
                mime="text/plain",
                key="download_log",
                use_container_width=True,
                help="导出详细的替换操作日志"
            )

    st.markdown("---")

    # 文件列表 - 紧凑显示
    st.markdown(f"**文件列表** ({len(st.session_state.replaced_files)}) ℹ️")
    with st.expander("", expanded=False):
        st.markdown("下方显示所有生成的文件，可以单个下载或查看详细日志")

    # 分页
    total_pages = (len(st.session_state.replaced_files) + PAGE_SIZE - 1) // PAGE_SIZE

    col_page1, col_page2, col_page3 = st.columns([2, 1, 2])

    with col_page2:
        current_page = st.number_input(
            "页",
            min_value=1,
            max_value=total_pages,
            value=1,
            key="current_page",
            label_visibility="collapsed",
            help=f"共{total_pages}页"
        )

    start_idx = (current_page - 1) * PAGE_SIZE
    end_idx = min(start_idx + PAGE_SIZE, len(st.session_state.replaced_files))
    current_files = st.session_state.replaced_files[start_idx:end_idx]

    st.caption(f"第 {current_page}/{total_pages} 页")

    # 文件表格显示
    file_data = []
    for idx, file in enumerate(current_files, start=start_idx + 1):
        is_valid = file.data and len(file.data.getvalue()) > 0
        status = "✅" if is_valid else "❌"
        file_data.append({
            "状态": status,
            "序号": idx,
            "文件名": file.filename[:25] + "..." if len(file.filename) > 25 else file.filename,
            "行号": file.row_idx + 1,
            "替换": file.replace_count
        })

    if file_data:
        file_df = pd.DataFrame(file_data)
        st.dataframe(file_df, use_container_width=True, hide_index=True)

    # 单个文件下载 - 修复文本提示
    st.markdown("**单个文件下载**")

    for idx, file in enumerate(current_files, start=start_idx + 1):
        is_valid = file.data and len(file.data.getvalue()) > 0

        col_name, col_log, col_download = st.columns([2, 1, 1], gap="small")

        with col_name:
            st.caption(f"#{idx} {file.filename}")

        with col_log:
            if st.button("📋 日志", key=f"log_{idx}", use_container_width=True, help="查看替换日志"):
                st.write(file.log)

        with col_download:
            st.download_button(
                label="⬇️ 下载文件",
                data=file.data,
                file_name=file.filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{idx}",
                disabled=not is_valid,
                use_container_width=True,
                help="下载此文件"
            )

else:
    st.info("💡 上传文件、添加规则后点击'开始替换'", icon="ℹ️")

# ==================== 底部帮助 ====================
st.markdown("---")

with st.expander("❓ 帮助指南", expanded=False):
    col_help1, col_help2 = st.columns(2, gap="medium")

    with col_help1:
        st.markdown("""
        **快速开始**
        1. 📤 上传Word和Excel文件
        2. 📋 添加替换规则
        3. ▶️ 点击"开始替换"
        4. 💾 下载结果

        **支持格式**
        • Word：.docx（不支持.doc）
        • Excel：.xlsx/.xls
        • 括号：【】（）()〔〕

        **文件限制**
        • Word最大50MB
        • Excel最大50MB
        • 建议行数<1000
        """)

    with col_help2:
        st.markdown("""
        **常见问题**

        ❓ **Word文件不支持.doc？**
        用Word打开文件 → 另存为.docx格式

        ❓ **怎样保留格式？**
        所有格式自动保留：字体、颜色、表格等

        ❓ **如何合并文档？**
        选择"合并为单个文档"导出方式

        ❓ **能否撤销操作？**
        点击"↶ 撤销"按钮撤销最后一次规则操作

        ❓ **如何加快速度？**
        • 分批处理（每批100-200行）
        • 使用SSD硬盘
        • 关闭其他程序
        """)

st.caption(f"Word+Excel批量替换工具 {VERSION} © 2026")