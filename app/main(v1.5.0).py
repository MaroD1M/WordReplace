# 导入标准库
import os
import sys
import tempfile
from tempfile import NamedTemporaryFile
import warnings
import shutil
import json
import io
import zipfile
import re
import unicodedata
import copy
from datetime import datetime
import hashlib
import base64

# 导入第三方库
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple, Set
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP

# 项目版本信息
VERSION = "v1.5.0"

# 配置常量
PAGE_SIZE = 10
WIDGET_HEIGHT = 250
PREVIEW_ROWS = 20
MAX_FILENAME_LENGTH = 200
MAX_WORD_FILE_SIZE = 50 * 1024 * 1024
MAX_EXCEL_FILE_SIZE = 50 * 1024 * 1024
CACHE_DIR = ".replace_cache"
HISTORY_FILE = ".replace_history.json"
MAX_HISTORY_ITEMS = 30

# 过滤警告
warnings.filterwarnings("ignore", category=UserWarning)

# 环境变量
os.environ["STREAMLIT_VERSION"] = "1.51.0"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# 设置页面配置
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 全局样式优化 - 紧凑布局
st.markdown("""
<style>
    /* 全局间距优化 */
    .main {
        padding: 0.5rem 1rem !important;
    }

    [data-testid="stMainBlockContainer"] {
        padding-top: 0.5rem !important;
        padding-bottom: 0.5rem !important;
    }

    /* 块容器紧凑 */
    .stContainer {
        padding: 0.75rem !important;
        margin-bottom: 0.5rem !important;
        border-radius: 6px;
        background-color: #ffffff;
    }

    /* 删除元素间多余间距 */
    .element-container {
        margin-bottom: 0.3rem !important;
    }

    .stColumn {
        gap: 0.5rem !important;
    }

    /* 按钮紧凑 */
    .stButton > button {
        border-radius: 5px;
        font-weight: 500;
        padding: 0.4rem 0.8rem !important;
        font-size: 13px !important;
        margin-bottom: 0.2rem !important;
    }

    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 2px 8px rgba(0,0,0,0.12);
    }

    /* 输入框紧凑 */
    .stTextInput, .stTextArea, .stSelectbox, .stNumberInput {
        margin-bottom: 0.3rem !important;
    }

    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select,
    .stNumberInput > div > div > input {
        border-radius: 5px;
        border: 1px solid #e0e0e0;
        font-size: 13px;
        padding: 0.5rem !important;
    }

    /* 标题紧凑 */
    h1 {
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #1f77b4;
        margin-bottom: 0.5rem;
    }

    h2 {
        margin-top: 0.5rem;
        margin-bottom: 0.3rem;
        color: #1f77b4;
        font-size: 1.2rem;
    }

    h3 {
        margin-top: 0.3rem;
        margin-bottom: 0.2rem;
        color: #333;
        font-size: 1.05rem;
    }

    .stSubheader {
        margin-bottom: 0.5rem !important;
        padding-bottom: 0.3rem;
        border-bottom: 1.5px solid #e0e0e0;
        font-size: 1.1rem !important;
    }

    /* 展开器紧凑 */
    .streamlit-expander {
        margin-bottom: 0.3rem !important;
        border-radius: 5px;
    }

    /* 标签页紧凑 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1px;
        margin-bottom: 0.5rem;
    }

    .stTabs [data-baseweb="tab"] {
        height: 40px;
        padding-top: 8px;
        border-radius: 5px 5px 0 0;
        font-size: 13px;
    }

    /* 数据框紧凑 */
    div[data-testid="stDataFrame"] {
        border-radius: 5px;
        border: 1px solid #e0e0e0;
        font-size: 12px;
    }

    /* 指标卡紧凑 */
    .metric-container {
        background-color: #f8f9fa;
        padding: 0.5rem !important;
        border-radius: 5px;
        border-left: 3px solid #1f77b4;
        margin-bottom: 0.3rem;
    }

    /* 信息框紧凑 */
    .stats-box, .success-box, .warning-box, .error-box {
        padding: 0.6rem !important;
        margin: 0.3rem 0 !important;
        border-radius: 5px;
        border-left-width: 3px;
        font-size: 13px;
    }

    .stats-box {
        background-color: #f0f9ff;
        border-left-color: #0ea5e9;
    }

    .success-box {
        background-color: #f0fdf4;
        border-left-color: #22c55e;
    }

    .warning-box {
        background-color: #fffbeb;
        border-left-color: #f59e0b;
    }

    .error-box {
        background-color: #fef2f2;
        border-left-color: #ef4444;
    }

    /* 分隔线 */
    hr {
        margin: 0.5rem 0 !important;
        border: none;
        border-top: 1px solid #e0e0e0;
    }

    /* 无线电和复选框紧凑 */
    .stRadio, .stCheckbox {
        margin-bottom: 0.2rem !important;
    }

    .stRadio > label, .stCheckbox > label {
        margin-bottom: 0.2rem !important;
        font-size: 13px;
    }

    /* 文件上传器紧凑 */
    .stFileUploader {
        margin-bottom: 0.3rem !important;
    }

    /* 进度条 */
    .stProgress {
        margin-bottom: 0.3rem !important;
    }

    /* 表格紧凑 */
    table {
        font-size: 12px !important;
    }

    td, th {
        padding: 0.4rem !important;
    }

    /* 响应式 */
    @media (max-width: 768px) {
        .main {
            padding: 0.3rem 0.5rem;
        }
        .stContainer {
            padding: 0.5rem;
        }
    }
</style>
""", unsafe_allow_html=True)


# ---------------------- 数据结构与初始化 ----------------------

@dataclass
class ReplacedFile:
    """存储替换后的文件数据结构"""
    filename: str
    data: io.BytesIO
    row_idx: int
    log: str
    replace_count: int = 0


@dataclass
class HistoryRecord:
    """历史记录数据结构"""
    timestamp: str
    word_file: str
    excel_file: str
    rules_count: int
    files_generated: int
    status: str


class CacheManager:
    """缓存管理器"""

    def __init__(self):
        self.cache_dir = CACHE_DIR
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)

    def save_rules(self, rules: List[Tuple[str, str]], filename: str):
        """保存规则到缓存"""
        try:
            rules_data = [{"keyword": old, "excel_column": col} for old, col in rules]
            cache_file = os.path.join(self.cache_dir, f"{filename}.json")
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(rules_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"⚠️ 保存缓存失败", icon="⚠️")

    def load_rules(self, filename: str) -> List[Tuple[str, str]]:
        """加载缓存的规则"""
        try:
            cache_file = os.path.join(self.cache_dir, f"{filename}.json")
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    rules_data = json.load(f)
                    return [(r["keyword"], r["excel_column"]) for r in rules_data]
        except:
            pass
        return []

    def get_cached_rules_list(self) -> List[str]:
        """获取所有缓存的规则文件"""
        try:
            if os.path.exists(self.cache_dir):
                files = [f.replace('.json', '') for f in os.listdir(self.cache_dir) if f.endswith('.json')]
                return sorted(files, reverse=True)[:10]  # 只返回最近10个
        except:
            pass
        return []


class HistoryManager:
    """历史记录管理器"""

    def __init__(self):
        self.history_file = HISTORY_FILE

    def add_record(self, record: HistoryRecord):
        """添加历史记录"""
        try:
            history = self.load_history()
            history.insert(0, {
                "timestamp": record.timestamp,
                "word_file": record.word_file,
                "excel_file": record.excel_file,
                "rules_count": record.rules_count,
                "files_generated": record.files_generated,
                "status": record.status
            })
            history = history[:MAX_HISTORY_ITEMS]
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except:
            pass

    def load_history(self) -> List[Dict]:
        """加载历史记录"""
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return []

    def clear_history(self):
        """清除历史记录"""
        try:
            if os.path.exists(self.history_file):
                os.remove(self.history_file)
                st.success("✅ 历史已清除", icon="✅")
        except:
            pass


def init_session_state():
    """初始化会话状态"""
    required_states = {
        "replace_rules": [],
        "replaced_files": [],
        "replace_log": [],
        "is_replacing": False,
        "replace_params": {},
        "replace_scope": "替换完整关键词",
        "export_mode_radio": "独立文件（ZIP压缩）",
        "undo_stack": [],
        "rule_filter": "",
        "show_advanced": False,
    }

    for key, default in required_states.items():
        if key not in st.session_state:
            st.session_state[key] = default


init_session_state()


# ---------------------- 核心工具函数 ----------------------

def clean_text(text: str) -> str:
    """清理文本"""
    if not isinstance(text, str):
        return ""
    text = text.strip()
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'[\u00A0\u2002-\u200B]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text


def clean_filename(filename: str) -> str:
    """清理文件名非法字符"""
    return re.sub(r'[\\/:*?"<>|]', "_", str(filename))


def generate_safe_filename(
        excel_row: pd.Series,
        file_name_col: str,
        file_prefix: str = "",
        file_suffix: str = "",
        row_idx: int = 0,
        max_length: int = MAX_FILENAME_LENGTH
) -> str:
    """安全生成文件名"""
    try:
        if file_name_col and file_name_col in excel_row.index:
            base_name = clean_text(str(excel_row[file_name_col]))
        else:
            base_name = f"文件_{row_idx + 1}"

        if not base_name or base_name.isspace():
            base_name = f"文件_{row_idx + 1}"

        if file_prefix and file_suffix:
            filename = f"{file_prefix}{base_name}{file_suffix}.docx"
        elif file_prefix:
            filename = f"{file_prefix}{base_name}.docx"
        elif file_suffix:
            filename = f"{base_name}{file_suffix}.docx"
        else:
            filename = f"{base_name}.docx"

        filename = clean_filename(filename)

        filename_bytes = filename.encode('utf-8')
        if len(filename_bytes) > max_length:
            truncated_base = base_name
            while len(f"{file_prefix}{truncated_base}{file_suffix}.docx".encode('utf-8')) > max_length:
                truncated_base = truncated_base[:-1]

            if file_prefix and file_suffix:
                filename = f"{file_prefix}{truncated_base}{file_suffix}.docx"
            elif file_prefix:
                filename = f"{file_prefix}{truncated_base}.docx"
            elif file_suffix:
                filename = f"{truncated_base}{file_suffix}.docx"
            else:
                filename = f"{truncated_base}.docx"

            filename = clean_filename(filename)

        return filename

    except:
        return f"文件_{row_idx + 1}.docx"


def precompute_replace_patterns(
        replace_rules: List[Tuple[str, str]],
        excel_row: pd.Series
) -> List[Tuple[str, str, str, str]]:
    """预计算替换模式"""
    replace_patterns = []

    for old_text, col_name in replace_rules:
        if col_name in excel_row.index:
            replacement = str(excel_row[col_name]).strip()
        else:
            replacement = ""

        cleaned_text = clean_text(old_text)

        if not cleaned_text:
            continue

        if st.session_state.replace_scope == "仅替换括号内内容":
            if cleaned_text.startswith("【") and cleaned_text.endswith("】"):
                new_format = f"【{replacement}】"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("（") and cleaned_text.endswith("）"):
                new_format = f"（{replacement}）"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("(") and cleaned_text.endswith(")"):
                new_format = f"({replacement})"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("〔") and cleaned_text.endswith("〕"):
                new_format = f"〔{replacement}〕"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            else:
                replace_patterns.append((old_text, col_name, cleaned_text, replacement))
        else:
            replace_patterns.append((old_text, col_name, cleaned_text, replacement))

    return replace_patterns


def process_paragraph(
        paragraph,
        replace_patterns: List[Tuple[str, str, str, str]],
        cleaned_para: str = None
) -> Dict:
    """处理段落替换"""
    para_text = paragraph.text
    if cleaned_para is None:
        cleaned_para = clean_text(para_text)
    replace_count = defaultdict(int)

    if not para_text or not replace_patterns:
        return replace_count

    has_keyword = False

    for old_text, col_name, format_keyword, replacement in replace_patterns:
        if format_keyword and format_keyword in cleaned_para:
            has_keyword = True
            break

    if has_keyword:
        new_text = para_text
        for old_text, col_name, format_keyword, replacement in replace_patterns:
            if format_keyword and format_keyword in cleaned_para:
                count = new_text.count(format_keyword)
                if count > 0:
                    new_text = new_text.replace(format_keyword, replacement)
                    replace_count[(old_text, col_name)] += count

        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ''

    return replace_count


def replace_word_with_format(
        word_file: st.runtime.uploaded_file_manager.UploadedFile,
        excel_row: pd.Series,
        replace_rules: List[Tuple[str, str]]
) -> Tuple[io.BytesIO, str, int]:
    """替换Word文件"""
    replace_count = defaultdict(int)
    total_replace = 0

    try:
        file_size = len(word_file.getvalue())
        if file_size > MAX_WORD_FILE_SIZE:
            raise ValueError(f"文件过大")

        doc = Document(io.BytesIO(word_file.getvalue()))

        replace_patterns = precompute_replace_patterns(replace_rules, excel_row)

        if not replace_patterns:
            output_file = io.BytesIO()
            doc.save(output_file)
            output_file.seek(0)
            return output_file, "⚠ 未找到匹配规则", 0

        for paragraph in doc.paragraphs:
            para_count = process_paragraph(paragraph, replace_patterns)
            for key, count in para_count.items():
                replace_count[key] += count
                total_replace += count

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_count = process_paragraph(paragraph, replace_patterns)
                        for key, count in para_count.items():
                            replace_count[key] += count
                            total_replace += count

        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)

        if replace_count:
            log_lines = [f"✓ {old}" for old, _ in replace_count.keys()]
            replace_log = ", ".join(log_lines[:3])
            if len(replace_count) > 3:
                replace_log += f" 等{len(replace_count) - 3}个"
        else:
            replace_log = "⚠ 无替换"

        return output_file, replace_log, total_replace

    except Exception as e:
        return io.BytesIO(), f"❌ 失败", 0


def merge_word_documents(
        replaced_files: List[ReplacedFile]
) -> io.BytesIO:
    """合并Word文档"""
    if not replaced_files:
        raise ValueError("没有文件")

    try:
        main_doc = Document(io.BytesIO(replaced_files[0].data.getvalue()))
        main_body = main_doc._body._element

        for idx in range(1, len(replaced_files)):
            try:
                file = replaced_files[idx]

                if not file.data or len(file.data.getvalue()) == 0:
                    continue

                sub_doc = Document(io.BytesIO(file.data.getvalue()))
                sub_body = sub_doc._body._element

                page_break_para = OxmlElement('w:p')
                page_break_pPr = OxmlElement('w:pPr')

                page_break_element = OxmlElement('w:pageBreakBefore')
                page_break_element.set(qn('w:val'), '1')

                page_break_pPr.append(page_break_element)
                page_break_para.append(page_break_pPr)
                main_body.append(page_break_para)

                for element in sub_body:
                    main_body.append(copy.deepcopy(element))

            except:
                continue

        output = io.BytesIO()
        main_doc.save(output)
        output.seek(0)
        return output

    except Exception as e:
        raise


def get_replace_params(
        word_file: Optional[st.runtime.uploaded_file_manager.UploadedFile],
        excel_df: Optional[pd.DataFrame],
        start_row: int,
        end_row: int,
        file_name_col: str,
        file_prefix: str,
        file_suffix: str
) -> Dict:
    """获取替换参数"""
    return {
        "word_filename": word_file.name if word_file else "",
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "rule_count": len(st.session_state.replace_rules),
        "rule_hash": hash(tuple(st.session_state.replace_rules))
    }


def fix_float_precision(x: str, column_name: Optional[str] = None) -> str:
    """修复浮点数精度"""
    if not x or not isinstance(x, str):
        return x

    x = x.strip()
    if not x:
        return ""

    try:
        if x.replace('.', '', 1).replace('-', '', 1).isdigit():
            pass
        else:
            return x
    except:
        return x

    float_pattern = r'^\s*[-+]?\d*\.?\d+\s*$'
    if not re.match(float_pattern, x):
        return x

    try:
        dec_value = Decimal(x)

        if dec_value.as_tuple().exponent >= 0:
            return str(int(dec_value))

        float_val = float(dec_value)
        float_str = str(float_val)

        if '999999' in float_str or '000000' in float_str:
            if '.' in x:
                orig_dec_part = x.split('.')[1]
                orig_dec_places = len(orig_dec_part.rstrip('0'))

                if orig_dec_places > 0:
                    try:
                        quantized = dec_value.quantize(
                            Decimal('1.' + '0' * orig_dec_places),
                            rounding=ROUND_HALF_UP
                        )
                        result = format(quantized, f'.{orig_dec_places}f')
                        return result.rstrip('0').rstrip('.')
                    except:
                        pass

        return x
    except:
        return x


def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """清理Excel数据类型"""
    df_clean = df.copy()

    for col in df_clean.columns:
        try:
            col_name = str(col)
            if col_name != col:
                df_clean = df_clean.rename(columns={col: col_name})
                col = col_name

            df_clean[col] = df_clean[col].fillna("")
            df_clean[col] = df_clean[col].astype(str).str.strip()
            df_clean[col] = df_clean[col].apply(lambda x: fix_float_precision(x, col))

        except:
            try:
                df_clean[col] = df_clean[col].astype(str).str.strip()
            except:
                pass

    return df_clean


def get_file_hash(file_data: bytes) -> str:
    """获取文件哈希值"""
    return hashlib.md5(file_data).hexdigest()[:6]


def export_statistics_to_csv(replaced_files: List[ReplacedFile]) -> str:
    """导出统计数据到CSV"""
    try:
        data = []
        for idx, file in enumerate(replaced_files, 1):
            data.append({
                "序号": idx,
                "文件名": file.filename,
                "行号": file.row_idx + 1,
                "替换次数": file.replace_count,
                "状态": "✅" if file.data and len(file.data.getvalue()) > 0 else "❌"
            })

        df = pd.DataFrame(data)
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
        return csv_buffer.getvalue()
    except:
        return ""


def get_keyword_statistics(replace_rules: List[Tuple[str, str]],
                           replaced_files: List[ReplacedFile]) -> Dict:
    """获取关键字替换统计"""
    stats = {}
    for keyword, _ in replace_rules:
        stats[keyword] = 0

    for file in replaced_files:
        for keyword, _ in replace_rules:
            if f"✓ {keyword}" in file.log:
                pattern = f"✓ {re.escape(keyword)}.*?\((\d+)次\)"
                matches = re.findall(pattern, file.log)
                if matches:
                    stats[keyword] += int(matches[0])

    return stats


# 创建管理器
cache_manager = CacheManager()
history_manager = HistoryManager()

# ===================== 侧栏 =====================
with st.sidebar:
    st.title("📚 快速导航")

    # 统计信息
    if st.session_state.replaced_files:
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 文件数", len(st.session_state.replaced_files), delta=None)
        with col2:
            st.metric("📋 规则数", len(st.session_state.replace_rules), delta=None)

    st.markdown("---")

    # 快速功能
    st.subheader("⚡ 快速功能")

    cached = cache_manager.get_cached_rules_list()
    if cached:
        selected = st.selectbox("📂 加载规则", ["选择..."] + cached, key="sidebar_cache")
        if selected and selected != "选择...":
            if st.button("✅ 加载", key="sidebar_load", use_container_width=True):
                loaded = cache_manager.load_rules(selected)
                if loaded:
                    st.session_state.replace_rules = loaded
                    st.success(f"✅ 加载{len(loaded)}条", icon="✅")
                    st.rerun()

    # 历史记录
    history = history_manager.load_history()
    if history:
        st.subheader("📜 最近操作")
        for h in history[:3]:
            status = "✅" if h["status"] == "success" else "❌"
            st.caption(f"{status} {h['timestamp']}\n{h['word_file'][:15]}...")

    st.markdown("---")

    # 工具
    st.subheader("🔧 工具")

    if st.button("🗑️ 清空所有", key="sidebar_clear", use_container_width=True):
        st.session_state.replace_rules = []
        st.session_state.replaced_files = []
        st.success("✅ 已清空", icon="✅")
        st.rerun()

    if history:
        if st.button("📜 清除历史", key="sidebar_clear_hist", use_container_width=True):
            history_manager.clear_history()
            st.rerun()

# ===================== 主页面 =====================
st.title("📋 Word+Excel批量替换工具")

# 进度显示（如果有替换任务）
if st.session_state.replaced_files and st.session_state.replace_params:
    progress_col, status_col = st.columns([3, 1])
    with progress_col:
        success_count = len([f for f in st.session_state.replaced_files
                             if f.data and len(f.data.getvalue()) > 0])
        total_count = len(st.session_state.replaced_files)
        st.progress(success_count / total_count if total_count > 0 else 0)
    with status_col:
        st.metric("成功率", f"{int(success_count / total_count * 100) if total_count > 0 else 0}%")

st.markdown("---")

# ==================== 主工作区 ====================
col_main_left, col_main_right = st.columns([2, 1], gap="medium")

# ==================== 左侧：文件上传和预览 ====================
with col_main_left:
    st.subheader("📤 文件上传")

    # 上传区域 - 两列并排
    col_upload1, col_upload2 = st.columns(2, gap="small")

    with col_upload1:
        st.markdown("**Word**")
        word_file = st.file_uploader(
            "选择文件",
            type=["docx"],
            key="word",
            label_visibility="collapsed",
            help="仅支持.docx"
        )
        if word_file:
            file_size_mb = len(word_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_WORD_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ 过大", icon="❌")
                word_file = None
            else:
                st.caption(f"✅ {file_size_mb:.1f}MB")

    with col_upload2:
        st.markdown("**Excel**")
        excel_file = st.file_uploader(
            "选择文件",
            type=["xlsx", "xls"],
            key="excel",
            label_visibility="collapsed",
            help="支持.xlsx/.xls"
        )
        if excel_file:
            file_size_mb = len(excel_file.getvalue()) / 1024 / 1024
            if file_size_mb > MAX_EXCEL_FILE_SIZE / 1024 / 1024:
                st.error(f"❌ 过大", icon="❌")
                excel_file = None
            else:
                st.caption(f"✅ {file_size_mb:.1f}MB")

    st.markdown("---")

    # 文件预览 - 可折叠
    with st.expander("👀 文件预览", expanded=True):
        col_prev1, col_prev2 = st.columns(2, gap="small")

        excel_df = None
        excel_cols = []

        with col_prev1:
            st.markdown("**Word内容**")
            if word_file:
                try:
                    doc = Document(io.BytesIO(word_file.getvalue()))
                    preview_text = ""
                    para_count = 0

                    for paragraph in doc.paragraphs:
                        if para_count >= 10:
                            preview_text += "\n... (内容已截断)"
                            break
                        if paragraph.text.strip():
                            preview_text += paragraph.text[:50] + "\n"
                            para_count += 1

                    st.text_area(
                        "预览",
                        value=preview_text[:500] if preview_text else "空文档",
                        height=120,
                        disabled=True,
                        label_visibility="collapsed"
                    )
                    st.caption(f"📄 {len(doc.paragraphs)}段落，{len(doc.tables)}表格")
                except:
                    st.warning("❌ 预览失败", icon="⚠️")
            else:
                st.info("请上传Word文件", icon="ℹ️")

        with col_prev2:
            st.markdown("**Excel内容**")
            if excel_file:
                try:
                    with NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_excel:
                        temp_excel.write(excel_file.getvalue())
                        excel_path = temp_excel.name

                    try:
                        with pd.ExcelFile(excel_path, engine="openpyxl") as excel_wb:
                            sheet_names = excel_wb.sheet_names
                            selected_sheet = sheet_names[0]

                            excel_df = pd.read_excel(
                                excel_wb,
                                sheet_name=selected_sheet,
                                dtype=str,
                                keep_default_na=False,
                                na_values=[]
                            )

                            if excel_df.empty:
                                st.warning("⚠️ 表格为空", icon="⚠️")
                            else:
                                excel_df = clean_excel_types(excel_df)
                                excel_cols = excel_df.columns.tolist()

                                # 显示摘要
                                preview_df = excel_df.head(5)
                                st.dataframe(
                                    preview_df,
                                    use_container_width=True,
                                    hide_index=True,
                                    height=120
                                )

                                col_s1, col_s2 = st.columns(2)
                                with col_s1:
                                    st.metric("行数", len(excel_df))
                                with col_s2:
                                    st.metric("列数", len(excel_cols))

                    finally:
                        try:
                            os.unlink(excel_path)
                        except:
                            pass

                except Exception as e:
                    st.warning("❌ 读取失败", icon="⚠️")
            else:
                st.info("请上传Excel文件", icon="ℹ️")

# ==================== 右侧：规则管理 ====================
with col_main_right:
    st.subheader("📋 规则管理")

    # 替换范围 - 紧凑显示
    st.markdown("**替换范围**")
    replace_scope = st.radio(
        "模式",
        options=["完整关键词", "括号内容"],
        key="replace_scope_compact",
        horizontal=True,
        label_visibility="collapsed"
    )
    st.session_state.replace_scope = ["替换完整关键词", "仅替换括号内内容"][
        ["完整关键词", "括号内容"].index(replace_scope)]

    st.markdown("---")

    # 规则列表 - 紧凑显示
    st.markdown(f"**规则列表** ({len(st.session_state.replace_rules)})")

    if st.session_state.replace_rules:
        # 规则快速显示
        with st.container(height=200, border=True):
            for idx, (old, col) in enumerate(st.session_state.replace_rules):
                col_del, col_rule = st.columns([0.5, 3], gap="small")
                with col_del:
                    if st.button("❌", key=f"del_{idx}", use_container_width=True,
                                 help="删除此规则"):
                        st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                        st.session_state.replace_rules.pop(idx)
                        st.session_state.replaced_files = []
                        st.rerun()
                with col_rule:
                    st.caption(f"**{old[:15]}** → {col[:15]}")

        # 规则操作按钮
        col_undo, col_clear = st.columns(2, gap="small")
        with col_undo:
            if st.session_state.undo_stack:
                if st.button("↶ 撤销", key="undo", use_container_width=True):
                    st.session_state.replace_rules = st.session_state.undo_stack.pop()
                    st.success("✅ 已撤销", icon="✅")
                    st.rerun()
        with col_clear:
            if st.button("🗑️ 清空", key="clear_rules", use_container_width=True):
                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                st.session_state.replace_rules.clear()
                st.session_state.replaced_files = []
                st.rerun()
    else:
        st.info("📁 暂无规则", icon="ℹ️")

    st.markdown("---")

    # 添加规则 - 紧凑
    st.markdown("**新增规则**")
    new_keyword = st.text_input(
        "关键字",
        placeholder="如：【姓名】",
        key="new_keyword",
        label_visibility="collapsed"
    )

    if excel_cols:
        new_column = st.selectbox(
            "列",
            options=excel_cols,
            key="new_column",
            label_visibility="collapsed"
        )
    else:
        new_column = None

    if st.button(
            "➕ 添加规则",
            key="add_rule",
            type="primary",
            disabled=not (new_keyword and new_column),
            use_container_width=True
    ):
        rule = (new_keyword.strip(), new_column)
        if rule in st.session_state.replace_rules:
            st.warning("⚠️ 已存在", icon="⚠️")
        else:
            st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
            st.session_state.replace_rules.append(rule)
            st.success("✅ 已添加", icon="✅")
            st.rerun()

    st.markdown("---")

    # 规则导入导出 - 折叠显示
    with st.expander("💾 导入/导出/缓存", expanded=False):
        # 导入
        import_file = st.file_uploader(
            "导入JSON",
            type=["json"],
            key="import_rules",
            label_visibility="collapsed"
        )

        if import_file:
            try:
                rules_data = json.load(import_file)
                valid_rules = []
                for rule in rules_data:
                    if isinstance(rule, dict) and "keyword" in rule and "excel_column" in rule:
                        keyword = str(rule["keyword"]).strip()
                        excel_col = str(rule["excel_column"]).strip()
                        if keyword and excel_col:
                            valid_rules.append((keyword, excel_col))

                st.session_state.undo_stack.append(st.session_state.replace_rules.copy())
                for rule in valid_rules:
                    if rule not in st.session_state.replace_rules:
                        st.session_state.replace_rules.append(rule)

                st.success(f"✅ 导入{len(valid_rules)}条", icon="✅")
                st.rerun()
            except:
                st.error("❌ 格式错误", icon="❌")

        # 导出
        if st.session_state.replace_rules:
            rules_data = [
                {"keyword": old, "excel_column": col}
                for old, col in st.session_state.replace_rules
            ]
            rules_json = json.dumps(rules_data, ensure_ascii=False, indent=2)

            col_exp1, col_exp2 = st.columns(2, gap="small")
            with col_exp1:
                st.download_button(
                    label="📥 导出JSON",
                    data=rules_json,
                    file_name="rules.json",
                    mime="application/json",
                    key="export_rules",
                    use_container_width=True
                )
            with col_exp2:
                if st.button("💾 保存缓存", key="save_cache", use_container_width=True):
                    cache_name = f"rules_{datetime.now().strftime('%m%d_%H%M')}"
                    cache_manager.save_rules(st.session_state.replace_rules, cache_name)
                    st.success("✅ 已保存", icon="✅")

st.markdown("---")

# ==================== 底部：执行替换和参数配置 ====================
st.subheader("⚙️ 替换参数配置")

col_config1, col_config2, col_config3, col_config4 = st.columns(4, gap="small")

with col_config1:
    st.markdown("**核心字段**")
    file_name_col = st.selectbox(
        "用于文件名",
        options=excel_cols if excel_cols else ["未选择"],
        key="file_name_col",
        disabled=not excel_cols,
        label_visibility="collapsed"
    )

with col_config2:
    st.markdown("**起始行**")
    start_row = st.number_input(
        "开始",
        min_value=1,
        max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        value=1,
        key="start_row",
        disabled=excel_df is None or len(excel_df) == 0,
        label_visibility="collapsed"
    )

with col_config3:
    st.markdown("**结束行**")
    end_row = st.number_input(
        "结束",
        min_value=1,
        max_value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        value=len(excel_df) if excel_df is not None and len(excel_df) > 0 else 1,
        key="end_row",
        disabled=excel_df is None or len(excel_df) == 0,
        label_visibility="collapsed"
    )

with col_config4:
    st.markdown("**文件前缀**")
    file_prefix = st.text_input(
        "前缀",
        value="",
        key="file_prefix",
        placeholder="可选",
        max_chars=15,
        label_visibility="collapsed"
    ).strip()

if start_row > end_row:
    st.error("❌ 起始行不能大于结束行", icon="❌")

st.markdown("---")

# ==================== 执行替换 ====================
can_replace = word_file and excel_df is not None and len(excel_df) > 0 and len(st.session_state.replace_rules) > 0

current_params = get_replace_params(
    word_file, excel_df, start_row, end_row, file_name_col, file_prefix, ""
)

need_replace = (
        len(st.session_state.replaced_files) == 0 or
        st.session_state.replace_params != current_params
)

col_exec1, col_exec2, col_exec3, col_exec4 = st.columns([2, 1.5, 1.5, 1], gap="small")

with col_exec1:
    replace_btn = st.button(
        "▶️ 开始替换",
        key="replace",
        disabled=not can_replace or st.session_state.is_replacing or start_row > end_row,
        type="primary",
        use_container_width=True
    )

with col_exec2:
    if st.session_state.is_replacing:
        st.info("🔄 进行中", icon="🔄")
    elif len(st.session_state.replaced_files) > 0 and not need_replace:
        st.success(f"✅ {len(st.session_state.replaced_files)}个", icon="✅")

with col_exec3:
    pass

with col_exec4:
    pass

# 执行替换逻辑
if replace_btn and not st.session_state.is_replacing:
    st.session_state.is_replacing = True
    st.session_state.replaced_files = []
    st.session_state.replace_log = []

    progress_bar = st.progress(0)
    progress_text = st.empty()

    try:
        actual_end_row = min(end_row, len(excel_df))
        if start_row > actual_end_row:
            st.error("❌ 行号超出范围", icon="❌")
        else:
            total_rows = actual_end_row - start_row + 1

            for idx, row_idx in enumerate(range(start_row - 1, actual_end_row)):
                try:
                    excel_row = excel_df.iloc[row_idx]

                    replaced_file, replace_log, replace_cnt = replace_word_with_format(
                        word_file, excel_row, st.session_state.replace_rules
                    )

                    filename = generate_safe_filename(
                        excel_row,
                        file_name_col if file_name_col != "未选择" else "",
                        file_prefix,
                        "",
                        row_idx
                    )

                    st.session_state.replaced_files.append(ReplacedFile(
                        filename=filename,
                        data=replaced_file,
                        row_idx=row_idx,
                        log=replace_log,
                        replace_count=replace_cnt
                    ))

                    st.session_state.replace_log.append(f"【{row_idx + 1}】{replace_log}")

                    progress = (idx + 1) / total_rows
                    progress_bar.progress(progress)
                    progress_text.text(f"{idx + 1}/{total_rows}")

                except Exception as e:
                    st.session_state.replace_log.append(f"【{row_idx + 1}】❌ 失败")
                    continue

            st.session_state.replace_params = current_params
            st.success(f"🎉 完成！{len(st.session_state.replaced_files)} 个文件", icon="✅")

            history_record = HistoryRecord(
                timestamp=datetime.now().strftime("%m-%d %H:%M"),
                word_file=word_file.name[:20],
                excel_file=excel_file.name[:20],
                rules_count=len(st.session_state.replace_rules),
                files_generated=len(st.session_state.replaced_files),
                status="success"
            )
            history_manager.add_record(history_record)

    except Exception as e:
        st.error(f"❌ 出错", icon="❌")
    finally:
        st.session_state.is_replacing = False
        progress_bar.empty()
        progress_text.empty()

st.markdown("---")

# ==================== 下载结果区 ====================
if len(st.session_state.replaced_files) > 0:
    st.subheader("💾 下载结果")

    # 选项卡：导出方式
    col_export_opt1, col_export_opt2, col_export_opt3 = st.columns([2, 2, 2], gap="small")

    with col_export_opt1:
        st.markdown("**导出方式**")

    export_mode = st.radio(
        "方式",
        options=["独立文件（ZIP）", "合并为单个文档"],
        key="export_mode_radio",
        horizontal=True,
        label_visibility="collapsed"
    )

    st.markdown("---")

    # 统计信息 - 紧凑显示
    col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4, gap="small")

    with col_stat1:
        st.metric("📄 总数", len(st.session_state.replaced_files))

    with col_stat2:
        success_count = len([f for f in st.session_state.replaced_files
                             if f.data and len(f.data.getvalue()) > 0])
        st.metric("✅ 成功", success_count)

    with col_stat3:
        total_replace = sum(f.replace_count for f in st.session_state.replaced_files)
        st.metric("🔄 替换次数", total_replace)

    with col_stat4:
        st.metric("📋 规则数", len(st.session_state.replace_rules))

    st.markdown("---")

    # 导出按钮
    col_down1, col_down2, col_down3 = st.columns(3, gap="small")

    with col_down1:
        if export_mode == "独立文件（ZIP）":
            try:
                valid_files = [f for f in st.session_state.replaced_files
                               if f.data and len(f.data.getvalue()) > 0]

                if valid_files:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for file in valid_files:
                            zipf.writestr(file.filename, file.data.getvalue())

                    zip_buffer.seek(0)
                    zip_filename = f"批量替换_{len(valid_files)}个.zip"

                    st.download_button(
                        label=f"📦 下载ZIP",
                        data=zip_buffer,
                        file_name=zip_filename,
                        mime="application/zip",
                        key="download_all_zip",
                        use_container_width=True,
                        type="primary"
                    )
            except:
                st.error("❌ 创建ZIP失败", icon="❌")
        else:
            valid_files = [f for f in st.session_state.replaced_files
                           if f.data and len(f.data.getvalue()) > 0]

            if valid_files:
                try:
                    merged_data = merge_word_documents(valid_files)

                    st.download_button(
                        label=f"📋 下载合并文档",
                        data=merged_data,
                        file_name="合并结果.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_merged",
                        use_container_width=True,
                        type="primary"
                    )
                except:
                    st.error("❌ 合并失败", icon="❌")

    with col_down2:
        # 导出统计
        if st.button("📊 导出统计", key="export_stats", use_container_width=True):
            csv_data = export_statistics_to_csv(st.session_state.replaced_files)
            st.download_button(
                label="📥 下载CSV",
                data=csv_data,
                file_name="统计.csv",
                mime="text/csv",
                key="download_stats",
                use_container_width=True
            )

    with col_down3:
        # 日志导出
        if st.session_state.replace_log:
            log_text = "\n".join(st.session_state.replace_log)
            st.download_button(
                label="📝 导出日志",
                data=log_text,
                file_name="替换日志.txt",
                mime="text/plain",
                key="download_log",
                use_container_width=True
            )

    st.markdown("---")

    # 文件列表 - 紧凑显示
    st.markdown(f"**文件列表** ({len(st.session_state.replaced_files)})")

    # 分页
    total_pages = (len(st.session_state.replaced_files) + PAGE_SIZE - 1) // PAGE_SIZE

    col_page1, col_page2, col_page3 = st.columns([2, 1, 2])

    with col_page2:
        current_page = st.number_input(
            "页",
            min_value=1,
            max_value=total_pages,
            value=1,
            key="current_page",
            label_visibility="collapsed"
        )

    start_idx = (current_page - 1) * PAGE_SIZE
    end_idx = min(start_idx + PAGE_SIZE, len(st.session_state.replaced_files))
    current_files = st.session_state.replaced_files[start_idx:end_idx]

    st.caption(f"第 {current_page}/{total_pages} 页")

    # 文件表格显示
    file_data = []
    for idx, file in enumerate(current_files, start=start_idx + 1):
        is_valid = file.data and len(file.data.getvalue()) > 0
        status = "✅" if is_valid else "❌"
        file_data.append({
            "状态": status,
            "序号": idx,
            "文件名": file.filename[:25] + "..." if len(file.filename) > 25 else file.filename,
            "行号": file.row_idx + 1,
            "替换": file.replace_count
        })

    if file_data:
        file_df = pd.DataFrame(file_data)
        st.dataframe(file_df, use_container_width=True, hide_index=True)

    # 单个文件下载
    st.markdown("**单个文件下载**")

    for idx, file in enumerate(current_files, start=start_idx + 1):
        is_valid = file.data and len(file.data.getvalue()) > 0

        col_name, col_log, col_btn = st.columns([2, 1, 1], gap="small")

        with col_name:
            st.caption(f"#{idx} {file.filename}")

        with col_log:
            with st.expander("📋", expanded=False):
                st.text(file.log)

        with col_btn:
            st.download_button(
                label="⬇️",
                data=file.data,
                file_name=file.filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{idx}",
                disabled=not is_valid,
                use_container_width=True
            )

else:
    st.info("💡 上传文件、添加规则后点击'开始替换'", icon="ℹ️")

# ==================== 底部帮助 ====================
st.markdown("---")

with st.expander("❓ 帮助", expanded=False):
    col_help1, col_help2 = st.columns(2, gap="medium")

    with col_help1:
        st.markdown("""
        **快速开始**
        1. 上传Word和Excel文件
        2. 添加替换规则
        3. 点击"开始替换"
        4. 下载结果

        **支持格式**
        • Word：.docx
        • Excel：.xlsx/.xls
        • 括号：【】（）()〔〕
        """)

    with col_help2:
        st.markdown("""
        **常见问题**
        • .doc文件：转换为.docx
        • 格式丢失：工具保留所有格式
        • 大文件：建议分批处理
        • 错误恢复：使用"撤销"功能
        """)

st.caption(f"v{VERSION} © 2024 Word+Excel批量替换工具")